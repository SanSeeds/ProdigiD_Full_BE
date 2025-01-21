import io
import os
import random
import tempfile
import threading
from threading import Thread
from deep_translator import GoogleTranslator
from django.shortcuts import get_object_or_404
from django.shortcuts import render
from django.contrib.auth import authenticate, login, logout
from django.http import FileResponse, HttpResponse, JsonResponse
from gtts import gTTS
from .email_llama3 import BHASHINI_API_KEY, BHASHINI_USER_ID, ask_question_chatbot, fetch_single_image, generate_blog, extract_document_content, generate_email, bhashini_translate,generate_bus_pro, generate_offer_letter, generate_summary, generate_content, generate_sales_script, get_templates, rephrasely, translate_multiple_texts, translate_with_retry, update_presentation_with_generated_content  
from django.contrib.auth.models import User
from django.utils.crypto import get_random_string
from django.utils import timezone
from rest_framework_api_key.permissions import HasAPIKey
from django.conf import settings
from datetime import date, datetime, timedelta
from .models import Cart, EmailVerificationOTP, GuestLogin, GuestWordsCount, PasswordResetRequest, Payment, Profile, TemporaryEmailVerificationOTP, UserPasswordHistory, UserService, UserSession, YearlyCart
from django.core.mail import send_mail
from django.contrib.auth import update_session_auth_hash
from django.views.decorators.csrf import csrf_exempt
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import authenticate
from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated 
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad, unpad
import base64
from django.contrib.auth.hashers import make_password  # Import the function
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.models import User
from django.core.exceptions import ValidationError
from rest_framework import status
from rest_framework.renderers import BaseRenderer
from pptx import Presentation
from django.utils.dateparse import parse_date
from django.shortcuts import render
import logging
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from datetime import timedelta
from django.utils import timezone
from dateutil.relativedelta import relativedelta
from django.conf import settings
from django.utils import timezone
from datetime import timedelta
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
import logging
from django.core.exceptions import ValidationError
from django.core.validators import validate_email
from django.contrib.auth.hashers import make_password
from django.core.mail import send_mail
from django.utils import timezone
import json
import razorpay
from io import BytesIO
from django.core.mail import EmailMessage
from reportlab.lib.pagesizes import letter
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.dateparse import parse_date
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
import json
from django.http import JsonResponse, HttpResponse
import json
from django.views.decorators.csrf import csrf_exempt
from .models import Payment
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from django.views.decorators.http import require_GET, require_http_methods , require_POST
# import googletrans 

logger = logging.getLogger(__name__)


def test_report(request):
    report_path = os.path.join(os.path.dirname(__file__), 'report.html')
    with open(report_path, 'r') as file:
        report_content = file.read()
    return render(request, 'test_report.html', {'report_content': report_content})


# Base64-encoded AES IV and Secret Key
AES_IV_b64 = settings.AES_IV
AES_SECRET_KEY_b64 = settings.AES_SECRET_KEY
ENCRYPTION_IV_b64 = settings.ENCRYPTION_IV
ENCRYPTION_SECRET_KEY_b64 = settings.ENCRYPTION_SECRET_KEY

# Decode Base64 strings to bytes
AES_IV = base64.b64decode(AES_IV_b64)
AES_SECRET_KEY = base64.b64decode(AES_SECRET_KEY_b64)
ENCRYPTION_IV = base64.b64decode(ENCRYPTION_IV_b64)
ENCRYPTION_SECRET_KEY = base64.b64decode(ENCRYPTION_SECRET_KEY_b64)

# Decode Base64 strings to bytes
AES_IV = base64.b64decode(AES_IV_b64)
AES_SECRET_KEY = base64.b64decode(AES_SECRET_KEY_b64)


# Ensure IV is 16 bytes long (128 bits)
if len(AES_IV) != 16:
    raise ValueError("AES IV must be 16 bytes long")

razorpay_client = razorpay.Client(auth=(settings.RAZORPAY_KEY_ID, settings.RAZORPAY_KEY_SECRET))


class CustomAesRenderer(BaseRenderer):
    media_type = 'application/octet-stream'
    format = 'aes'

    def render(self, data, media_type=None, renderer_context=None):
        plaintext = json.dumps(data)
        padded_plaintext = pad(plaintext.encode(), 16)
        cipher = AES.new(AES_SECRET_KEY, AES.MODE_CBC, AES_IV)
        ciphertext = cipher.encrypt(padded_plaintext)
        ciphertext_b64 = base64.b64encode(ciphertext).decode()
        response = {'ciphertext': ciphertext_b64}
        return json.dumps(response)

def landing(request):
    return render(request, 'landing.html')

def about(request):
    return render(request, 'about.html')

def invoice(request):
    return render(request, 'invoice_template.html')


def guest_generate_otp():
    """Generate a 6-digit random OTP."""
    return str(random.randint(100000, 999999))



def guest_otp_expiry_time():
    """Set OTP validity to 30 minutes."""
    return timezone.now() + timedelta(minutes=30)



@csrf_exempt
@permission_classes([HasAPIKey])
@require_POST
def create_razorpay_order(request):
    if request.method == "POST":
        try:
            # Get encrypted content from the request and decrypt it
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({"error": "No encrypted content found in the request."}, status=400)
            
            # Decrypt the content to obtain original data
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            
            # Extract amount and email
            amount = data.get('amount', 0)  # Get the amount in rupees
            print(amount)
            email = data.get('email')  # Extract email from the request
            
            if not email:
                return JsonResponse({"error": "Email is required"}, status=400)
            
            # Convert amount to paise
            amount_in_paise = int(amount * 100)

            # Create Razorpay order
            razorpay_order = razorpay_client.order.create({
                "amount": amount_in_paise,  # Use amount in paise
                "currency": "INR",
                "payment_capture": "1"
            })

            # Save order details to the Payment table including the email
            Payment.objects.create(
                order_id=razorpay_order['id'],
                amount=amount,  # Store amount in rupees in the database
                currency="INR",
                payment_capture=True,
                email=email  # Store the email
            )

            # Encrypt the response data before returning it
            response_data = {
                "order_id": razorpay_order['id'],
                "amount": amount,  # Return amount in rupees for response
                "currency": "INR",
                "razorpay_key_id": settings.RAZORPAY_KEY_ID
            }
            encrypted_response = encrypt_data(response_data)
            
            return JsonResponse({"encrypted_content": encrypted_response}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format."}, status=400)
        except ValueError as e:
            return JsonResponse({"error": str(e)}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)

@csrf_exempt
@require_POST
def verify_payment(request):
    try:
        # Extract and decrypt the encrypted content from the request
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        
        razorpay_order_id = data.get('razorpay_order_id')
        razorpay_payment_id = data.get('razorpay_payment_id')
        razorpay_signature = data.get('razorpay_signature')
        selected_services = data.get('selected_services')
        email = data.get('email')
        
        logger.info(f"Received payment verification request with order_id: {razorpay_order_id}, payment_id: {razorpay_payment_id}, signature: {razorpay_signature}")

        # Verify payment signature
        params_dict = {
            'razorpay_order_id': razorpay_order_id,
            'razorpay_payment_id': razorpay_payment_id,
            'razorpay_signature': razorpay_signature
        }

        try:
            # Verify the payment signature
            razorpay_client.utility.verify_payment_signature(params_dict)
            logger.info("Payment signature verification successful")

            # Fetch and update the payment record
            payment = Payment.objects.get(order_id=razorpay_order_id)
            payment.payment_id = razorpay_payment_id
            payment.signature = razorpay_signature
            payment.email = email
            payment.verified = True

            # Process selected services
            if not selected_services or not email:
                return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

            user = get_object_or_404(User, email=email)
            user_services, created = UserService.objects.get_or_create(user=user)
            subscribed_services = []

            # Check if "Introductory Offer" is selected
            if selected_services.get("introductory_offer_service", False):
                user_services.email_service = user_services.offer_letter_service = 1
                user_services.business_proposal_service = user_services.sales_script_service = 1
                user_services.content_generation_service = user_services.summarize_service = 1
                user_services.ppt_generation_service = user_services.blog_generation_service = 1
                user_services.rephrasely_service = 1
                subscribed_services = [
                    "Email Service", "Offer Letter Service", "Business Proposal Service",
                    "Sales Script Service", "Content Generation Service", "Summarize Service",
                    "PPT Generation Service", "Blog Generation Service", "Rephrasely Service"
                ]
            else:
                # Update services based on data
                for service_key, service_attr in [
                    ("email_service", "Email Service"), ("offer_letter_service", "Offer Letter Service"),
                    ("business_proposal_service", "Business Proposal Service"), ("sales_script_service", "Sales Script Service"),
                    ("content_generation_service", "Content Generation Service"), ("summarize_service", "Summarize Service"),
                    ("ppt_generation_service", "PPT Generation Service"), ("blog_generation_service", "Blog Generation Service"),
                    ("rephrasely_service", "Rephrasely Service")
                ]:
                    if selected_services.get(service_key, 0) > 0:
                        setattr(user_services, service_key, 1)
                        subscribed_services.append(service_attr)

            # Save the updated services
            user_services.save()
            payment.order_datetime = datetime.now()
            payment.subscribed_services = selected_services
            payment.service = user_services
            payment.subscription_duration = 'monthly'
            payment.save()

            # Send subscription confirmation email
            subject = 'Subscription Confirmation - ProdigiDesk Services'
            services_list = ''.join(f"<li>{service}</li>" for service in subscribed_services)
            message = f"""
            <html><body>
            <p>Dear {user.get_full_name()},</p>
            <p>We are pleased to inform you that your purchase has been successfully processed. Below is a summary of the services activated for you:</p>
            <p><strong>Activated Services for a month:</strong></p>
            <ul>{services_list}</ul>
            <p><strong>Order Details:</strong></p>
            <ul>
                <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                <li><strong>Order Date and Time:</strong> {payment.order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                <li><strong>Registered Email:</strong> {email}</li>
            </ul>
            <p>Should you have any queries or require assistance, feel free to contact us at contact@espritanalytique.com. We look forward to serving you!</p>
            <p>Thank you for choosing ProdigiDesk.</p>
            <br>Best regards
            <br>The ProdigiDesk Team
            <br><a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>   
            </body></html>
            """

            email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
            email_message.content_subtype = 'html'
            email_message.send(fail_silently=False)



            # Encrypt the success response
            response_data = {'message': 'Payment and service save successful'}
            encrypted_response = encrypt_data(response_data)
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except razorpay.errors.SignatureVerificationError:
            logger.error("Payment signature verification failed")
            encrypted_response = encrypt_data({'error': 'Payment verification failed'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format")
        encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
        return JsonResponse({'encrypted_content': encrypted_response}, status=400)
    except Exception as e:
        logger.error(f"Exception occurred: {str(e)}")
        encrypted_response = encrypt_data({'error': str(e)})
        return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=400)


@csrf_exempt
@require_POST
def verify_payment_yearly(request):
    if request.method == "POST":
        try:
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
            
            # Decrypt request content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            # Extract payment details
            razorpay_order_id = data.get('razorpay_order_id')
            razorpay_payment_id = data.get('razorpay_payment_id')
            razorpay_signature = data.get('razorpay_signature')
            selected_services = data.get('selected_services')
            email = data.get('email')

            logger.info(f"Received yearly payment verification request with order_id: {razorpay_order_id}")

            # Verify payment signature
            params_dict = {
                'razorpay_order_id': razorpay_order_id,
                'razorpay_payment_id': razorpay_payment_id,
                'razorpay_signature': razorpay_signature
            }

            try:
                razorpay_client.utility.verify_payment_signature(params_dict)
                logger.info("Payment signature verification successful for yearly subscription")

                # Update payment and user service details
                payment = Payment.objects.get(order_id=razorpay_order_id)
                payment.payment_id = razorpay_payment_id
                payment.signature = razorpay_signature
                payment.email = email
                payment.verified = True

                if not selected_services or not email:
                    return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

                user = get_object_or_404(User, email=email)
                user_services, created = UserService.objects.get_or_create(user=user)
                subscribed_services = []
                expiration_date = timezone.now().date() + relativedelta(years=1)  # Default expiration date set to 1 year

                # Check if the introductory offer is selected
                if selected_services.get("introductory_offer_service", False):
                    # Activate multiple services as part of the introductory offer
                    services = [
                        "email_service", "offer_letter_service", "business_proposal_service",
                        "sales_script_service", "content_generation_service", "summarize_service",
                        "ppt_generation_service", "blog_generation_service", "rephrasely_service"
                    ]
                    for service in services:
                        setattr(user_services, service, 1)
                        setattr(user_services, f"{service.replace('_service', '')}_end_date", expiration_date)
                    subscribed_services = [service.replace("_service", "").replace("_", " ").title() for service in services]
                else:
                    # Activate individual services based on selection
                    for service, end_date_field in [
                        ("email_service", "email_end_date"),
                        ("offer_letter_service", "offer_letter_end_date"),
                        ("business_proposal_service", "business_proposal_end_date"),
                        ("sales_script_service", "sales_script_end_date"),
                        ("content_generation_service", "content_generation_end_date"),
                        ("summarize_service", "summarize_end_date"),
                        ("ppt_generation_service", "ppt_generation_end_date"),
                        ("blog_generation_service", "blog_generation_end_date"),
                        ("rephrasely_service", "rephrasely_end_date"),
                    ]:
                        if selected_services.get(service, 0) > 0:
                            setattr(user_services, service, 1)
                            setattr(user_services, end_date_field, expiration_date)
                            subscribed_services.append(service.replace("_service", "").replace("_", " ").title())

                user_services.save()
                payment.order_datetime = datetime.now()
                payment.subscribed_services = selected_services
                payment.subscription_duration = 'yearly'  # Set subscription_duration as yearly
                payment.service = user_services
                payment.save()

                # Send subscription confirmation email
                subject = 'Subscription Confirmation - ProdigiDesk Services'
                services_list = ''.join(f"<li>{service}</li>" for service in subscribed_services)
                message = f"""
                <html><body>
                <p>Dear {user.get_full_name()},</p>
                <p>We are pleased to inform you that your subscription has been successfully processed. Below is a summary of the services activated for you:</p>
                <p><strong>Activated Services for a year:</strong></p>
                <ul>{services_list}</ul>
                <p><strong>Order Details:</strong></p>
                <ul>
                    <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                    <li><strong>Order Date and Time:</strong> {payment.order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                    <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                    <li><strong>Registered Email:</strong> {email}</li>
                </ul>
                <p>Should you have any queries or require assistance, feel free to contact us at contact@espritanalytique.com. We look forward to serving you!</p>
                <p>Thank you for choosing ProdigiDesk.</p>
                <br>
                <br>Best regards
                <br>The ProdigiDesk Team
                <br><a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>   
                
                </body></html>
                """

                email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
                email_message.content_subtype = 'html'
                email_message.send(fail_silently=False)

                # Prepare success response with encryption
                response_content = {
                    'message': 'Yearly payment and service save successful',
                    'subscribed_services': subscribed_services,
                    'subscription_duration': 'yearly'
                }
                encrypted_response = encrypt_data(response_content)

                return JsonResponse({'encrypted_content': encrypted_response}, status=200)

            except razorpay.errors.SignatureVerificationError:
                logger.error("Yearly payment signature verification failed")
                return JsonResponse({"status": "Payment verification failed"}, status=400)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format")
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            logger.error(f"Exception occurred: {str(e)}")
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)



from dateutil.relativedelta import relativedelta

@csrf_exempt
@require_POST
def extend_service(request):
    if request.method == "POST":
        try:
            # Retrieve and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the encrypted content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f"Decrypted content: {data}")

            # Extract relevant fields
            email = data.get('email')
            selected_services = data.get('selected_services')  
            razorpay_order_id = data.get('razorpay_order_id')
            razorpay_payment_id = data.get('razorpay_payment_id')
            razorpay_signature = data.get('razorpay_signature')

            # Logging for debugging purposes
            if not selected_services or not email:
                logger.error(f'Missing fields - selected_services: {selected_services}, email: {email}')
                return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

            logger.info(f"Received service extension request with order_id: {razorpay_order_id}, payment_id: {razorpay_payment_id}, signature: {razorpay_signature}")

            # Verify payment signature
            params_dict = {
                'razorpay_order_id': razorpay_order_id,
                'razorpay_payment_id': razorpay_payment_id,
                'razorpay_signature': razorpay_signature
            }

            try:
                # Verify the payment signature using Razorpay's utility function
                razorpay_client.utility.verify_payment_signature(params_dict)
                logger.info("Payment signature verification successful")

                # Update the Payment record in the database
                payment = Payment.objects.get(order_id=razorpay_order_id)
                payment.payment_id = razorpay_payment_id
                payment.signature = razorpay_signature
                payment.email = email
                payment.verified = True  # Mark the payment as verified
                payment.subscription_duration = 'monthly' 

                # Find the user and user services
                user = get_object_or_404(User, email=email)
                user_services, created = UserService.objects.get_or_create(user=user)

                extended_services = []
                current_date = timezone.now().date()
                new_expiry_date = current_date + relativedelta(months=1)

                # Extend each service by one month
                for service_key, service_value in selected_services.items():
                    if service_value:
                        if service_key == "email_service":
                            user_services.email_service = True
                            user_services.email_end_date = new_expiry_date if user_services.email_end_date is None else user_services.email_end_date + relativedelta(months=1)
                            extended_services.append("Email Service")
                        elif service_key == "offer_letter_service":
                            user_services.offer_letter_service = True
                            user_services.offer_letter_end_date = new_expiry_date if user_services.offer_letter_end_date is None else user_services.offer_letter_end_date + relativedelta(months=1)
                            extended_services.append("Offer Letter Service")
                        elif service_key == "business_proposal_service":
                            user_services.business_proposal_service = True
                            user_services.business_proposal_end_date = new_expiry_date if user_services.business_proposal_end_date is None else user_services.business_proposal_end_date + relativedelta(months=1)
                            extended_services.append("Business Proposal Service")
                        elif service_key == "sales_script_service":
                            user_services.sales_script_service = True
                            user_services.sales_script_end_date = new_expiry_date if user_services.sales_script_end_date is None else user_services.sales_script_end_date + relativedelta(months=1)
                            extended_services.append("Sales Script Service")
                        elif service_key == "content_generation_service":
                            user_services.content_generation_service = True
                            user_services.content_generation_end_date = new_expiry_date if user_services.content_generation_end_date is None else user_services.content_generation_end_date + relativedelta(months=1)
                            extended_services.append("Content Generation Service")
                        elif service_key == "summarize_service":
                            user_services.summarize_service = True
                            user_services.summarize_end_date = new_expiry_date if user_services.summarize_end_date is None else user_services.summarize_end_date + relativedelta(months=1)
                            extended_services.append("Summarize Service")
                        elif service_key == "ppt_generation_service":
                            user_services.ppt_generation_service = True
                            user_services.ppt_generation_end_date = new_expiry_date if user_services.ppt_generation_end_date is None else user_services.ppt_generation_end_date + relativedelta(months=1)
                            extended_services.append("PPT Generation Service")
                        elif service_key == "blog_generation_service":
                            user_services.blog_generation_service = True
                            user_services.blog_generation_end_date = new_expiry_date if user_services.blog_generation_end_date is None else user_services.blog_generation_end_date + relativedelta(months=1)
                            extended_services.append("Blog Generation Service")
                        elif service_key == "rephrasely_service":
                            user_services.rephrasely_service = True
                            user_services.rephrasely_end_date = new_expiry_date if user_services.rephrasely_end_date is None else user_services.rephrasely_end_date + relativedelta(months=1)
                            extended_services.append("Rephrasely Service")

                user_services.save()

                # Update the Payment record with the new services
                order_datetime = datetime.now()
                payment.order_datetime = order_datetime
                payment.subscribed_services = list(selected_services.keys())
                payment.service = user_services
                payment.save()

                # Send confirmation email
                subject = 'Service Extension Confirmation - ProdigiDesk Services'
                services_list = ''.join([f"<li>{service}</li>" for service in extended_services])
                message = f"""
                <html>
                <body>
                <p>Dear {user.get_full_name()},</p>
                <p>Your monthly service extension to ProdigiDesk has been successfully processed.</p>
                <p>The following services have been extended:</p>
                <ul>{services_list}</ul>
                <p><strong>Order Details:</strong></p>
                <ul>
                    <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                    <li><strong>Order Date and Time:</strong> {order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                    <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                    <li><strong>Registered Email:</strong> {email}</li>
                </ul>
                <p>To see more details of the transaction and to get the invoice, click <a href="https://prodigidesk.ai/userSummary">here</a>.</p>
                <p>Thank you for choosing us. We look forward to supporting you further.</p>
                <br>Best regards,
                <br>The ProdigiDesk Team
                <br>contact@espritanalytique.com
                <br><a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>               
                  </body>
                </html>
                """

                email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
                email_message.content_subtype = 'html'
                email_message.send()

                response_data = {'message': 'Payment and service extension save successful'}
                encrypted_response = encrypt_data(response_data)
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)

            except razorpay.errors.SignatureVerificationError:
                logger.error("Payment signature verification failed")
                response_data = {"status": "Payment verification failed"}
                return JsonResponse({'encrypted_content': encrypt_data(response_data)}, status=400)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format")
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            logger.error(f"Exception occurred: {str(e)}")
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)


@csrf_exempt
@require_POST
def extend_service_yearly(request):
    if request.method == "POST":
        try:
            # Retrieve and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the encrypted content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f"Decrypted content: {data}")

            # Extract relevant fields
            razorpay_order_id = data.get('razorpay_order_id')
            razorpay_payment_id = data.get('razorpay_payment_id')
            razorpay_signature = data.get('razorpay_signature')
            service_ids = data.get('service_ids')
            email = data.get('email')

            logger.info(f"Received service extension request with order_id: {razorpay_order_id}, payment_id: {razorpay_payment_id}, signature: {razorpay_signature}")

            # Verify payment signature
            params_dict = {
                'razorpay_order_id': razorpay_order_id,
                'razorpay_payment_id': razorpay_payment_id,
                'razorpay_signature': razorpay_signature
            }

            try:
                # Verify the payment signature using Razorpay's utility function
                razorpay_client.utility.verify_payment_signature(params_dict)
                logger.info("Payment signature verification successful")

                # Update the Payment record in the database
                payment = Payment.objects.get(order_id=razorpay_order_id)
                payment.payment_id = razorpay_payment_id
                payment.signature = razorpay_signature
                payment.email = email
                payment.verified = True  # Mark the payment as verified
                payment.subscription_duration = 'yearly'

                if not service_ids or not email:
                    logger.error('No services or email found in the request.')
                    return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

                # Fetch user and initialize user services
                user = get_object_or_404(User, email=email)
                user_services, created = UserService.objects.get_or_create(user=user)

                extended_services = []
                current_date = timezone.now().date()
                new_expiry_date = current_date + relativedelta(years=1)

                # Extend each service by one year
                for service_key in service_ids:
                    if service_key == "email_service":
                        user_services.email_service = True
                        user_services.email_end_date = new_expiry_date if user_services.email_end_date is None else user_services.email_end_date + relativedelta(years=1)
                        extended_services.append("Email Service")
                    elif service_key == "offer_letter_service":
                        user_services.offer_letter_service = True
                        user_services.offer_letter_end_date = new_expiry_date if user_services.offer_letter_end_date is None else user_services.offer_letter_end_date + relativedelta(years=1)
                        extended_services.append("Offer Letter Service")
                    elif service_key == "business_proposal_service":
                        user_services.business_proposal_service = True
                        user_services.business_proposal_end_date = new_expiry_date if user_services.business_proposal_end_date is None else user_services.business_proposal_end_date + relativedelta(years=1)
                        extended_services.append("Business Proposal Service")
                    elif service_key == "sales_script_service":
                        user_services.sales_script_service = True
                        user_services.sales_script_end_date = new_expiry_date if user_services.sales_script_end_date is None else user_services.sales_script_end_date + relativedelta(years=1)
                        extended_services.append("Sales Script Service")
                    elif service_key == "content_generation_service":
                        user_services.content_generation_service = True
                        user_services.content_generation_end_date = new_expiry_date if user_services.content_generation_end_date is None else user_services.content_generation_end_date + relativedelta(years=1)
                        extended_services.append("Content Generation Service")
                    elif service_key == "summarize_service":
                        user_services.summarize_service = True
                        user_services.summarize_end_date = new_expiry_date if user_services.summarize_end_date is None else user_services.summarize_end_date + relativedelta(years=1)
                        extended_services.append("Summarize Service")
                    elif service_key == "ppt_generation_service":
                        user_services.ppt_generation_service = True
                        user_services.ppt_generation_end_date = new_expiry_date if user_services.ppt_generation_end_date is None else user_services.ppt_generation_end_date + relativedelta(years=1)
                        extended_services.append("PPT Generation Service")
                    elif service_key == "blog_generation_service":
                        user_services.blog_generation_service = True
                        user_services.blog_generation_end_date = new_expiry_date if user_services.blog_generation_end_date is None else user_services.blog_generation_end_date + relativedelta(years=1)
                        extended_services.append("Blog Generation Service")
                    elif service_key == "rephrasely_service":
                        user_services.rephrasely_service = True
                        user_services.rephrasely_end_date = new_expiry_date if user_services.rephrasely_end_date is None else user_services.rephrasely_end_date + relativedelta(years=1)
                        extended_services.append("Rephrasely Service")

                user_services.save()

                # Update the Payment record with the new services
                order_datetime = datetime.now()
                payment.order_datetime = order_datetime
                payment.subscribed_services = service_ids
                payment.service = user_services
                payment.save()

                # Send confirmation email
                subject = 'Service Extension Confirmation - ProdigiDesk Services'
                services_list = ''.join([f"<li>{service}</li>" for service in extended_services])
                message = f"""
                <html>
                <body>
                <p>Dear {user.get_full_name()},</p>
                <p>Your yearly service extension to ProdigiDesk has been successfully processed.</p>
                <p>The following services have been extended:</p>
                <ul>{services_list}</ul>
                <p>Order Details:</p>
                <ul>
                    <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                    <li><strong>Order Date and Time:</strong> {payment.order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                    <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                    <li><strong>Registered Email:</strong> {email}</li>
                </ul>
                <p>To see more details of the transaction and to get the invoice, click <a href="https://prodigidesk.ai/userSummary">here</a>.</p>
                <p>Thank you for choosing us. We look forward to supporting you further.</p>
                <p>Best regards,<br>
                The ProdigiDesk Team<br>
                contact@espritanalytique.com<br>
                <a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>
                </p>
                </body>
                </html>
                """

                email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
                email_message.content_subtype = 'html'
                email_message.send()

                response_data = {'message': 'Payment and service extension save successful'}
                encrypted_response = encrypt_data(response_data)
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)

            except razorpay.errors.SignatureVerificationError:
                logger.error("Payment signature verification failed")
                response_data = {"status": "Payment verification failed"}
                return JsonResponse({'encrypted_content': encrypt_data(response_data)}, status=400)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format")
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            logger.error(f"Exception occurred: {str(e)}")
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)

@csrf_exempt
@require_POST
def generate_invoice(request):
    if request.method == "POST":
        data = json.loads(request.body)
        order_id = data.get('order_id')
        payment_id = data.get('payment_id')
        email = data.get('email')

        try:
            # Fetch the payment record based on order_id and payment_id
            payment = Payment.objects.get(order_id=order_id, payment_id=payment_id)
            buffer = io.BytesIO()

            # Set up the PDF document using SimpleDocTemplate
            doc = SimpleDocTemplate(buffer, pagesize=letter)

            # Use ReportLab's styles
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            title_style.alignment = TA_CENTER
            normal_style = styles['Normal']

            # Title
            elements = []
            elements.append(Paragraph("ProdigiDesk Invoice", title_style))
            elements.append(Spacer(1, 12))  # Adds space after title

            # Company info or logo
            elements.append(Paragraph("Esprit Analltique", normal_style))
            elements.append(Paragraph("NO 4B, CENTURION AVENUE 3RD CROSS STREET, Ayanambakkam, Poonamallee, Tiruvallur-600095", normal_style))
            elements.append(Spacer(1, 12))

            # Invoice Info
            invoice_data = [
                ["Order ID:", order_id],
                ["Payment ID:", payment_id],
                ["Amount:", f"{payment.amount} {payment.currency}"],
                ["Email:", email],
                ["Payment Status:", "Success" if payment.verified else "Failed"],
                ["Payment Capture:", 'Yes' if payment.payment_capture else 'No'],
                ["Order Date:", payment.order_datetime.strftime('%Y-%m-%d %H:%M:%S') if payment.order_datetime else 'N/A']
            ]
            
            table = Table(invoice_data, colWidths=[150, 300])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)
            elements.append(Spacer(1, 12))

            # Only include names of services that have a value of True
            if payment.subscribed_services:
                elements.append(Paragraph("Purchased Services:", normal_style))
                for service_name, subscribed in payment.subscribed_services.items():
                    if subscribed:
                        elements.append(Paragraph(f"- {service_name}", normal_style))

            # Build the document
            doc.build(elements)

            # Return the PDF file as response
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename=invoice_{order_id}.pdf'
            return response

        except Payment.DoesNotExist:
            return JsonResponse({"error": "Payment not found"}, status=404)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)

@csrf_exempt
@require_GET
def invoice_details(request):
    if request.method == "GET":
        payment_id = request.GET.get('payment_id')  # Fetch payment_id from query parameters

        if not payment_id:
            return JsonResponse({"error": "Payment ID is required"}, status=400)

        try:
            payment = Payment.objects.get(payment_id=payment_id)

            payment_details = {
                "order_id": payment.order_id,
                "payment_id": payment.payment_id,
                "signature": payment.signature,
                "email": payment.email,
                "amount": str(payment.amount),  # Convert Decimal to string for JSON serialization
                "currency": payment.currency,
                "payment_capture": payment.payment_capture,
                "created_at": payment.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                "verified": payment.verified,
                "order_datetime": payment.order_datetime.strftime('%Y-%m-%d %H:%M:%S') if payment.order_datetime else None,
                "subscribed_services": payment.subscribed_services,
                # Include any other fields you want to return
            }

            return JsonResponse(payment_details, status=200)

        except Payment.DoesNotExist:
            return JsonResponse({"error": "Payment not found"}, status=404)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)

def encrypt_data(data):
    plaintext = json.dumps(data)
    padded_plaintext = pad(plaintext.encode(), 16)
    cipher = AES.new(ENCRYPTION_SECRET_KEY, AES.MODE_CBC, ENCRYPTION_IV)
    ciphertext = cipher.encrypt(padded_plaintext)
    ciphertext_b64 = base64.b64encode(ciphertext).decode()
    return ciphertext_b64


def decrypt_data(encrypted_data):
    try:
        encrypted_data_bytes = base64.b64decode(encrypted_data)
        cipher = AES.new(ENCRYPTION_SECRET_KEY, AES.MODE_CBC, ENCRYPTION_IV)
        decrypted_data = unpad(cipher.decrypt(encrypted_data_bytes), 16)
        return decrypted_data.decode('utf-8')
    except Exception as e:
        raise ValueError(f"Decryption error: {e}")


from django.contrib.auth.hashers import make_password, check_password

@csrf_exempt
@permission_classes([HasAPIKey])
@require_POST
def add_user(request):
    if request.method == 'POST':
        try:
            # Load untrusted domains from the text file
            with open('./domains.txt', 'r') as file:
                untrusted_domains = {line.strip().lower() for line in file}

            # Load and decode the request body
            body = request.body.decode('utf-8')
            logger.debug(f"Request body received: {body}")

            # Extract and decrypt the incoming payload
            data = json.loads(body)
            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                logger.warning("No encrypted content found in the request.")
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            logger.debug(f"Encrypted content received: {encrypted_content}")
            decrypted_content = decrypt_data(encrypted_content)  # Decrypt the content using your custom decrypt_data function
            logger.debug(f"Decrypted content: {decrypted_content}")

            content = json.loads(decrypted_content)

            if not content:
                logger.warning('No content found in the request.')
                return JsonResponse({'error': 'No content found in the request.'}, status=400)

            # Extract required fields
            first_name = content.get('first_name')
            last_name = content.get('last_name')
            username = content.get('username')
            email = content.get('email')
            password = content.get('password')
            confirm_password = content.get('confirm_password')
            state = content.get('state')  # Extract the state field

            # Check if username and email are provided
            if not username:
                return JsonResponse({'error': 'Username is required.'}, status=400)
            if not email:
                return JsonResponse({'error': 'Email is required.'}, status=400)

            # Normalize username and email to lowercase
            username = username.lower()
            email = email.lower()

            # Extract domain from the email
            try:
                email_domain = email.split('@')[1].lower()
            except IndexError:
                return JsonResponse({'error': 'Invalid email format.'}, status=400)

            # Check if the email domain is in the untrusted list
            if email_domain in untrusted_domains:
                return JsonResponse({
                    'error': 'It seems you are using an untrusted email domain service. Please try with another email.'}, 
                    status=400)

            # Check if passwords match
            if password != confirm_password:
                return JsonResponse({'error': 'Passwords do not match.'}, status=400)

            # Check if the password is the same as the email
            if password == email:
                return JsonResponse({'error': 'Password cannot be the same as the email address.'}, status=400)

            if password == username:
                return JsonResponse({'error': 'Password cannot be the same as the username.'}, status=400)

            # Check if username already exists
            if User.objects.filter(username=username).exists():
                return JsonResponse({'error': 'Username already exists.'}, status=400)

            # Validate email
            try:
                validate_email(email)
            except ValidationError:
                return JsonResponse({'error': 'Invalid email address.'}, status=400)

            # Check if email already exists
            if User.objects.filter(email=email).exists():
                user = User.objects.get(email=email)
                # Check if the password matches any of the last 5 passwords for this email
                recent_passwords = UserPasswordHistory.get_recent_passwords(user)
                for recent_password in recent_passwords:
                    if check_password(password, recent_password.hashed_password):  # Check if the password matches any of the recent passwords
                        return JsonResponse({'error': 'New password cannot be the same as any of the last 5 passwords.'}, status=400)

            # Check if state is provided
            if not state:
                return JsonResponse({'error': 'State is required.'}, status=400)

            # Create user
            user = User.objects.create(
                first_name=first_name,
                last_name=last_name,
                username=username,
                email=email,
                password=make_password(password),  # Hash the password
                state=state
            )
            user.save()

            # Store the new hashed password in the history table
            UserPasswordHistory.store_password(user, email, make_password(password))

            # Prepare the response
            response_data = {
                'message': 'User created successfully',
                'user_id': user.id,
                'email': email
            }

            # Encrypt the response content
            encrypted_response_content = encrypt_data(response_data)  # Encrypt the response using your custom encrypt_data function

            # Return the encrypted response
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=201)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format in request")
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)
    else:
        logger.error("Invalid request method")
        return JsonResponse({'error': 'Invalid request method'}, status=405)


@csrf_exempt
@permission_classes([HasAPIKey])
@require_POST
def send_email_verification_otp(request):
    if request.method == 'POST':
        try:
            # Load untrusted domains from the text file
            with open('./domains.txt', 'r') as file:
                untrusted_domains = {line.strip().lower() for line in file}

            # Load and decode the request body
            body = request.body.decode('utf-8')
            logger.debug(f"Request body received: {body}")

            # Extract and decrypt the incoming payload
            data = json.loads(body)
            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                logger.warning("No encrypted content found in the request.")
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            logger.debug(f"Encrypted content received: {encrypted_content}")
            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f"Decrypted content: {decrypted_content}")

            data = json.loads(decrypted_content)

            # Extract email from the decrypted payload
            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email is required'}, status=400)

            # Extract domain from the email
            try:
                email_domain = email.split('@')[1].lower()
            except IndexError:
                return JsonResponse({'error': 'Invalid email format.'}, status=400)

            # Check if the email domain is in the untrusted list
            if email_domain in untrusted_domains:
                return JsonResponse({
                    'error': 'It seems you are using an untrusted email domain service. Please try with another email.'},
                    status=400)

            # Check if the email is already registered
            if User.objects.filter(email=email).exists():
                return JsonResponse({'error': 'Email is already registered'}, status=400)

            # Generate OTP and set expiry time
            otp = generate_otp()
            expiry_time = timezone.now() + timedelta(minutes=10)

            # Save OTP and its expiry in the database
            TemporaryEmailVerificationOTP.objects.update_or_create(
                email=email,
                defaults={
                    'otp': otp,
                    'expiry_time': expiry_time
                }
            )

            # Send OTP via email
            subject = 'Welcome to ProdigiDesk'
            plain_message = f"""
Dear Sir/Madam,

We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is {otp} and will remain valid for a period of 10 minutes.

Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.

This notification is intended to provide you with the necessary information to complete your Email Verification. If you have any concerns or require assistance, please contact our support team through the appropriate channels.

Thank you for your understanding and cooperation.

Sincerely,
The ProdigiDesk Team
"""
            html_message = f"""
<p>Dear Sir/Madam,</p>
<p>We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is <strong>{otp}</strong> and will remain valid for a period of 10 minutes.</p>
<p>Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.</p>
<p>This notification is intended to provide you with the necessary information to complete your Email Verification. If you have any concerns or require assistance, please contact our support team through the appropriate channels.</p>
<p>Thank you for your understanding and cooperation.</p>
<p>Sincerely,<br>The ProdigiDesk Team</p>
"""

            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [email]

            try:
                send_mail(
                    subject, 
                    plain_message, 
                    from_email, 
                    recipient_list, 
                    fail_silently=False,
                    html_message=html_message  # Add HTML content here
                )
            except Exception as e:
                return JsonResponse({'error': f'Error sending email: {str(e)}'}, status=500)

            # Encrypt the response content
            encrypted_response_content = encrypt_data({'success': 'OTP sent successfully'})

            # Return the encrypted response
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format received.")
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            logger.error(f"An unexpected error occurred: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=405)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def fetch_filtered_payments(request):
    if request.method == 'POST':
        try:
            # Decrypt the request payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({"error": "No encrypted content found in the request."}, status=400)
            
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            # Extract parameters from decrypted data
            order_id = data.get('order_id')
            payment_id = data.get('payment_id')
            email = data.get('email')
            
            # Filter payments based on provided parameters
            filters = {}
            if order_id:
                filters['order_id'] = order_id
            if payment_id:
                filters['payment_id'] = payment_id
            if email:
                filters['email'] = email

            # Fetch the filtered payments
            payments = Payment.objects.filter(**filters)

            # Serialize the records into a list of dictionaries
            payment_list = []
            for payment in payments:
                payment_data = {
                    'order_id': payment.order_id,
                    'payment_id': payment.payment_id,
                    'signature': payment.signature,
                    'email': payment.email,
                    'amount': str(payment.amount),  # Converting Decimal to string
                    'currency': payment.currency,
                    'payment_capture': payment.payment_capture,
                    'created_at': payment.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                    'verified': payment.verified,
                    'order_datetime': payment.order_datetime.strftime('%Y-%m-%d %H:%M:%S') if payment.order_datetime else None,
                    'subscribed_services': payment.subscribed_services,
                    'service': payment.service.id if payment.service else None,
                    'invoice_number': payment.invoice_number if payment.invoice_number else None,
                    'subscription_duration': payment.subscription_duration
                }
                payment_list.append(payment_data)

            # Encrypt the response data
            encrypted_response = encrypt_data(payment_list)
            
            # Return the encrypted response as JSON
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except Payment.DoesNotExist:
            return JsonResponse({"error": "No matching payment found"}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)


# Backend OTP Verification API
@csrf_exempt
@require_POST
def otp_verify(request):
    if request.method == 'POST':
        try:
            # Load and decode the request body
            body = request.body.decode('utf-8')
            data = json.loads(body)

            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the incoming payload
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            # Extract OTP from the decrypted payload
            otp = data.get('otp')
            email = data.get('email')

            if not otp or not email:
                return JsonResponse({'error': 'OTP and email are required'}, status=400)

            # Fetch the OTP entry from the database
            try:
                otp_entry = TemporaryEmailVerificationOTP.objects.get(email=email)
            except TemporaryEmailVerificationOTP.DoesNotExist:
                return JsonResponse({'error': 'Incorrect OTP'}, status=400)

            # Check if the OTP is correct and not expired
            if otp_entry.otp != otp:
                return JsonResponse({'error': 'Incorrect OTP'}, status=400)

            if otp_entry.expiry_time < timezone.now():
                return JsonResponse({'error': 'OTP Expired'}, status=400)

            # OTP is correct and not expired
            # Encrypt the response content and send success
            encrypted_response_content = encrypt_data({'success': 'OTP verified successfully'})
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=405)



@csrf_exempt
@require_POST
def send_feedback(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            # print(f'Encrypted content received: {encrypted_content}')  # Debug print

            if not encrypted_content:
                # print('No encrypted content found in the request.')  # Debug print
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            # print(f'Decrypted content: {decrypted_content}')  # Debug print
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)

            # Extract feedback and userEmail from decrypted content
            feedback_text = data.get('feedback')
            user_email = data.get('userEmail')

            # print(f'Feedback Text: {feedback_text}, User Email: {user_email}')  # Debug print

            if not feedback_text or not user_email:
                # print('Feedback text or userEmail is missing.')  # Debug print
                logger.warning('Feedback text or userEmail is missing.')
                return JsonResponse({'error': 'Missing feedback text or userEmail'}, status=400)

            # Compose the email
            subject = 'New Feedback Submission'
            message = f'User Email: {user_email}\n\nFeedback:\n{feedback_text}\n\n'

            # Send the email
            # print('Sending email...')  # Debug print
            send_mail(
                subject,
                message,
                settings.EMAIL_HOST_USER, 
                ['prodigidesk@prodigidesk.ai'],  
                fail_silently=False
            )
            # print(f"Feedback email sent successfully from {user_email}")  # Debug print

            # Encrypt the response message
            encrypted_response = encrypt_data({'message': 'Feedback sent successfully'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            # print('Invalid JSON format in request')  # Debug print
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
        except Exception as e:
            # print(f"Unexpected error: {str(e)}")  # Debug print
            logger.error(f"Unexpected error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)
    else:
        # print('Invalid request method')  # Debug print
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)


@csrf_exempt
@require_POST
def save_selected_services(request):
    if request.method == "POST":
        try:
            # Decode the request body
            body = request.body.decode('utf-8')

            # Extract and decrypt the incoming payload
            data = json.loads(body)
            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            # Extract email and selected services
            email = data.get("email")
            selected_services = data.get("selected_services", {})

            # Get or create the user and user services
            user = get_object_or_404(User, email=email)
            user_services, created = UserService.objects.get_or_create(user=user)

            # Check if "Introductory Offer" is selected
            if selected_services.get("introductory_offer_service", False):
                # Set all services to 1
                user_services.email_service = 1
                user_services.offer_letter_service = 1
                user_services.business_proposal_service = 1
                user_services.sales_script_service = 1
                user_services.content_generation_service = 1
                user_services.summarize_service = 1
                user_services.ppt_generation_service = 1
                user_services.blog_generation_service = 1
                user_services.rephrasely_service = 1
            else:
                # Update services based on the data
                user_services.email_service = selected_services.get("email_service", user_services.email_service)
                user_services.offer_letter_service = selected_services.get("offer_letter_service", user_services.offer_letter_service)
                user_services.business_proposal_service = selected_services.get("business_proposal_service", user_services.business_proposal_service)
                user_services.sales_script_service = selected_services.get("sales_script_service", user_services.sales_script_service)
                user_services.content_generation_service = selected_services.get("content_generation_service", user_services.content_generation_service)
                user_services.summarize_service = selected_services.get("summarize_service", user_services.summarize_service)
                user_services.ppt_generation_service = selected_services.get("ppt_generation_service", user_services.ppt_generation_service)
                user_services.blog_generation_service = selected_services.get("blog_generation_service", user_services.blog_generation_service)
                user_services.rephrasely_service = selected_services.get("rephrasely_service", user_services.rephrasely_service)

            # Save the updated user services
            user_services.save()

            # Encrypt the success response
            encrypted_response_content = encrypt_data({'message': 'Services saved successfully'})

            # Return the encrypted response
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)


@csrf_exempt
@require_POST
def update_user_services(request, email):
    if request.method == "POST":
        # Retrieve the user and their services
        user = get_object_or_404(User, email=email)
        user_services = get_object_or_404(UserService, user=user)
        
        # Parse the JSON data from the request body
        try:
            services_data = json.loads(request.body).get("selected_services", {})
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON data"}, status=400)

        # Update the services dynamically based on the mapping from the frontend
        service_mapping = {
            "email_service": "email_service",
            "offer_letter_service": "offer_letter_service",
            "business_proposal_service": "business_proposal_service",
            "sales_script_service": "sales_script_service",
            "content_generation_service": "content_generation_service",
            "summarize_service": "summarize_service",
            "ppt_generation_service": "ppt_generation_service",
            "blog_generation_service": "blog_generation_service",  # New Service
            "rephrasely_service": "rephrasely_service",            # New Service
        }

        # Update user services based on received data
        for service_key, db_field in service_mapping.items():
            if service_key in services_data:
                setattr(user_services, db_field, services_data[service_key])

        # Save the updated user services
        user_services.save()

        return JsonResponse({"success": True, "message": "Services updated successfully."})

    return JsonResponse({"error": "Invalid request method"}, status=400)


#Encrypted API To get all user Services
@require_GET
@api_view(['GET'])
@permission_classes([IsAuthenticated, HasAPIKey])
def get_user_services(request, email):
    if request.method == "GET":
        try:
            email = email.lower()
            user = get_object_or_404(User, email=email)
            user_services = get_object_or_404(UserService, user=user)

            services = {
                "email_service": {
                    "id": 1,
                    "is_active": user_services.email_service and (user_services.email_end_date >= date.today()),
                    "end_date": user_services.email_end_date.isoformat() if user_services.email_end_date else None,
                },
                "offer_letter_service": {
                    "id": 2,
                    "is_active": user_services.offer_letter_service and (user_services.offer_letter_end_date >= date.today()),
                    "end_date": user_services.offer_letter_end_date.isoformat() if user_services.offer_letter_end_date else None,
                },
                "business_proposal_service": {
                    "id": 3,
                    "is_active": user_services.business_proposal_service and (user_services.business_proposal_end_date >= date.today()),
                    "end_date": user_services.business_proposal_end_date.isoformat() if user_services.business_proposal_end_date else None,
                },
                "sales_script_service": {
                    "id": 4,
                    "is_active": user_services.sales_script_service and (user_services.sales_script_end_date >= date.today()),
                    "end_date": user_services.sales_script_end_date.isoformat() if user_services.sales_script_end_date else None,
                },
                "content_generation_service": {
                    "id": 5,
                    "is_active": user_services.content_generation_service and (user_services.content_generation_end_date >= date.today()),
                    "end_date": user_services.content_generation_end_date.isoformat() if user_services.content_generation_end_date else None,
                },
                "summarize_service": {
                    "id": 6,
                    "is_active": user_services.summarize_service and (user_services.summarize_end_date >= date.today()),
                    "end_date": user_services.summarize_end_date.isoformat() if user_services.summarize_end_date else None,
                },
                "ppt_generation_service": {
                    "id": 7,
                    "is_active": user_services.ppt_generation_service and (user_services.ppt_generation_end_date >= date.today()),
                    "end_date": user_services.ppt_generation_end_date.isoformat() if user_services.ppt_generation_end_date else None,
                },
                "blog_generation_service": {
                    "id": 9,
                    "is_active": user_services.blog_generation_service and (user_services.blog_generation_end_date >= date.today()),
                    "end_date": user_services.blog_generation_end_date.isoformat() if user_services.blog_generation_end_date else None,
                },
                "rephrasely_service": {
                    "id": 10,
                    "is_active": user_services.rephrasely_service and (user_services.rephrasely_end_date >= date.today()),
                    "end_date": user_services.rephrasely_end_date.isoformat() if user_services.rephrasely_end_date else None,
                },
            }

            # Encrypt the response content
            encrypted_services = encrypt_data({"user_id": user.id, "services": services})

            return JsonResponse({'encrypted_content': encrypted_services}, status=200)

        except User.DoesNotExist:
            return JsonResponse({"error": "User not found"}, status=404)
        except UserService.DoesNotExist:
            return JsonResponse({"error": "User services not found"}, status=404)

    return JsonResponse({"error": "Invalid request method"}, status=400)



def generate_otp():
    return ''.join(random.choices('0123456789', k=6))


# Encrypted API to send OTP for Password Reset
@csrf_exempt
@require_POST
def send_otp(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f"Encrypted content received: {encrypted_content}")

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f"Decrypted content: {decrypted_content}")
            data = json.loads(decrypted_content)

            email = data.get('email')
            logger.debug(f"Received OTP request for email: {email}")

            try:
                user = User.objects.get(email=email)
            except User.DoesNotExist:
                logger.warning(f"Email does not exist: {email}")
                encrypted_response = encrypt_data({'error': 'Email does not exist'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=404)

            # Generate OTP
            otp = generate_otp()
            expiry_time = timezone.now() + timedelta(minutes=10)

            # Store OTP and expiry time in the database
            PasswordResetRequest.objects.update_or_create(
                user=user,
                defaults={
                    'otp': otp,
                    'expiry_time': expiry_time
                }
            )
            logger.info(f"Generated OTP for user {user.username}")

            # Send OTP via email
            subject = 'Confidential OTP for Password Reset'
            message = f"""
Dear Sir/Madam,

We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is {otp} and will remain valid for a period of 10 minutes.

Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.

This notification is intended to provide you with the necessary information to complete your Password Reset. If you have any concerns or require assistance, please contact our support team through the appropriate channels.

Thank you for your understanding and cooperation.

Sincerely,
The ProdigiDesk Team
"""

            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [email]

            try:
                send_mail(subject, message, from_email, recipient_list, fail_silently=False)
                logger.info(f"OTP email sent to {email}")
            except Exception as e:
                logger.error(f"Error sending email: {str(e)}")
                encrypted_response = encrypt_data({'error': f'Error sending email: {str(e)}'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=500)

            encrypted_response = encrypt_data({'success': 'OTP sent successfully'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format in request")
            encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            encrypted_response = encrypt_data({'error': str(e)})
            return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    logger.error("Invalid request method")
    encrypted_response = encrypt_data({'error': 'Invalid request method'})
    return JsonResponse({'encrypted_content': encrypted_response}, status=405)


# @csrf_exempt
# @api_view(['POST'])
# @permission_classes([])
# def reset_password_with_otp(request):
#     if request.method == 'POST':
#         try:
#             # Extract and decrypt the incoming payload
#             encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
#             logger.debug(f"Encrypted content received: {encrypted_content}")

#             if not encrypted_content:
#                 logger.warning('No encrypted content found in the request.')
#                 return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

#             decrypted_content = decrypt_data(encrypted_content)
#             logger.debug(f"Decrypted content: {decrypted_content}")
#             data = json.loads(decrypted_content)

#             email = data.get('email')
#             otp = data.get('otp')
#             new_password = data.get('new_password')
#             confirm_new_password = data.get('confirm_new_password')

#             logger.debug(f"Received password reset request for email: {email} with OTP: {otp}")

#             try:
#                 user = User.objects.get(email=email)
#             except User.DoesNotExist:
#                 logger.warning(f"Email does not exist: {email}")
#                 encrypted_response = encrypt_data({'error': 'Email does not exist'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=404)

#             # Verify OTP
#             try:
#                 reset_request = PasswordResetRequest.objects.get(user=user, otp=otp)
#                 if reset_request.expiry_time < timezone.now():
#                     logger.warning(f"OTP expired for user {user.username}")
#                     encrypted_response = encrypt_data({'error': 'OTP expired'})
#                     return JsonResponse({'encrypted_content': encrypted_response}, status=400)
#             except PasswordResetRequest.DoesNotExist:
#                 logger.warning(f"Invalid OTP for user {user.username}")
#                 encrypted_response = encrypt_data({'error': 'Invalid OTP'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Ensure new password is not the same as the current password
#             if user.check_password(new_password):
#                 logger.warning(f"User {user.username} tried to use the same new password as the current password.")
#                 encrypted_response = encrypt_data({'error': 'New password cannot be the same as the current password.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Check if new passwords match
#             if new_password != confirm_new_password:
#                 logger.warning(f"User {user.username} provided non-matching new passwords.")
#                 encrypted_response = encrypt_data({'error': 'New passwords do not match.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Update password
#             user.set_password(new_password)
#             user.save()
#             logger.info(f"User {user.username} successfully reset their password.")

#             # Remove the OTP request after successful password reset
#             reset_request.delete()

#             encrypted_response = encrypt_data({'success': 'Password reset successfully'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=200)

#         except json.JSONDecodeError:
#             logger.error("Invalid JSON format in request")
#             encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=400)
#         except Exception as e:
#             logger.error(f"Unexpected error: {str(e)}")
#             encrypted_response = encrypt_data({'error': str(e)})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=500)

#     logger.error("Invalid request method")
#     encrypted_response = encrypt_data({'error': 'Invalid request method'})
#     return JsonResponse({'encrypted_content': encrypted_response}, status=405)

@csrf_exempt
@api_view(['POST'])
@permission_classes([])
def reset_password_with_otp(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f"Encrypted content received: {encrypted_content}")

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f"Decrypted content: {decrypted_content}")
            data = json.loads(decrypted_content)

            email = data.get('email')
            otp = data.get('otp')
            new_password = data.get('new_password')
            confirm_new_password = data.get('confirm_new_password')

            logger.debug(f"Received password reset request for email: {email} with OTP: {otp}")

            try:
                user = User.objects.get(email=email)
            except User.DoesNotExist:
                logger.warning(f"Email does not exist: {email}")
                encrypted_response = encrypt_data({'error': 'Email does not exist'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=404)

            # Verify OTP
            try:
                reset_request = PasswordResetRequest.objects.get(user=user, otp=otp)
                if reset_request.expiry_time < timezone.now():
                    logger.warning(f"OTP expired for user {user.username}")
                    encrypted_response = encrypt_data({'error': 'OTP expired'})
                    return JsonResponse({'encrypted_content': encrypted_response}, status=400)
            except PasswordResetRequest.DoesNotExist:
                logger.warning(f"Invalid OTP for user {user.username}")
                encrypted_response = encrypt_data({'error': 'Invalid OTP'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Ensure new passwords match
            if new_password != confirm_new_password:
                logger.warning(f"User {user.username} provided non-matching new passwords.")
                encrypted_response = encrypt_data({'error': 'New passwords do not match.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Ensure new password is not the same as the current password
            if user.check_password(new_password):
                logger.warning(f"User {user.username} tried to use the same new password as the current password.")
                encrypted_response = encrypt_data({'error': 'New password cannot be the same as the current password.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Check if the new password matches any of the last 5 passwords
            recent_passwords = UserPasswordHistory.get_recent_passwords(user)
            for recent_password in recent_passwords:
                if check_password(new_password, recent_password.hashed_password):
                    logger.warning(f"User {user.username} tried to use a password from their last 5 passwords.")
                    encrypted_response = encrypt_data({'error': 'New password cannot match any of the last 5 passwords.'})
                    return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Update password
            user.set_password(new_password)
            user.save()
            logger.info(f"User {user.username} successfully reset their password.")

            # Store the new hashed password in the password history
            UserPasswordHistory.store_password(user, email, make_password(new_password))

            # Remove the OTP request after successful password reset
            reset_request.delete()

            encrypted_response = encrypt_data({'success': 'Password reset successfully'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format in request")
            encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            encrypted_response = encrypt_data({'error': str(e)})
            return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    logger.error("Invalid request method")
    encrypted_response = encrypt_data({'error': 'Invalid request method'})
    return JsonResponse({'encrypted_content': encrypted_response}, status=405)


#Encrypted API to Sign the user in
@csrf_exempt
@require_POST
def signin(request):
    if request.method == 'POST':
        try:
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
            
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            login_input = data.get('login_input').lower()
            password = data.get('password')
            logout_from_all = data.get('logout_from_all', False)  # Check if the checkbox is set

            if not login_input or not password:
                logger.warning('Login input and password are required')
                return JsonResponse({'error': 'Login input and password are required'}, status=400)

            try:
                if '@' in login_input:
                    user = User.objects.get(email=login_input)
                else:
                    user = User.objects.get(username=login_input)
            except User.DoesNotExist:
                logger.warning('Username or email not found')
                return JsonResponse({'error': 'Username or email not found'}, status=401)

            # Check if the user has an active session
            if logout_from_all:
                # Mark all previous sessions as inactive
                UserSession.objects.filter(user=user, active=True).update(active=False)

            active_session = UserSession.objects.filter(user=user, active=True).first()
            if active_session and not logout_from_all:
                logger.warning('User already logged in with an active session')
                return JsonResponse({'error': 'User already logged in'}, status=403)

            user = authenticate(request, username=user.username, password=password)
            if user is not None:
                login(request, user)
                refresh = RefreshToken.for_user(user)
                logger.info(f'User {user.username} authenticated successfully')

                # Create a new session
                session_id = get_random_string(length=32)
                UserSession.objects.create(user=user, session_id=session_id, email=user.email, active=True)  # Include email here

                logger.debug(f'Session created with ID: {session_id}')

                encrypted_response = encrypt_data({
                    'success': 'User authenticated',
                    'access': str(refresh.access_token),
                    'refresh': str(refresh),
                    'user_id': user.id,
                    'session_id': session_id  # Include session_id in the response
                })

                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            else:
                logger.warning('Password not correct')
                return JsonResponse({'error': 'Password not correct'}, status=401)
        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            return JsonResponse({'error': 'Internal server error'}, status=500)
    else:
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)



#Encrypted API to check session status of the user
@csrf_exempt
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def check_session_status(request):
    if request.method == 'POST':
        try:
            # Decrypt incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f'Encrypted content received: {encrypted_content}')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the content
            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)

            # Extract session ID from the decrypted content
            session_id = data.get('session_id')

            if not session_id:
                logger.warning('Session ID not provided')
                encrypted_response = encrypt_data({'error': 'Session ID not provided'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            session = UserSession.objects.filter(session_id=session_id).first()

            if not session:
                logger.warning('Session not found')
                encrypted_response = encrypt_data({'error': 'Session not found'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=404)

            # Log the session status
            logger.info(f'Session {session_id} status is {session.active}')

            # Encrypt the session status and return it
            encrypted_response = encrypt_data({
                'session_id': session_id,
                'active': session.active
            })
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            encrypted_response = encrypt_data({'error': 'Invalid JSON'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)

        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            encrypted_response = encrypt_data({'error': 'Internal server error'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    else:
        logger.error('Invalid request method')
        encrypted_response = encrypt_data({'error': 'Invalid request method'})
        return JsonResponse({'encrypted_content': encrypted_response}, status=405)



#Encrypted API to logout the user from all devices
@csrf_exempt
@require_POST
def logout_from_all_devices(request):
    if request.method == 'POST':
        try:
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
            
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            login_input = data.get('login_input').lower()
            password = data.get('password')
            
            if not login_input or not password:
                logger.warning('Login input and password are required')
                return JsonResponse({'error': 'Login input and password are required'}, status=400)

            try:
                if '@' in login_input:
                    user = User.objects.get(email=login_input)
                else:
                    user = User.objects.get(username=login_input)
            except User.DoesNotExist:
                logger.warning('Username or email not found')
                return JsonResponse({'error': 'Username or email not found'}, status=401)

            user = authenticate(request, username=user.username, password=password)
            if user is not None:
                # Mark all previous sessions as inactive
                UserSession.objects.filter(user=user, active=True).update(active=False)

                encrypted_response = encrypt_data({
                    'success': 'User logged out from all devices'
                })

                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            else:
                logger.warning('Password not correct')
                return JsonResponse({'error': 'Password not correct'}, status=401)
        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            return JsonResponse({'error': 'Internal server error'}, status=500)
    else:
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)


#Only Comment out this code when you want access tokens for testing (in case of postman)
# @csrf_exempt
# def signin(request):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body.decode('utf-8'))
#             login_input = data.get('login_input').lower()
#             password = data.get('password')
            
#             if not login_input or not password:
#                 logger.warning('Login input and password are required')
#                 return JsonResponse({'error': 'Login input and password are required'}, status=400)

#             try:
#                 if '@' in login_input:
#                     user = User.objects.get(email=login_input)
#                 else:
#                     user = User.objects.get(username=login_input)
#             except User.DoesNotExist:
#                 logger.warning('Username or email not found')
#                 return JsonResponse({'error': 'Username or email not found'}, status=401)

#             user = authenticate(request, username=user.username, password=password)
#             if user is not None:
#                 login(request, user)
#                 # Generate JWT tokens or similar mechanism for authentication
#                 # For example, using Django Rest Framework JWT or similar library
#                 # Here, `RefreshToken` is just a placeholder
#                 refresh = RefreshToken.for_user(user)
#                 logger.info(f'User {user.username} authenticated successfully')

#                 return JsonResponse({
#                     'success': 'User authenticated',
#                     'access': str(refresh.access_token),
#                     'refresh': str(refresh),
#                     'user_id': user.id  # Include user ID in the response
#                 }, status=200)
#             else:
#                 logger.warning('Password not correct')
#                 return JsonResponse({'error': 'Password not correct'}, status=401)
#         except json.JSONDecodeError:
#             logger.error('Invalid JSON format in request')
#             return JsonResponse({'error': 'Invalid JSON'}, status=400)
#         except Exception as e:
#             logger.error(f'Internal server error: {str(e)}')
#             return JsonResponse({'error': 'Internal server error'}, status=500)
#     else:
#         logger.error('Invalid request method')
#         return JsonResponse({'error': 'Invalid request method'}, status=405)


@csrf_exempt
@require_POST
def session_logout(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Request data: {data}')

            session_id = data.get('session_id')
            if not session_id:
                logger.warning('Session ID is required')
                return JsonResponse({'error': 'Session ID is required'}, status=400)

            session = UserSession.objects.filter(session_id=session_id, active=True).first()
            if session:
                session.active = False
                session.save()
                logger.info(f'Session {session_id} marked as inactive')
                return JsonResponse({'success': 'Logged out successfully'}, status=200)
            else:
                logger.warning(f'Session {session_id} not found or already inactive')
                return JsonResponse({'error': 'Session not found or already inactive'}, status=404)
        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            return JsonResponse({'error': 'Internal server error'}, status=500)
    else:
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)


#Encrypted API to Reset the user's password
# @csrf_exempt
# @require_POST
# def reset_password(request):
#     if request.method == 'POST':
#         try:
#             # Extract and decrypt the incoming payload
#             encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
#             logger.debug(f"Encrypted content received: {encrypted_content}")

#             if not encrypted_content:
#                 logger.warning('No encrypted content found in the request.')
#                 encrypted_response = encrypt_data({'error': 'No encrypted content found in the request.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             decrypted_content = decrypt_data(encrypted_content)
#             logger.debug(f"Decrypted content: {decrypted_content}")
#             data = json.loads(decrypted_content)
            
#             email = data.get('email')
#             otp = data.get('otp')
#             new_password = data.get('new_password')
#             confirm_password = data.get('confirm_password')

#             if not all([email, otp, new_password, confirm_password]):
#                 logger.warning('All fields are required')
#                 encrypted_response = encrypt_data({'error': 'All fields are required'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             if new_password != confirm_password:
#                 logger.warning('Passwords do not match')
#                 encrypted_response = encrypt_data({'error': 'Passwords do not match'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             try:
#                 user = User.objects.get(email=email)
#                 logger.info(f'User found: {user.username}')
#             except User.DoesNotExist:
#                 logger.warning(f'User with email {email} does not exist')
#                 encrypted_response = encrypt_data({'error': 'User with this email does not exist'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=404)

#             try:
#                 password_reset_request = PasswordResetRequest.objects.get(user=user, otp=otp)
#                 logger.info('Password reset request found')
#             except PasswordResetRequest.DoesNotExist:
#                 logger.warning('Invalid OTP')
#                 encrypted_response = encrypt_data({'error': 'Invalid OTP'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             if password_reset_request.expiry_time < timezone.now():
#                 logger.warning('OTP has expired')
#                 encrypted_response = encrypt_data({'error': 'OTP has expired'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             user.set_password(new_password)
#             user.save()
#             logger.info(f'Password for user {user.username} reset successfully')

#             password_reset_request.delete()
#             logger.info('Password reset request deleted')

#             encrypted_response = encrypt_data({'success': 'Password reset successfully'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=200)

#         except json.JSONDecodeError:
#             logger.error('Invalid JSON format in request')
#             encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=400)
#         except Exception as e:
#             logger.error(f"Unexpected error: {str(e)}")
#             encrypted_response = encrypt_data({'error': str(e)})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=500)

#     else:
#         logger.error('Invalid request method')
#         encrypted_response = encrypt_data({'error': 'Invalid request method'})
#         return JsonResponse({'encrypted_content': encrypted_response}, status=405)

@csrf_exempt
@require_POST
def reset_password(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f"Encrypted content received: {encrypted_content}")

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                encrypted_response = encrypt_data({'error': 'No encrypted content found in the request.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f"Decrypted content: {decrypted_content}")
            data = json.loads(decrypted_content)
            
            email = data.get('email')
            otp = data.get('otp')
            new_password = data.get('new_password')
            confirm_password = data.get('confirm_password')

            if not all([email, otp, new_password, confirm_password]):
                logger.warning('All fields are required')
                encrypted_response = encrypt_data({'error': 'All fields are required'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            if new_password != confirm_password:
                logger.warning('Passwords do not match')
                encrypted_response = encrypt_data({'error': 'Passwords do not match'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            try:
                user = User.objects.get(email=email)
                logger.info(f'User found: {user.username}')
            except User.DoesNotExist:
                logger.warning(f'User with email {email} does not exist')
                encrypted_response = encrypt_data({'error': 'User with this email does not exist'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=404)

            try:
                password_reset_request = PasswordResetRequest.objects.get(user=user, otp=otp)
                logger.info('Password reset request found')
            except PasswordResetRequest.DoesNotExist:
                logger.warning('Invalid OTP')
                encrypted_response = encrypt_data({'error': 'Invalid OTP'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            if password_reset_request.expiry_time < timezone.now():
                logger.warning('OTP has expired')
                encrypted_response = encrypt_data({'error': 'OTP has expired'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Check if the new password matches any of the last 5 passwords
            recent_passwords = UserPasswordHistory.get_recent_passwords(user)
            for recent_password in recent_passwords:
                if check_password(new_password, recent_password.hashed_password):  # Compare the hashed passwords
                    logger.warning('New password matches one of the last 5 passwords')
                    encrypted_response = encrypt_data({'error': 'New password cannot match any of the last 5 passwords'})
                    return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Reset the password
            user.set_password(new_password)
            user.save()
            logger.info(f'Password for user {user.username} reset successfully')

            # Store the new hashed password in the password history
            UserPasswordHistory.store_password(user, email, make_password(new_password))

            password_reset_request.delete()
            logger.info('Password reset request deleted')

            encrypted_response = encrypt_data({'success': 'Password reset successfully'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            encrypted_response = encrypt_data({'error': 'Invalid JSON format'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            encrypted_response = encrypt_data({'error': str(e)})
            return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    else:
        logger.error('Invalid request method')
        encrypted_response = encrypt_data({'error': 'Invalid request method'})
        return JsonResponse({'encrypted_content': encrypted_response}, status=405)


@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def email_generator(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f'Encrypted content received: {encrypted_content}')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=status.HTTP_400_BAD_REQUEST)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)

            # Define the fields that need language detection and translation
            fields_to_check = [
                'purpose', 'subject', 'to', 'tone', 'contextualBackground', 
                'callToAction', 'additionalDetails', 'priorityLevel', 'closingRemarks'
            ]

            # Define the Indian languages mapping
            indian_languages = {
                "English": "en",
                "Hindi": "hi",
                "Tamil": "ta",
                "Telugu": "te",
                "Marathi": "mr",
                "Kannada": "kn",
                "Bengali": "bn",
                "Odia": "or",
                "Assamese": "as",
                "Punjabi": "pa",
                "Malayalam": "ml",
                "Gujarati": "gu",
                "Urdu": "ur",
                "Sanskrit": "sa",
                "Nepali": "ne",
                "Bodo": "brx",
                "Maithili": "mai",
                "Sindhi": "sd",
                "Kashmiri": "ks",
                "Konkani": "kok",
                "Dogri": "doi",
                "Goan Konkani": "gom",
                "Santali": "sat",
            }

            # Translate non-English content, including keywords
            for field in fields_to_check:
                value = data.get(field)
                if value:
                    try:
                        # Detect language of the field value
                        detected_language, confidence = classify(value)
                        language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                        logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

                        # If detected language is not English, translate
                        if detected_language != 'en':
                            logger.info(f"Translating {field} from {language_name} to English.")
                            translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                            logger.debug(f"Translated {field}: {translated_text}")
                            data[field] = translated_text
                        else:
                            logger.info(f"{field} is already in English. No translation needed.")
                    except Exception as e:
                        logger.error(f"Error processing field {field}: {str(e)}")

            # Translate keywords
            keywords = data.get('keywords', [])
            translated_keywords = []
            for keyword in keywords:
                try:
                    detected_language, confidence = classify(keyword)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Keyword: {keyword} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

                    if detected_language != 'en':
                        logger.info(f"Translating keyword from {language_name} to English.")
                        translated_keyword = GoogleTranslator(source=detected_language, target='en').translate(keyword)
                        logger.debug(f"Translated keyword: {translated_keyword}")
                        translated_keywords.append(translated_keyword)
                    else:
                        translated_keywords.append(keyword)
                except Exception as e:
                    logger.error(f"Error translating keyword {keyword}: {str(e)}")
                    translated_keywords.append(keyword)
            data['keywords'] = translated_keywords

            # Extract data from the processed content
            purpose = data.get('purpose')
            if purpose == 'Other':
                purpose = data.get('otherPurpose')
            num_words = data.get('num_words')
            subject = data.get('subject')
            rephrase = data.get('rephraseSubject', False)
            to = data.get('to')
            tone = data.get('tone')
            contextual_background = data.get('contextualBackground')
            call_to_action = data.get('callToAction')
            if call_to_action == 'Other':
                call_to_action = data.get('otherCallToAction')
            additional_details = data.get('additionalDetails')
            priority_level = data.get('priorityLevel')
            closing_remarks = data.get('closingRemarks')

            logger.info(f'Generating email with the following data: {data}')

            generated_content = generate_email(
                purpose, num_words, subject, rephrase, to, tone, data['keywords'],
                contextual_background, call_to_action, additional_details,
                priority_level, closing_remarks
            )

            if generated_content:
                logger.info('Email content generated successfully.')
                # Encrypt the response content
                encrypted_response = encrypt_data({'generated_content': generated_content})
                logger.debug(f'Encrypted response: {encrypted_response}')

                return JsonResponse({'encrypted_content': encrypted_response})
            else:
                logger.error('Failed to generate email content.')
                return JsonResponse({'error': 'Failed to generate email content.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            logger.error(f'Error processing request: {e}')
            return JsonResponse({'error': 'An error occurred while processing the request.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def translate_content(request):
    translated_content = None
    error = None
    language = ""

    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f'Encrypted content received: {encrypted_content}')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)

            generated_content = data.get('generated_content')
            language = data.get('language')

            if not generated_content or not language:
                logger.warning('Both generated_content and language are required fields.')
                return JsonResponse({'error': 'Both generated_content and language are required fields.'}, status=400)

            logger.info(f'Translating content: {generated_content} to language: {language}')
            response = bhashini_translate(generated_content, language)

            if response["status_code"] == 200:
                translated_content = response["translated_content"]
                logger.info(f'Content translated successfully: {translated_content}')  # Log the translated content

                # Ensure translated content is properly encoded
                translated_content = translated_content.encode('utf-8').decode('utf-8')

                # Encrypt the response content
                encrypted_response = encrypt_data({
                    'generated_content': generated_content,
                    'translated_content': translated_content,
                    'selected_language': language
                })
                logger.debug(f'Encrypted response: {encrypted_response}')
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            else:
                logger.error('Translation failed with status code: {}'.format(response["status_code"]))
                return JsonResponse({'error': 'Translation failed.'}, status=500)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.warning(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Unexpected error: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.error('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

SUPPORTED_LANGUAGES = {
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Hindi": "hi",
    "Chinese": "zh-CN",
    "Italian": "it",
    "Japanese": "ja",
    "Korean": "ko",
    "Portuguese": "pt",
    "Dutch": "nl",
    "Polish": "pl",
    "Russian": "ru",
    "Indonesian": "id",
    "Turkish": "tr",

}

@require_POST
@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def translate_content_google(request):
    translated_content = None
    language_name = ""

    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f'Encrypted content received: {encrypted_content}')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)
            generated_content = data.get('generated_content')
            language_name = data.get('language')

            if not generated_content or not language_name:
                logger.warning('Both generated_content and language are required fields.')
                return JsonResponse({'error': 'Both generated_content and language are required fields.'}, status=400)

            # Validate the language
            if language_name not in SUPPORTED_LANGUAGES:
                logger.warning(f'Unsupported language requested: {language_name}')
                return JsonResponse({'error': f'Unsupported language. Supported languages are: {list(SUPPORTED_LANGUAGES.keys())}'}, status=400)

            # Get the corresponding language code
            language_code = SUPPORTED_LANGUAGES[language_name]
            logger.info(f'Translating content: {generated_content} to language: {language_name} ({language_code})')

            # Use deep_translator's GoogleTranslator for translation
            translated_content = GoogleTranslator(source='auto', target=language_code).translate(generated_content)
            logger.info(f'Content translated successfully: {translated_content}')  # Log the translated content
            # Encrypt the response content
            encrypted_response = encrypt_data({
                'generated_content': generated_content,
                'translated_content': translated_content,
                'selected_language': language_name
            })
            logger.debug(f'Encrypted response: {encrypted_response}')

            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.warning(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Unexpected error: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.error('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


import json
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from django.http import JsonResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated


MAX_WORKERS = 50  # Number of threads for concurrent processing
RETRY_LIMIT = 1000  # Maximum retries for translation API
MAX_SENTENCES_PER_CHUNK = 30  # Number of sentences to process in a single chunk

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def translate_content_formatted(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

    try:
        # Extract and decrypt the incoming payload
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)

        generated_content = data.get('generated_content')
        language = data.get('language')

        if not generated_content or not language:
            return JsonResponse({'error': 'Both generated_content and language are required fields.'}, status=400)

        # Initialize the translated content list
        translated_paragraphs = []

        # Break content into paragraphs or smaller chunks
        paragraphs = generated_content.split('\n\n')  # Split by paragraphs

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = []
            for paragraph in paragraphs:
                sentences = paragraph.split('. ')  # Split paragraph into sentences
                chunked_sentences = [
                    sentences[i:i + MAX_SENTENCES_PER_CHUNK]
                    for i in range(0, len(sentences), MAX_SENTENCES_PER_CHUNK)
                ]

                for chunk in chunked_sentences:
                    futures.append(
                        executor.submit(
                            lambda c: '. '.join([translate_with_retry(s, language) for s in c]),
                            chunk
                        )
                    )

            # Process completed translations
            for future in as_completed(futures):
                try:
                    translated_paragraphs.append(future.result())
                except ValueError as e:
                    return JsonResponse({'error': str(e)}, status=500)

        # Combine the translated paragraphs back into the full content
        translated_content = '\n\n'.join(translated_paragraphs)

        # Log the translated content
        logger.info(f'Translated content: {translated_content}')

        # Encrypt the response content
        encrypted_response = encrypt_data({
            'generated_content': generated_content,
            'translated_content': translated_content,
            'selected_language': language
        })

        return JsonResponse({'encrypted_content': encrypted_response}, status=200)

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)



import concurrent.futures

@csrf_exempt
@require_POST
def translate(request):
    if request.method != 'POST':
        logger.warning('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)

    try:
        # Extract and decrypt the incoming payload
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        
        input_text = data.get('input_text', '')
        from_language = data.get('from_language', '')
        to_language = data.get('to_language', '')

        if not (input_text and from_language and to_language):
            logger.warning('Missing input_text, from_language, or to_language.')
            return JsonResponse({'error': 'Please provide the input text and select both languages.'}, status=400)

        def perform_translation():
            # Perform the translation
            return bhashini_translate(input_text, to_language, from_language)["translated_content"]

         # Use ThreadPoolExecutor to run the translation in a separate thread
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(perform_translation)
            translated_text = future.result()
    

        logger.info('Translation successful')

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format received.'}, status=400)
    except KeyError as e:
        logger.error(f'Missing key in translation response: {str(e)}')
        return JsonResponse({'error': f'Missing key in translation response: {str(e)}'}, status=500)
    except Exception as e:
        logger.error(f'Error during request handling: {str(e)}')
        return JsonResponse({'error': f'Error during request handling: {str(e)}'}, status=500)

    # Prepare and encrypt the response
    response_data = {
        'translated_text': translated_text,
        'input_text': input_text,
        'from_language': from_language,
        'to_language': to_language
    }
    encrypted_response = encrypt_data(response_data)
    logger.debug(f'Encrypted response: {encrypted_response}')

    return JsonResponse({'encrypted_content': encrypted_response}, status=200)


@csrf_exempt
@require_POST
def translate_android(request):
    if request.method != 'POST':
        logger.warning('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)

    try:
        # Parse the incoming JSON payload
        data = json.loads(request.body.decode('utf-8'))
        
        input_text = data.get('input_text', '')
        from_language = data.get('from_language', '')
        to_language = data.get('to_language', '')

        if not (input_text and from_language and to_language):
            logger.warning('Missing input_text, from_language, or to_language.')
            return JsonResponse({'error': 'Please provide the input text and select both languages.'}, status=400)

        # Function to perform translation
        def perform_translation():
            # Perform the translation
            return bhashini_translate(input_text, to_language, from_language)["translated_content"]

        # Use ThreadPoolExecutor to run the translation in a separate thread
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(perform_translation)
            translated_text = future.result()
        
        logger.info('Translation successful')

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format received.'}, status=400)
    except KeyError as e:
        logger.error(f'Missing key in translation response: {str(e)}')
        return JsonResponse({'error': f'Missing key in translation response: {str(e)}'}, status=500)
    except Exception as e:
        logger.error(f'Error during request handling: {str(e)}')
        return JsonResponse({'error': f'Error during request handling: {str(e)}'}, status=500)

    # Prepare the plain JSON response
    response_data = {
        'translated_text': translated_text,
        'input_text': input_text,
        'from_language': from_language,
        'to_language': to_language
    }
    return JsonResponse(response_data, status=200)



SUPPORTED_LANGUAGES = {
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Hindi": "hi",
    "Chinese": "zh-CN",
    "Italian": "it",
    "Japanese": "ja",
    "Korean": "ko",
    "Portuguese": "pt",
    "Dutch": "nl",
    "Polish": "pl",
    "Russian": "ru",
    "Indonesian": "id",
    "Turkish": "tr",
}


@csrf_exempt
@require_POST
def translate_international(request):
    if request.method != 'POST':
        logger.warning('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)

    try:
        # Extract and decrypt the incoming payload
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)

        # If decrypted content is a string, parse it into a dictionary
        if isinstance(decrypted_content, str):
            decrypted_content = json.loads(decrypted_content)

        # Ensure decrypted content is a dictionary
        if not isinstance(decrypted_content, dict):
            logger.error('Decrypted content is not a dictionary')
            return JsonResponse({'error': 'Decrypted content is invalid.'}, status=400)

        input_text = decrypted_content.get('input_text', '')
        from_language_name = decrypted_content.get('from_language', '')
        to_language_name = decrypted_content.get('to_language', '')

        if not (input_text and from_language_name and to_language_name):
            logger.warning('Missing input_text, from_language, or to_language.')
            return JsonResponse({'error': 'Please provide the input text and select both languages.'}, status=400)

        # Validate the languages
        if from_language_name not in SUPPORTED_LANGUAGES or to_language_name not in SUPPORTED_LANGUAGES:
            logger.warning('Invalid language selection.')
            return JsonResponse({'error': 'One or both of the languages are not supported.'}, status=400)

        from_language = SUPPORTED_LANGUAGES[from_language_name]
        to_language = SUPPORTED_LANGUAGES[to_language_name]

        def perform_translation():
            # Perform the translation using GoogleTranslator
            return GoogleTranslator(source=from_language, target=to_language).translate(input_text)

        # Use ThreadPoolExecutor to run the translation in a separate thread
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(perform_translation)
            translated_text = future.result()

        logger.info('Translation successful')

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format received.'}, status=400)
    except KeyError as e:
        logger.error(f'Missing key in translation response: {str(e)}')
        return JsonResponse({'error': f'Missing key in translation response: {str(e)}'}, status=500)
    except Exception as e:
        logger.error(f'Error during request handling: {str(e)}')
        return JsonResponse({'error': f'Error during request handling: {str(e)}'}, status=500)

    # Prepare the response data
    response_data = {
        'translated_text': translated_text,
        'input_text': input_text,
        'from_language': from_language_name,
        'to_language': to_language_name
    }

    # Encrypt the response data before sending it back
    encrypted_response = encrypt_data(response_data)

    return JsonResponse({'encrypted_content': encrypted_response}, status=200)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def business_proposal_generator(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")
        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        # Parse decrypted JSON
        data = json.loads(decrypted_content)
        logger.debug(f"Parsed decrypted JSON: {data}")

        # Define Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Fields to check for language detection and translation
        fields_to_check = [
            'businessIntroduction', 'proposalObjective','otherObjective', 'scopeOfWork', 'projectPhases',
            'expectedOutcomes', 'technologiesAndInnovations', 'targetAudience',
            'budgetInformation', 'timeline', 'benefitsToRecipient', 'closingRemarks'
        ]

        # Translate only non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    logger.debug(f"Original Value: {value}")

                    # Translate if not English
                    if detected_language != 'en':
                        logger.info(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")

        # Extract processed fields
        business_intro = data.get('businessIntroduction')
        proposal_objective = data.get('proposalObjective')
        other_objective = data.get('otherObjective')

        # Handle 'Other' objective
        if proposal_objective == 'Others' and other_objective:
            proposal_objective = other_objective

        num_words = data.get('numberOfWords')
        scope_of_work = data.get('scopeOfWork')
        project_phases = data.get('projectPhases')
        expected_outcomes = data.get('expectedOutcomes')
        tech_innovations = data.get('technologiesAndInnovations')
        target_audience = data.get('targetAudience')
        budget_info = data.get('budgetInformation')
        timeline = data.get('timeline')
        benefits = data.get('benefitsToRecipient')
        closing_remarks = data.get('closingRemarks')

        # Generate the business proposal
        logger.info("Generating business proposal content.")
        proposal_content = generate_bus_pro(
            business_intro, proposal_objective, num_words, scope_of_work,
            project_phases, expected_outcomes, tech_innovations, target_audience,
            budget_info, timeline, benefits, closing_remarks
        )

        # Encrypt the response content
        encrypted_content = encrypt_data({'generated_content': proposal_content})
        logger.info("Business proposal content generated successfully.")

        return JsonResponse({'encrypted_content': encrypted_content}, status=200)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

#Encrypted API For Offer Letter Service
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def offer_letter_generator(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        encrypted_content = json.loads(body).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        logger.debug(f'Decrypted content: {data}')

        # Define the fields that need language detection and translation
        fields_to_check = [
            'companyDetails', 'candidateFullName', 'positionTitle', 'department',
            'status', 'location', 'compensationBenefits', 'workHours', 'termsConditions',
            'deadline', 'contactInfo', 'documentsNeeded', 'closingRemarks'
        ]

        # Define the Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Translate non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language of the field value
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    print(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    print(f"Original Value: {value}")

                    # If detected language is not English, translate
                    if detected_language != 'en':
                        print(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        print(f"Translated Value for {field}: {translated_text}")
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")

        # Extract fields from the processed data
        company_details = data.get('companyDetails')
        candidate_name = data.get('candidateFullName')
        position_title = data.get('positionTitle')
        department = data.get('department')
        status = data.get('status')
        location = data.get('location')
        start_date = data.get('expectedStartDate')
        compensation_benefits = data.get('compensationBenefits')  # Merged field
        work_hours = data.get('workHours')
        terms = data.get('termsConditions')
        acceptance_deadline = data.get('deadline')
        contact_info = data.get('contactInfo')
        documents_needed = data.get('documentsNeeded')
        closing_remarks = data.get('closingRemarks')

        logger.info('Generating offer letter content.')
        offer_letter_content = generate_offer_letter(
            company_details, candidate_name, position_title, department, status,
            location, start_date, compensation_benefits, work_hours,
            terms, acceptance_deadline, contact_info, documents_needed, closing_remarks
        )

        if offer_letter_content:
            encrypted_content = encrypt_data({'generated_content': offer_letter_content})
            logger.info('Offer letter content generated successfully.')
            return JsonResponse({'encrypted_content': encrypted_content}, status=200)

        logger.error('Failed to generate offer letter content.')
        return JsonResponse({'error': 'Failed to generate offer letter. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f'ValueError: {str(e)}')
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f'Exception: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@require_http_methods(["POST", "GET"])
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def profile(request):
    user = request.user
    profile = Profile.objects.get(user=user)
    errors = []

    if request.method == 'POST':
        try:
            # Decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            # Update user and profile data based on received JSON
            user.first_name = data.get('first_name', user.first_name)
            user.last_name = data.get('last_name', user.last_name)
            user.email = data.get('email', user.email)
            user.state = data.get('state', user.state)  # Update state field
            profile.bio = data.get('bio', profile.bio)
            profile.location = data.get('location', profile.location)
            profile.user_gst = data.get('user_gst', profile.user_gst)

            if not user.first_name:
                errors.append("First name is required.")
            if not user.last_name:
                errors.append("Last name is required.")
            if not user.email:
                errors.append("Email is required.")

            if not errors:
                user.save()
                profile.save()
                response_data = {'message': 'Profile updated successfully.'}
            else:
                response_data = {'errors': errors}

            encrypted_response = encrypt_data(response_data)
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    # Handle GET request
    response_data = {
        'user': {
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email,
            'state': user.state,  # Include state field
        },
        'profile': {
            'bio': profile.bio,
            'location': profile.location,
            'birth_date': profile.birth_date.isoformat() if profile.birth_date else None,
            'user_gst': profile.user_gst
        }
    }

    encrypted_response = encrypt_data(response_data)
    return JsonResponse({'encrypted_content': encrypted_response})



from django.core.exceptions import ObjectDoesNotExist



from django.http import JsonResponse
from django.contrib.auth.models import User
from rest_framework.decorators import api_view, permission_classes
from django.core.exceptions import ObjectDoesNotExist

@require_GET
@api_view(['GET'])
@permission_classes([IsAuthenticated, HasAPIKey])
def profile_info(request):
    if request.method == 'GET':
        try:
            # Get the encrypted content from the query parameter
            encrypted_content = request.GET.get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
            
            # Decrypt the content to get the original data
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            
            # Extract email from decrypted data
            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email parameter is required.'}, status=400)

            # Fetch the user based on the provided email
            user = User.objects.get(email=email)

            # Get the user's subscribed services
            user_service = UserService.objects.get(user=user)

            # Get all payments made by the user
            payments = Payment.objects.filter(email=user.email)

            # Prepare the payment details
            payment_info = []
            for payment in payments:
                payment_info.append({
                    'order_id': payment.order_id,
                    'payment_id': payment.payment_id,
                    'amount': str(payment.amount),
                    'currency': payment.currency,
                    'created_at': payment.created_at.isoformat(),  # Convert to ISO string
                    'verified': payment.verified,
                })

            # Prepare the response data with date fields converted to ISO format
            response_data = {
                'user_info': {
                    'first_name': user.first_name,
                    'last_name': user.last_name,
                    'email': user.email,
                    'last_login': user.last_login.isoformat() if user.last_login else None,
                    'username': user.username,
                    'date_joined': user.date_joined.isoformat(),
                    'state': user.state
                },
                'services': {
                    'email_service': user_service.email_service,
                    'offer_letter_service': user_service.offer_letter_service,
                    'business_proposal_service': user_service.business_proposal_service,
                    'sales_script_service': user_service.sales_script_service,
                    'content_generation_service': user_service.content_generation_service,
                    'summarize_service': user_service.summarize_service,
                    'ppt_generation_service': user_service.ppt_generation_service,
                    'blog_generation_service': user_service.blog_generation_service,
                    'rephrasely_service': user_service.rephrasely_service,
                    'service_start_dates': {
                        'email_service_start': user_service.email_end_date.isoformat() if user_service.email_end_date else None,
                        'offer_letter_service_start': user_service.offer_letter_end_date.isoformat() if user_service.offer_letter_end_date else None,
                        'business_proposal_service_start': user_service.business_proposal_end_date.isoformat() if user_service.business_proposal_end_date else None,
                        'sales_script_service_start': user_service.sales_script_end_date.isoformat() if user_service.sales_script_end_date else None,
                        'content_generation_service_start': user_service.content_generation_end_date.isoformat() if user_service.content_generation_end_date else None,
                        'summarize_service_start': user_service.summarize_end_date.isoformat() if user_service.summarize_end_date else None,
                        'ppt_generation_service_start': user_service.ppt_generation_end_date.isoformat() if user_service.ppt_generation_end_date else None,
                        'blog_generation_service_start': user_service.blog_generation_end_date.isoformat() if user_service.blog_generation_end_date else None,
                        'rephrasely_service_start': user_service.rephrasely_end_date.isoformat() if user_service.rephrasely_end_date else None,
                    },
                    'service_end_dates': {
                        'email_service_end': user_service.email_end_date.isoformat() if user_service.email_end_date else None,
                        'offer_letter_service_end': user_service.offer_letter_end_date.isoformat() if user_service.offer_letter_end_date else None,
                        'business_proposal_service_end': user_service.business_proposal_end_date.isoformat() if user_service.business_proposal_end_date else None,
                        'sales_script_service_end': user_service.sales_script_end_date.isoformat() if user_service.sales_script_end_date else None,
                        'content_generation_service_end': user_service.content_generation_end_date.isoformat() if user_service.content_generation_end_date else None,
                        'summarize_service_end': user_service.summarize_end_date.isoformat() if user_service.summarize_end_date else None,
                        'ppt_generation_service_end': user_service.ppt_generation_end_date.isoformat() if user_service.ppt_generation_end_date else None,
                        'blog_generation_service_end': user_service.blog_generation_end_date.isoformat() if user_service.blog_generation_end_date else None,
                        'rephrasely_service_end': user_service.rephrasely_end_date.isoformat() if user_service.rephrasely_end_date else None,
                    },
                },
                'payments': payment_info,
            }

            # Encrypt the response data
            encrypted_response = encrypt_data(response_data)
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ObjectDoesNotExist:
            return JsonResponse({'error': 'User not found.'}, status=404)
        except UserService.DoesNotExist:
            return JsonResponse({'error': 'User services not found.'}, status=404)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

# @csrf_exempt
# @require_POST
# @api_view(['POST'])
# @permission_classes([IsAuthenticated, HasAPIKey])
# def change_password(request):
#     if request.method == 'POST':
#         try:
#             # Decrypt incoming request body
#             encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
#             if not encrypted_content:
#                 logger.warning('No encrypted content found in the request.')
#                 return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

#             decrypted_content = decrypt_data(encrypted_content)
#             data = json.loads(decrypted_content)
#             logger.debug(f"Decrypted content: {data}")

#             current_password = data.get('current_password')
#             new_password = data.get('new_password')
#             confirm_new_password = data.get('confirm_new_password')

#             # Validate the current password
#             if not request.user.check_password(current_password):
#                 logger.warning(f"User {request.user.username} provided incorrect current password.")
#                 encrypted_response = encrypt_data({'error': 'Current password is incorrect.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Check if new passwords match
#             if new_password != confirm_new_password:
#                 logger.warning(f"User {request.user.username} provided non-matching new passwords.")
#                 encrypted_response = encrypt_data({'error': 'New passwords do not match.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Prevent using the same new password as the current password
#             if new_password == current_password:
#                 logger.warning(f"User {request.user.username} attempted to use the same new password as the current password.")
#                 encrypted_response = encrypt_data({'error': 'New password cannot be the same as the current password.'})
#                 return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#             # Update password
#             request.user.set_password(new_password)
#             request.user.save()

#             # Keep the user logged in after password change
#             update_session_auth_hash(request, request.user)
#             logger.info(f"User {request.user.username} successfully changed their password.")

#             encrypted_response = encrypt_data({'message': 'Password changed successfully.'})
#             return JsonResponse({'encrypted_content': encrypted_response})

#         except json.JSONDecodeError:
#             logger.error("Invalid JSON received.")
#             encrypted_response = encrypt_data({'error': 'Invalid JSON.'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=400)

#         except Exception as e:
#             logger.error(f'Internal server error: {str(e)}')
#             encrypted_response = encrypt_data({'error': 'Internal server error.'})
#             return JsonResponse({'encrypted_content': encrypted_response}, status=500)

#     logger.error("Invalid request method used.")
#     encrypted_response = encrypt_data({'error': 'Invalid request method.'})
#     return JsonResponse({'encrypted_content': encrypted_response}, status=405)

@csrf_exempt
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def change_password(request):
    if request.method == 'POST':
        try:
            # Decrypt incoming request body
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f"Decrypted content: {data}")

            current_password = data.get('current_password')
            new_password = data.get('new_password')
            confirm_new_password = data.get('confirm_new_password')

            # Validate the current password
            if not request.user.check_password(current_password):
                logger.warning(f"User {request.user.username} provided incorrect current password.")
                encrypted_response = encrypt_data({'error': 'Current password is incorrect.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Check if new passwords match
            if new_password != confirm_new_password:
                logger.warning(f"User {request.user.username} provided non-matching new passwords.")
                encrypted_response = encrypt_data({'error': 'New passwords do not match.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Prevent using the same new password as the current password
            if new_password == current_password:
                logger.warning(f"User {request.user.username} attempted to use the same new password as the current password.")
                encrypted_response = encrypt_data({'error': 'New password cannot be the same as the current password.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Check if the new password matches any of the last 5 passwords
            recent_passwords = UserPasswordHistory.get_recent_passwords(request.user)
            for recent_password in recent_passwords:
                if check_password(new_password, recent_password.hashed_password):
                    logger.warning(f"User {request.user.username} tried to use a password from their last 5 passwords.")
                    encrypted_response = encrypt_data({'error': 'New password cannot match any of the last 5 passwords.'})
                    return JsonResponse({'encrypted_content': encrypted_response}, status=400)

            # Update password
            request.user.set_password(new_password)
            request.user.save()

            # Store the new hashed password in the password history
            UserPasswordHistory.store_password(request.user, request.user.email, make_password(new_password))

            # Keep the user logged in after password change
            update_session_auth_hash(request, request.user)
            logger.info(f"User {request.user.username} successfully changed their password.")

            encrypted_response = encrypt_data({'message': 'Password changed successfully.'})
            return JsonResponse({'encrypted_content': encrypted_response})

        except json.JSONDecodeError:
            logger.error("Invalid JSON received.")
            encrypted_response = encrypt_data({'error': 'Invalid JSON.'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=400)

        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            encrypted_response = encrypt_data({'error': 'Internal server error.'})
            return JsonResponse({'encrypted_content': encrypted_response}, status=500)

    logger.error("Invalid request method used.")
    encrypted_response = encrypt_data({'error': 'Invalid request method.'})
    return JsonResponse({'encrypted_content': encrypted_response}, status=405)

# @require_POST
# @api_view(['POST'])
# @permission_classes([IsAuthenticated, HasAPIKey])
# def summarize_document(request):
#     try:
#         # Extract encrypted content from request.POST
#         encrypted_content = request.POST.get('encrypted_content')
#         if not encrypted_content:
#             logger.warning('No encrypted content found in the request.')
#             return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

#         # Decrypt the JSON payload
#         decrypted_content = decrypt_data(encrypted_content)
#         data = json.loads(decrypted_content)
#         logger.debug(f'Decrypted content: {data}')

#         # Define Indian languages mapping
#         indian_languages = {
#             "English": "en",
#             "Hindi": "hi",
#             "Tamil": "ta",
#             "Telugu": "te",
#             "Marathi": "mr",
#             "Kannada": "kn",
#             "Bengali": "bn",
#             "Odia": "or",
#             "Assamese": "as",
#             "Punjabi": "pa",
#             "Malayalam": "ml",
#             "Gujarati": "gu",
#             "Urdu": "ur",
#             "Sanskrit": "sa",
#             "Nepali": "ne",
#             "Bodo": "brx",
#             "Maithili": "mai",
#             "Sindhi": "sd",
#             "Kashmiri": "ks",
#             "Konkani": "kok",
#             "Dogri": "doi",
#             "Goan Konkani": "gom",
#             "Santali": "sat",
#         }

#         # Fields to check for language detection and translation
#         fields_to_check = [
#             'documentContext', 'mainSubject', 'summaryPurpose', 'lengthDetail',
#             'importantElements', 'audience', 'tone', 'format', 'additionalInstructions', 'text'
#         ]

#         # Translate only non-English content
#         for field in fields_to_check:
#             value = data.get(field)
#             if value:
#                 try:
#                     # Detect language
#                     detected_language, confidence = classify(value)
#                     language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
#                     logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
#                     logger.debug(f"Original Value: {value}")

#                     # Translate if not English
#                     if detected_language != 'en':
#                         logger.info(f"Translating {field} from {language_name} to English.")
#                         translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
#                         logger.debug(f"Translated {field}: {translated_text}")
#                         data[field] = translated_text
#                     else:
#                         logger.info(f"{field} is already in English. No translation needed.")
#                 except Exception as e:
#                     logger.error(f"Error processing field {field}: {str(e)}")

#         # Extract form fields from decrypted data
#         document_context = data.get('documentContext')
#         main_subject = data.get('mainSubject')
#         summary_purpose = data.get('summaryPurpose')
#         length_detail = data.get('lengthDetail')
#         important_elements = data.get('importantElements')
#         audience = data.get('audience')
#         tone = data.get('tone')
#         format_ = data.get('format')
#         additional_instructions = data.get('additionalInstructions')

#         # Extract the uploaded file or text from request
#         document_file = request.FILES.get('documentFile')
        
#         text = data.get('text')

#         # Ensure we have either documentFile or text
#         if not document_file and not text:
#             logger.warning('No document file or text provided.')
#             return JsonResponse({'error': 'No document file or text provided.'}, status=400)

#         # If documentFile is provided, use it for summarization
#         if document_file:
#             logger.info('Using uploaded document file for summarization.')
#             summary = generate_summary(
#                 document_context, main_subject, summary_purpose, length_detail,
#                 important_elements, audience, tone, format_, additional_instructions, document_file
#             )

#         # If only text is provided, use it for summarization
#         elif text:
#             logger.info('Using provided text for summarization.')
#             summary = generate_summary(
#                 document_context, main_subject, summary_purpose, length_detail,
#                 important_elements, audience, tone, format_, additional_instructions, text=text
#             )

#         # Handle specific error scenarios from generate_summary
#         if summary.startswith("Error:"):
#             if "Uploaded file too large" in summary:
#                 logger.warning(summary)
#                 return JsonResponse({'error': summary}, status=413)
#             else:
#                 logger.error(summary)
#                 return JsonResponse({'error': summary}, status=500)

#         # Encrypt the response content
#         encrypted_response = encrypt_data({'summary': summary})
#         logger.info('Summary generated and encrypted successfully.')

#         return JsonResponse({'encrypted_content': encrypted_response}, status=200)

#     except json.JSONDecodeError:
#         logger.error('Invalid JSON format received.')
#         return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
#     except ValueError as e:
#         logger.error(f'ValueError: {str(e)}')
#         return JsonResponse({'error': str(e)}, status=400)
#     except Exception as e:
#         logger.error(f'Exception: {str(e)}')
#         return JsonResponse({'error': str(e)}, status=500)


@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def summarize_document(request):
    try:
        # Extract encrypted content from request.POST
        encrypted_content = request.POST.get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
 
        # Decrypt the JSON payload
        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        print(decrypted_content);
        logger.debug(f'Decrypted content: {data}')
 
        # Define Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }
 
        # Fields to check for language detection and translation
        fields_to_check = [
            'documentContext', 'mainSubject', 'summaryPurpose', 'lengthDetail',
            'importantElements', 'audience', 'tone', 'format', 'additionalInstructions','otherDocumentContext',
            'otherMainSubject', 'otherPurpose', 'otherLengthDetail', 'otherFormat', 'customTone'
        ]
 
        # Translate only non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    logger.debug(f"Original Value: {value}")
 
                    # Translate if not English
                    if detected_language != 'en':
                        logger.info(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")
 
        # Extract form fields from decrypted data
        document_context = data.get('documentContext')
        if document_context == "Other":
            document_context = data.get('otherDocumentContext')
        else:
            document_context = document_context
        main_subject = data.get('mainSubject')
        if(main_subject == "Other"):
            main_subject = data.get('otherMainSubject')
        else:
            main_subject = main_subject
        summary_purpose = data.get('summaryPurpose')
        if(summary_purpose == "Other"):
            summary_purpose = data.get('otherPurpose')
        else:
            summary_purpose = summary_purpose
        length_detail = data.get('lengthDetail')
        if(length_detail == "Other"):
            length_detail = data.get('otherLengthDetail')
        else:
            length_detail = length_detail
        important_elements = data.get('importantElements')
        audience = data.get('audience')
        tone = data.get('tone')
        if(tone == "Other"):
            tone = data.get('customTone')
        else:
            tone = tone
        format_ = data.get('format')
        if(format_ == "Other"):
            format_ = data.get('otherFormat')
        else:
            format_ = format_
        additional_instructions = data.get('additionalInstructions')
 
        # Extract the uploaded file or text from request
        document_file = request.FILES.get('documentFile')
        text = data.get('text')
 
        # Ensure we have either documentFile or text
        if not document_file and not text:
            logger.warning('No document file or text provided.')
            return JsonResponse({'error': 'No document file or text provided.'}, status=400)
 
        # If documentFile is provided, use it for summarization
        if document_file:
            logger.info('Using uploaded document file for summarization.')
            summary = generate_summary(
                document_context, main_subject, summary_purpose, length_detail,
                important_elements, audience, tone, format_, additional_instructions, document_file
            )
 
        # If only text is provided, use it for summarization
        elif text:
            logger.info('Using provided text for summarization.')
            summary = generate_summary(
                document_context, main_subject, summary_purpose, length_detail,
                important_elements, audience, tone, format_, additional_instructions, text=text
            )
 
        # Handle specific error scenarios from generate_summary
        if summary.startswith("Error:"):
            if "Uploaded file too large" in summary:
                logger.warning(summary)
                return JsonResponse({'error': summary}, status=413)
            else:
                logger.error(summary)
                return JsonResponse({'error': summary}, status=500)
 
        # Encrypt the response content
        encrypted_response = encrypt_data({'summary': summary})
        logger.info('Summary generated and encrypted successfully.')
 
        return JsonResponse({'encrypted_content': encrypted_response}, status=200)
 
    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f'ValueError: {str(e)}')
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f'Exception: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)

#Encrypted API For contnet generation Service
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def content_generator(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")

        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        data = json.loads(decrypted_content)

        # Define the fields that need language detection and translation
        fields_to_check = [
            'company_info', 'content_purpose', 'desired_action', 'topic_details', 
            'keywords', 'audience_profile', 'format_structure', 'seo_keywords', 'references'
        ]

        # Define the Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Translate non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language of the field value
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    print(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    print(f"Original Value: {value}")

                    # If detected language is not English, translate
                    if detected_language != 'en':
                        print(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        print(f"Translated Value for {field}: {translated_text}")
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")   

        # Extract fields from the processed data
        company_info = data.get('company_info')
        content_purpose = data.get('content_purpose')
        desired_action = data.get('desired_action')
        topic_details = data.get('topic_details')
        keywords = data.get('keywords')
        audience_profile = data.get('audience_profile')
        format_structure = data.get('format_structure')
        num_words = data.get('num_words')
        seo_keywords = data.get('seo_keywords')
        references = data.get('references')

        logger.debug(f"Data extracted for content generation: company_info={company_info}, content_purpose={content_purpose}, desired_action={desired_action}")

        # Generate the content
        logger.info("Generating content...")
        content = generate_content(
            company_info,
            content_purpose,
            desired_action,
            topic_details,
            keywords,
            audience_profile,
            format_structure,
            num_words,
            seo_keywords,
            references
        )

        if content:
            logger.info("Content generated successfully.")
            encrypted_response_content = encrypt_data({'generated_content': content})
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        logger.error("Failed to generate content.")
        return JsonResponse({'error': 'Failed to generate content. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


from langid import classify

from deep_translator import GoogleTranslator

#With Language Detection and Translation using Google Translate API
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated, HasAPIKey])
def sales_script_generator(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")

        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        # Parse the decrypted JSON
        data = json.loads(decrypted_content)
        logger.debug(f"Parsed decrypted JSON: {data}")

        # Define the Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Fields that require language detection and potential translation
        fields_to_check = [
            'company_details', 'product_descriptions', 'features_benefits', 
            'pricing_info', 'promotions', 'target_audience', 
            'sales_objectives', 'competitive_advantage', 'compliance'
        ]

        # Translate only non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language of the field value
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    print(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")
                    print(f"Original Value: {value}")
                    
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

                    # Translate if the detected language is not English
                    if detected_language != 'en':
                        print(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        print(f"Translated Value for {field}: {translated_text}")
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        print(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    print(f"Error processing field {field}: {str(e)}")
                    logger.error(f"Error processing field {field}: {str(e)}")

        logger.debug(f"Data after translation: {data}")

        # Extract fields from the processed data
        num_words = data.get('num_words')
        company_details = data.get('company_details')
        product_descriptions = data.get('product_descriptions')
        features_benefits = data.get('features_benefits')
        pricing_info = data.get('pricing_info')
        promotions = data.get('promotions')
        target_audience = data.get('target_audience')
        sales_objectives = data.get('sales_objectives')
        competitive_advantage = data.get('competitive_advantage')
        compliance = data.get('compliance')

        logger.debug(f"Data extracted for sales script generation: num_words={num_words}, company_details={company_details}")

        # Generate the sales script
        logger.info("Generating sales script...")
        sales_script = generate_sales_script(
            company_details,
            num_words,
            product_descriptions,
            features_benefits,
            pricing_info,
            promotions,
            target_audience,
            sales_objectives,
            competitive_advantage,
            compliance,
        )

        if sales_script:
            logger.info("Sales script generated successfully.")
            encrypted_response_content = encrypt_data({'generated_content': sales_script})
            return JsonResponse({
                'encrypted_content': encrypted_response_content,
                'language': 'en'
            }, status=200)

        logger.error("Failed to generate sales script.")
        return JsonResponse({'error': 'Failed to generate sales script. Please try again.'}, status=500)
    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def logout_view(request):
    try:
        # Fetch the active session for the logged-in user
        user_session = UserSession.objects.filter(user=request.user, active=True).first()

        if user_session:
            # Set the active field to False
            user_session.active = False
            user_session.save()

        # Perform Django logout operation
        logout(request)
        logger.info(f"User {request.user.username} logged out successfully.")
        return JsonResponse({'success': 'Logged out successfully'}, status=200)
    except Exception as e:
        logger.error(f"Error during logout: {str(e)}")
        return JsonResponse({'error': 'An error occurred during logout.'}, status=500)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_presentation(request):
    try:
        # Handle the multipart form data
        encrypted_content = request.POST.get('encrypted_content')
        if not encrypted_content:
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        # Decrypt the content
        decrypted_content = decrypt_data(encrypted_content)

        # Parse decrypted JSON data
        data = json.loads(decrypted_content)

        # Extract fields from the decrypted data
        title = data.get('title')
        num_slides = data.get('num_slides')
        special_instructions = data.get('special_instructions')
        template_name = data.get('template_name', 'default')
        document = request.FILES.get('document')

        if not title or not num_slides:
            return JsonResponse({'error': 'Title and number of slides are required.'}, status=400)

        # Handle document content optionally
        document_content = extract_document_content(document) if document else ""
        word_count = len(document_content.split())
        print(f"Word count: {word_count}")
        if word_count > 20000:
            return JsonResponse({'error': 'Document content exceeds the word limit of 20000 words.'}, status=413)

        # Get template path
        templates = get_templates()
        template_path = templates.get(template_name, templates['default'])
        output_path = "SmartAssistant_Presentation.pptx"

        # Generate presentation with the provided data
        prs = update_presentation_with_generated_content(
            template_path, output_path, document_content, title, num_slides, special_instructions
        )

        # Save presentation to a BytesIO object
        buffer = BytesIO()
        prs.save(buffer)
        buffer.seek(0)  # Rewind the buffer

        # Return file response
        response = FileResponse(buffer, as_attachment=True, filename='SmartOffice_Assistant_Presentation.pptx')
        return response

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_presentation_english(request):
    try:
        # Handle the multipart form data
        encrypted_content = request.POST.get('encrypted_content')
        if not encrypted_content:
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        # Decrypt the content
        decrypted_content = decrypt_data(encrypted_content)

        # Parse decrypted JSON data
        data = json.loads(decrypted_content)

        # Define the fields that need language detection and translation
        fields_to_check = ['title', 'special_instructions']

        # Define the Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Translate non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language of the field value
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

                    # If detected language is not English, translate
                    if detected_language != 'en':
                        logger.info(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")

        # Extract fields from the decrypted and processed data
        title = data.get('title')
        num_slides = data.get('num_slides')
        special_instructions = data.get('special_instructions')
        template_name = data.get('template_name', 'default')
        document = request.FILES.get('document')

        if not title or not num_slides:
            return JsonResponse({'error': 'Title and number of slides are required.'}, status=400)

        # Handle document content optionally
        document_content = extract_document_content(document) if document else ""
        word_count = len(document_content.split())
        logger.debug(f"Word count: {word_count}")
        if word_count > 2000:
            return JsonResponse({'error': 'Document content exceeds the word limit of 2000 words.'}, status=413)

        # Get template path
        templates = get_templates()
        template_path = templates.get(template_name, templates['default'])
        output_path = "SmartAssistant_Presentation.pptx"

        # Generate presentation with the provided data
        prs = update_presentation_with_generated_content(
            template_path, output_path, document_content, title, num_slides, special_instructions
        )

        # Save presentation to a BytesIO object
        buffer = BytesIO()
        prs.save(buffer)
        buffer.seek(0)  # Rewind the buffer

        # Return file response
        response = FileResponse(buffer, as_attachment=True, filename='SmartOffice_Assistant_Presentation.pptx')
        return response

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)
    
# @require_POST
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def generate_blog_view(request):
#     try:
#         # Load and decode the request body
#         body = request.body.decode('utf-8')
#         logger.debug(f"Request body received: {body}")

#         # Extract and decrypt the incoming payload
#         data = json.loads(body)
#         encrypted_content = data.get('encrypted_content')
#         if not encrypted_content:
#             logger.warning("No encrypted content found in the request.")
#             return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

#         logger.debug(f"Encrypted content received: {encrypted_content}")
#         decrypted_content = decrypt_data(encrypted_content)
#         logger.debug(f"Decrypted content: {decrypted_content}")

#         # Parse the decrypted JSON
#         data = json.loads(decrypted_content)

#         # Define the Indian languages mapping
#         indian_languages = {
#             "English": "en",
#             "Hindi": "hi",
#             "Tamil": "ta",
#             "Telugu": "te",
#             "Marathi": "mr",
#             "Kannada": "kn",
#             "Bengali": "bn",
#             "Odia": "or",
#             "Assamese": "as",
#             "Punjabi": "pa",
#             "Malayalam": "ml",
#             "Gujarati": "gu",
#             "Urdu": "ur",
#             "Sanskrit": "sa",
#             "Nepali": "ne",
#             "Bodo": "brx",
#             "Maithili": "mai",
#             "Sindhi": "sd",
#             "Kashmiri": "ks",
#             "Konkani": "kok",
#             "Dogri": "doi",
#             "Goan Konkani": "gom",
#             "Santali": "sat",
#         }

#         # Fields that require language detection and potential translation
#         fields_to_check = ['title', 'tone', 'keywords', 'customTone']

#         # Translate only non-English content
#         for field in fields_to_check:
#             value = data.get(field)
#             if value:
#                 try:
#                     # Detect language of the field value
#                     detected_language, confidence = classify(value)
#                     language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
#                     logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

#                     # Translate if the detected language is not English
#                     if detected_language != 'en':
#                         translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
#                         logger.debug(f"Translated {field}: {translated_text}")
#                         data[field] = translated_text
#                     else:
#                         print(f"{field} is already in English. No translation needed.")
#                 except Exception as e:
#                     print(f"Error processing field {field}: {str(e)}")
#                     logger.error(f"Error processing field {field}: {str(e)}")

#         logger.debug(f"Data after translation: {data}")

#         # Extract the required fields
#         title = data.get('title')
#         tone = data.get('tone')
#         custom_tone = data.get('customTone')  # Extract the custom tone
#         keywords = data.get('keywords', None) 

#         # Ensure required fields are present
#         if not title or not tone:
#             return JsonResponse({"error": "Missing 'title' or 'tone'."}, status=400)

#         # Call the generate_blog function
#         logger.info("Generating blog content...")

#         blog_content = generate_blog(title, tone, custom_tone, keywords)


#         if blog_content:
#             logger.info("Blog content generated successfully.")
#             encrypted_response_content = encrypt_data({'blog_content': blog_content})
#             return JsonResponse({
#                 'encrypted_content': encrypted_response_content,
#                 'language': 'en'
#             }, status=200)

#         logger.error("Failed to generate blog content.")
#         return JsonResponse({'error': 'Failed to generate blog content. Please try again.'}, status=500)

#     except json.JSONDecodeError:
#         logger.error("Invalid JSON format received.")
#         return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
#     except ValueError as e:
#         logger.error(f"ValueError occurred: {str(e)}")
#         return JsonResponse({"error": str(e)}, status=400)
#     except Exception as e:
#         logger.error(f"An unexpected error occurred: {str(e)}")
#         return JsonResponse({"error": str(e)}, status=500)

#     # If not a POST request, return an error
#     return JsonResponse({"error": "Only POST method is allowed."}, status=405)

# @require_POST
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def generate_blog_view(request):
#     try:
#         # Load and decode the request body
#         body = request.body.decode('utf-8')
#         logger.debug(f"Request body received: {body}")

#         # Extract and decrypt the incoming payload
#         data = json.loads(body)
#         encrypted_content = data.get('encrypted_content')
#         if not encrypted_content:
#             logger.warning("No encrypted content found in the request.")
#             return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

#         logger.debug(f"Encrypted content received: {encrypted_content}")
#         decrypted_content = decrypt_data(encrypted_content)
#         logger.debug(f"Decrypted content: {decrypted_content}")

#         # Parse the decrypted JSON
#         data = json.loads(decrypted_content)

#         # Define the Indian languages mapping
#         indian_languages = {
#             "English": "en",
#             "Hindi": "hi",
#             "Tamil": "ta",
#             "Telugu": "te",
#             "Marathi": "mr",
#             "Kannada": "kn",
#             "Bengali": "bn",
#             "Odia": "or",
#             "Assamese": "as",
#             "Punjabi": "pa",
#             "Malayalam": "ml",
#             "Gujarati": "gu",
#             "Urdu": "ur",
#             "Sanskrit": "sa",
#             "Nepali": "ne",
#             "Bodo": "brx",
#             "Maithili": "mai",
#             "Sindhi": "sd",
#             "Kashmiri": "ks",
#             "Konkani": "kok",
#             "Dogri": "doi",
#             "Goan Konkani": "gom",
#             "Santali": "sat",
#         }


#         # Fields that require language detection and potential translation
#         fields_to_check = ['title', 'tone', 'keywords', 'customTone']

#         # Translate only non-English content
#         for field in fields_to_check:
#             value = data.get(field)
#             if value:
#                 try:
#                     # Detect language of the field value
#                     detected_language, confidence = classify(value)
#                     language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
#                     logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

#                     # Translate if the detected language is not English
#                     if detected_language != 'en':
#                         translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
#                         logger.debug(f"Translated {field}: {translated_text}")
#                         data[field] = translated_text
#                 except Exception as e:
#                     logger.error(f"Error processing field {field}: {str(e)}")

#         logger.debug(f"Data after translation: {data}")

#         # Extract the required fields
#         title = data.get('title')
#         tone = data.get('tone')
#         custom_tone = data.get('customTone')  # Extract the custom tone
#         keywords = data.get('keywords', None) 

#         # Ensure required fields are present
#         if not title or not tone:
#             return JsonResponse({"error": "Missing 'title' or 'tone'."}, status=400)

#         # Generate blog content
#         logger.info("Generating blog content...")
#         blog_content = generate_blog(title, tone, custom_tone, keywords)

#         # Generate image
#         logger.info("Fetching related image...")
#         query = title if not keywords else f"{title}, {', '.join(keywords)}"
#         image = fetch_single_image(query, width=800, height=600)
#         print(image)
#         if blog_content:
#             logger.info("Blog content generated successfully.")
#             response_data = {
#                 'blog_content': blog_content,
#                 'image_url': image['url'] if image and 'url' in image else None
#             }
#             encrypted_response_content = encrypt_data(response_data)
#             return JsonResponse({
#                 'encrypted_content': encrypted_response_content,
#                 'language': 'en'
#             }, status=200)

#         logger.error("Failed to generate blog content.")
#         return JsonResponse({'error': 'Failed to generate blog content. Please try again.'}, status=500)

#     except json.JSONDecodeError:
#         logger.error("Invalid JSON format received.")
#         return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
#     except ValueError as e:
#         logger.error(f"ValueError occurred: {str(e)}")
#         return JsonResponse({"error": str(e)}, status=400)
#     except Exception as e:
#         logger.error(f"An unexpected error occurred: {str(e)}")
#         return JsonResponse({"error": str(e)}, status=500)

#     # If not a POST request, return an error
#     return JsonResponse({"error": "Only POST method is allowed."}, status=405)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_blog_view(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")
        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        # Parse the decrypted JSON
        data = json.loads(decrypted_content)

        # Define the Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Fields that require language detection and potential translation
        fields_to_check = ['title', 'tone', 'keywords', 'customTone']

        # Translate only non-English content
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language of the field value
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")
                    logger.info(f"Field: {field} - Detected Language: {language_name} (Confidence: {confidence:.2f})")

                    # Translate if the detected language is not English
                    if detected_language != 'en':
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")

        logger.debug(f"Data after translation: {data}")

        # Extract the required fields
        title = data.get('title')
        tone = data.get('tone')
        custom_tone = data.get('customTone')  # Extract the custom tone
        keywords = data.get('keywords', None)

        # Ensure required fields are present
        if not title or not tone:
            return JsonResponse({"error": "Missing 'title' or 'tone'."}, status=400)

        # Generate blog content
        logger.info("Generating blog content...")
        blog_content = generate_blog(title, tone, custom_tone, keywords)

        # Generate image
        logger.info("Fetching related image...")
        query = title + ","+ keywords if keywords else title
        image = fetch_single_image(query, width=800, height=600)

        if blog_content:
            logger.info("Blog content generated successfully.")
            response_data = {
                'blog_content': blog_content,
                'image_url': image['url'] if image and 'url' in image else None,
                'image_base64': image['base64_image'] if image and 'base64_image' in image else None
            }
            encrypted_response_content = encrypt_data(response_data)
            return JsonResponse({
                'encrypted_content': encrypted_response_content,
                'language': 'en'
            }, status=200)

        logger.error("Failed to generate blog content.")
        return JsonResponse({'error': 'Failed to generate blog content. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=500)

@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def regenerate_image(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")
        decrypted_content = decrypt_data(encrypted_content)
        print(decrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        # Parse the decrypted JSON
        data = json.loads(decrypted_content)

        # Extract title and keywords
        title = data.get('title')
        keywords = data.get('keywords')

        # Ensure required fields are present
        if not title:
            return JsonResponse({"error": "Missing 'title' in the request."}, status=400)

        # Generate the query string for the image
        query = title + ("," + keywords if keywords else "")
        logger.debug(f"Generated query for image: {query}")

        # Fetch related image
        logger.info("Fetching related image...")
        image = fetch_single_image(query, width=800, height=600)

        if image:
            logger.info("Image fetched successfully.")
            response_data = {
                'imagebase64': image['base64_image'] if 'base64_image' in image else None
            }
            encrypted_response_content = encrypt_data(response_data)
            return JsonResponse({
                'encrypted_content': encrypted_response_content
            }, status=200)

        logger.error("Failed to fetch image.")
        return JsonResponse({'error': 'Failed to fetch image. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=500)


#Encrypted API For rephrase Service
@require_POST
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def rephrasely_view(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')
        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")
        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        # Parse decrypted JSON
        data = json.loads(decrypted_content)
        logger.debug(f"Parsed decrypted JSON: {data}")

        # Define Indian languages mapping
        indian_languages = {
            "English": "en",
            "Hindi": "hi",
            "Tamil": "ta",
            "Telugu": "te",
            "Marathi": "mr",
            "Kannada": "kn",
            "Bengali": "bn",
            "Odia": "or",
            "Assamese": "as",
            "Punjabi": "pa",
            "Malayalam": "ml",
            "Gujarati": "gu",
            "Urdu": "ur",
            "Sanskrit": "sa",
            "Nepali": "ne",
            "Bodo": "brx",
            "Maithili": "mai",
            "Sindhi": "sd",
            "Kashmiri": "ks",
            "Konkani": "kok",
            "Dogri": "doi",
            "Goan Konkani": "gom",
            "Santali": "sat",
        }

        # Check and translate fields
        fields_to_check = ['text_to_rephrase', 'tone', 'target_audience']
        for field in fields_to_check:
            value = data.get(field)
            if value:
                try:
                    # Detect language
                    detected_language, confidence = classify(value)
                    language_name = next((k for k, v in indian_languages.items() if v == detected_language), "Unknown")

                    # Translate if not English
                    if detected_language != 'en':
                        logger.info(f"Translating {field} from {language_name} to English.")
                        translated_text = GoogleTranslator(source=detected_language, target='en').translate(value)
                        logger.debug(f"Translated {field}: {translated_text}")
                        data[field] = translated_text
                    else:
                        logger.info(f"{field} is already in English. No translation needed.")
                except Exception as e:
                    logger.error(f"Error processing field {field}: {str(e)}")

        # Extract required fields
        text_to_rephrase = data.get('text_to_rephrase')
        tone = data.get('tone')
        target_audience = data.get('target_audience')
        num_words = data.get('num_words', "default")  # Optional

        # Ensure required fields are present
        if not text_to_rephrase or not tone:
            return JsonResponse({'error': "Missing 'text_to_rephrase' or 'tone'."}, status=400)

        # Call the rephrasely function
        logger.info("Generating rephrased text...")
        rephrased_text = rephrasely(text_to_rephrase, tone, target_audience, num_words)
        logger.info("Rephrased text generated successfully.")

        # Encrypt the response content
        encrypted_response = encrypt_data({'rephrased_text': rephrased_text})

        # Return the encrypted response
        return JsonResponse({'encrypted_content': encrypted_response}, status=200)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

# List of greetings
GREETING_MESSAGES = [
    "Hey there! 👋 I’m Advika, your friendly AI champion. Got a question about our AI services? Let’s brighten your day with the perfect solution! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Greetings, human! I’m Advika, your digital assistant built for speed ⚡. Ask me anything about our AI services, and let’s get things done in a flash! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hello and welcome! 🌟 I’m Advika, here to assist with all your AI-related queries. What’s on your mind? Let’s dive into our exciting services together! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hi there! I’m Advika, your guide to exploring the world of AI. Have a question about our services? Let’s explore it together—just ask away! 🤔 You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hello, superstar! 🌟 I’m Advika, your AI sidekick. Curious about our amazing AI services? I’ve got all the answers, let’s get started! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "✨ Welcome, adventurer! ✨ I’m Advika, your AI guide on this exciting journey. Ask me about any of our cutting-edge AI services, and let’s unlock some magic together! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hey! I’m Advika, here to support you on your AI journey. Got questions about our services? Don’t worry, we’ll tackle them together—let’s get started! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hi there! ⏱ I’m Advika, and I’m here to help you quickly explore our AI services. Ask away, and I’ll provide the answers in no time! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Greetings! I’m Advika, your dedicated AI assistant. Have any questions about our AI offerings? I’m here to guide you—how may I assist you today? You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk).",
    "Hey, awesome human! 🎉 I’m Advika, your go-to AI for all things related to our services. Got a question? Let’s make it happen—go champion! You can also check out our [Resource Hub here](https://prodigidesk.ai/ProdigiDesk)."
]



require_http_methods(["GET", "POST"])
@api_view(['GET', 'POST'])
@permission_classes([])  
def chatbot_view(request):
    if settings.FAISS_VECTOR_STORE is None:
        return JsonResponse({"error": "Vector store not initialized"}, status=500)

    if request.method == 'GET':
        # Randomly select a greeting message
        greeting_message = random.choice(GREETING_MESSAGES)
        encrypted_message = encrypt_data({'answer': greeting_message})
        return JsonResponse({'encrypted_content': encrypted_message}, status=200)

    elif request.method == 'POST':
        try:
            # Extract and decrypt the incoming data
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            question = data.get('question')

            if not question:
                return JsonResponse({'error': 'No question provided.'}, status=400)

            # Handle user input (follow-up question)
            result = ask_question_chatbot(question)

            # Encrypt the response
            encrypted_response = encrypt_data({'answer': result})
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    else:
        return JsonResponse({'error': 'Invalid request method'}, status=405)

@api_view(['POST'])
@permission_classes([])
def speech_api(request):
    if request.method == 'POST':
        try:
            # Parse the JSON data from the request body
            data = json.loads(request.body)
            text = data.get('text', 'No text provided')
           
            # Generate speech from text using gTTS
            tts = gTTS(text=text, lang='en', slow=False)
 
            # In-memory file object for the audio
            mp3_fp = io.BytesIO()
            tts.write_to_fp(mp3_fp)
 
            mp3_fp.seek(0)  # Reset the pointer to the beginning of the file
 
            # Prepare the HTTP response with the audio file
            response = HttpResponse(mp3_fp, content_type='audio/mpeg')
            response['Content-Disposition'] = 'attachment; filename="speech.mp3"'
 
            return response
       
        except Exception as e:
            # Handle any exception and return error message
            return JsonResponse({'status': 'error', 'message': str(e)})
   
    else:
        # Handle non-POST requests
        return JsonResponse({'status': 'error', 'message': 'Only POST requests are allowed.'})

from django.views.decorators.http import require_POST

@csrf_exempt
@require_POST
# @api_view(['POST'])
def email_generator_guest(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')

            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            purpose = data.get('purpose')
            if purpose == 'Other':
                purpose = data.get('otherPurpose')
            num_words = data.get('num_words')
            subject = data.get('subject')
            rephrase = data.get('rephraseSubject', False)
            to = data.get('to')
            tone = data.get('tone')
            keywords = data.get('keywords', [])
            contextual_background = data.get('contextualBackground')
            call_to_action = data.get('callToAction')
            if call_to_action == 'Other':
                call_to_action = data.get('otherCallToAction')
            additional_details = data.get('additionalDetails')
            priority_level = data.get('priorityLevel')
            closing_remarks = data.get('closingRemarks')


            # Retrieve the guest email and existing word count if it exists
            guest_email = data.get('email')  # Assuming the email is part of the data
            if guest_email:
                guest_word_count = GuestLogin.objects.filter(email=guest_email).first()
                if guest_word_count and guest_word_count.word_count >= 2000:
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            # Generate email content based on the data
            generated_content = generate_email(
                purpose, num_words, subject, rephrase, to, tone, keywords,
                contextual_background, call_to_action, additional_details,
                priority_level, closing_remarks
            )

            if generated_content:

                # Calculate the word count for the generated content
                word_count = len(generated_content.split())

                # Encrypt the response content
                encrypted_response = encrypt_data({'generated_content': generated_content})

                # Update or create the word count in the GuestLogin model
                if guest_email:
                    if guest_word_count:
                        # Update existing word count
                        guest_word_count.word_count += word_count
                        guest_word_count.save()
                    else:
                        # Create new record if no existing word count found
                        guest_word_count = GuestLogin.objects.create(
                            email=guest_email,
                            word_count=word_count
                        )
                else:
                    print('No email found in the request data.')

                # Return the encrypted response with the generated content
                return JsonResponse({'encrypted_content': encrypted_response})

            else:
                return JsonResponse({'error': 'Failed to generate email content.'}, status=500)

        except Exception as e:
            return JsonResponse({'error': 'An error occurred while processing the request.'}, status=500)



@csrf_exempt
@require_POST
def business_proposal_generator_guest(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            # Extract fields for the business proposal
            business_intro = data.get('businessIntroduction')
            proposal_objective = data.get('proposalObjective')
            # Handle otherObjective if proposalObjective is 'Others'
            other_objective = data.get('otherObjective')
            if proposal_objective == 'Others' and other_objective:
                proposal_objective = other_objective

            num_words = data.get('numberOfWords')
            scope_of_work = data.get('scopeOfWork')
            project_phases = data.get('projectPhases')
            expected_outcomes = data.get('expectedOutcomes')
            tech_innovations = data.get('technologiesAndInnovations')  # Combined field
            target_audience = data.get('targetAudience')
            budget_info = data.get('budgetInformation')
            timeline = data.get('timeline')
            benefits = data.get('benefitsToRecipient')
            closing_remarks = data.get('closingRemarks')

            # Retrieve the guest email and existing word count if it exists
            guest_email = data.get('email')
            if guest_email:
                guest_word_count = GuestLogin.objects.filter(email=guest_email).first()
                if guest_word_count and guest_word_count.word_count >= 2000:
                    logger.info('Word Count limit already exceeded')
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            logger.info('Generating business proposal content.')
            proposal_content = generate_bus_pro(
                business_intro, proposal_objective, num_words, scope_of_work,
                project_phases, expected_outcomes, tech_innovations, target_audience,
                budget_info, timeline, benefits, closing_remarks
            )

            # Calculate the word count for the generated content
            word_count = len(proposal_content.split())
            logger.debug(f"Generated proposal content word count: {word_count}")
            print(word_count)

            # Encrypt the response content
            encrypted_content = encrypt_data({'generated_content': proposal_content})
            logger.info('Business proposal content generated successfully.')

            # Update or create the word count in the GuestLogin model
            if guest_email:
                if guest_word_count:
                    # Update existing word count
                    guest_word_count.word_count += word_count
                    guest_word_count.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {guest_word_count.word_count}')
                else:
                    # Create new record if no existing word count found
                    guest_word_count = GuestLogin.objects.create(
                        email=guest_email,
                        word_count=word_count
                    )
                    logger.info(f'Word count for {guest_email} saved successfully. Word count: {word_count}')
            else:
                logger.warning('No email found in the request data.')

            # Return the encrypted response with the generated content
            return JsonResponse({'encrypted_content': encrypted_content}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def offer_letter_generator_guest(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            # Extract fields for the offer letter
            company_details = data.get('companyDetails')
            candidate_name = data.get('candidateFullName')
            position_title = data.get('positionTitle')
            department = data.get('department')
            status = data.get('status')
            location = data.get('location')
            start_date = data.get('expectedStartDate')
            compensation_benefits = data.get('compensationBenefits')
            work_hours = data.get('workHours')
            terms = data.get('termsConditions')
            acceptance_deadline = data.get('deadline')
            contact_info = data.get('contactInfo')
            documents_needed = data.get('documentsNeeded')
            closing_remarks = data.get('closingRemarks')

            # Retrieve the guest email and existing word count if it exists
            guest_email = data.get('email')
            if guest_email:
                guest_record = GuestLogin.objects.filter(email=guest_email).first()
                if guest_record and guest_record.word_count >= 2000:
                    logger.info('Word Count limit already exceeded')
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            logger.info('Generating offer letter content.')
            offer_letter_content = generate_offer_letter(
                company_details, candidate_name, position_title, department, status,
                location, start_date, compensation_benefits, work_hours,
                terms, acceptance_deadline, contact_info, documents_needed, closing_remarks
            )

            # Calculate the word count for the generated content
            word_count = len(offer_letter_content.split())
            logger.debug(f"Generated offer letter content word count: {word_count}")

            # Encrypt the response content
            encrypted_content = encrypt_data({'generated_content': offer_letter_content})
            logger.info('Offer letter content generated successfully.')

            # Update or create the word count in the GuestLogin model
            if guest_email:
                if guest_record:
                    # Update existing word count
                    guest_record.word_count += word_count
                    guest_record.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {guest_record.word_count}')
                else:
                    # Create new record if no existing word count found
                    guest_record = GuestLogin.objects.create(
                        email=guest_email,
                        word_count=word_count
                    )
                    logger.info(f'Word count for {guest_email} saved successfully. Word count: {word_count}')
            else:
                logger.warning('No email found in the request data.')

            # Return the encrypted response with the generated content
            return JsonResponse({'encrypted_content': encrypted_content}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def sales_script_generator_guest(request):
    if request.method == 'POST':
        try:
            # Load and decode the request body
            body = request.body.decode('utf-8')
            logger.debug(f"Request body received: {body}")

            # Extract and decrypt the incoming payload
            data = json.loads(body)
            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                logger.warning("No encrypted content found in the request.")
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            logger.debug(f"Encrypted content received: {encrypted_content}")

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f"Decrypted content: {decrypted_content}")

            data = json.loads(decrypted_content)

            # Extract fields from the decrypted JSON data
            num_words = data.get('num_words')
            company_details = data.get('company_details')
            product_descriptions = data.get('product_descriptions')
            features_benefits = data.get('features_benefits')
            pricing_info = data.get('pricing_info')
            promotions = data.get('promotions')
            target_audience = data.get('target_audience')
            sales_objectives = data.get('sales_objectives')
            competitive_advantage = data.get('competitive_advantage')
            compliance = data.get('compliance')

            logger.debug(f"Data extracted for sales script generation: num_words={num_words}, company_details={company_details}")

            # Retrieve the guest email and existing word count if it exists
            guest_email = data.get('email')
            if guest_email:
                guest_record = GuestLogin.objects.filter(email=guest_email).first()
                if guest_record and guest_record.word_count >= 2000:
                    logger.info('Word Count limit already exceeded')
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            # Generate the sales script
            logger.info("Generating sales script...")
            sales_script = generate_sales_script(
                company_details,
                num_words,
                product_descriptions,
                features_benefits,
                pricing_info,
                promotions,
                target_audience,
                sales_objectives,
                competitive_advantage,
                compliance,
            )

            # Calculate the word count for the generated content
            word_count = len(sales_script.split())
            logger.debug(f"Generated sales script word count: {word_count}")

            # Encrypt the response content
            encrypted_response_content = encrypt_data({'generated_content': sales_script})
            logger.info("Sales script generated successfully.")

            # Update or create the word count in the GuestLogin model
            if guest_email:
                if guest_record:
                    # Update existing word count
                    guest_record.word_count += word_count
                    guest_record.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {guest_record.word_count}')
                else:
                    # Create new record if no existing word count found
                    guest_record = GuestLogin.objects.create(
                        email=guest_email,
                        word_count=word_count
                    )
                    logger.info(f'Word count for {guest_email} saved successfully. Word count: {word_count}')
            else:
                logger.warning('No email found in the request data.')

            # Return the encrypted response with the generated content
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format received.")
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f"ValueError occurred: {str(e)}")
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f"An unexpected error occurred: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def summarize_document_guest(request):
    if request.method == 'POST':
        try:
            # Extract form data
            document_context = request.POST.get('documentContext')
            main_subject = request.POST.get('mainSubject')
            summary_purpose = request.POST.get('summaryPurpose')
            length_detail = request.POST.get('lengthDetail')
            important_elements = request.POST.get('importantElements')
            audience = request.POST.get('audience')
            tone = request.POST.get('tone')
            format_ = request.POST.get('format')
            additional_instructions = request.POST.get('additionalInstructions')
            document_file = request.FILES.get('documentFile')
            guest_email = request.POST.get('email')  # Extract email from payload

            # Check if the file is provided
            if not document_file:
                return JsonResponse({'error': 'No document file provided.'}, status=400)

            # Retrieve the guest email and existing word count if it exists
            if guest_email:
                guest_record = GuestLogin.objects.filter(email=guest_email).first()
                if guest_record and guest_record.word_count >= 2000:
                    logger.info('Word Count limit already exceeded')
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            # Generate summary (ensure this function handles the file correctly)
            summary = generate_summary(
                document_context, main_subject, summary_purpose, length_detail,
                important_elements, audience, tone, format_, additional_instructions, document_file
            )

            # Handle errors in the summary generation process
            if summary.startswith("Error:"):
                logger.error(summary)
                return JsonResponse({'error': summary}, status=500)

            # Calculate the word count of the generated summary
            word_count = len(summary.split())
            logger.debug(f"Generated summary word count: {word_count}")

            # Encrypt and send back the generated summary
            # Update the word count for the guest email if provided
            if guest_email:
                if guest_record:
                    # Update the word count in the record
                    new_word_count = guest_record.word_count + word_count
                    guest_record.word_count = new_word_count
                    guest_record.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {new_word_count}')
                else:
                    logger.warning('No active session found for the given email.')

            # Return the generated summary in the response
            return JsonResponse({'summary': summary}, status=200)

        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def content_generator_guest(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')

        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")

        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        data = json.loads(decrypted_content)

        # Extract fields from the decrypted JSON data
        company_info = data.get('company_info')
        content_purpose = data.get('content_purpose')
        desired_action = data.get('desired_action')
        topic_details = data.get('topic_details')
        keywords = data.get('keywords')
        audience_profile = data.get('audience_profile')
        format_structure = data.get('format_structure')
        num_words = data.get('num_words')
        seo_keywords = data.get('seo_keywords')
        references = data.get('references')
        guest_email = data.get('email')  # Extract email from the payload

        logger.debug(f"Data extracted for content generation: company_info={company_info}, content_purpose={content_purpose}, desired_action={desired_action}")

        # Retrieve the guest email and existing word count if it exists
        if guest_email:
            guest_record = GuestLogin.objects.filter(email=guest_email).first()
            if guest_record and guest_record.word_count >= 2000:
                logger.info('Word Count limit already exceeded')
                return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

        # Generate the content
        logger.info("Generating content...")
        content = generate_content(
            company_info,
            content_purpose,
            desired_action,
            topic_details,
            keywords,
            audience_profile,
            format_structure,
            num_words,
            seo_keywords,
            references
        )

        if content:
            logger.info("Content generated successfully.")
            
            # Calculate word count of generated content
            word_count = len(content.split())
            logger.debug(f"Generated content word count: {word_count}")

            # Encrypt and send back the generated content
            # Update the word count for the guest email if provided
            if guest_email:
                if guest_record:
                    # Update the word count in the record
                    new_word_count = guest_record.word_count + word_count
                    guest_record.word_count = new_word_count

                    guest_record.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {new_word_count}')
                else:
                    logger.warning('No active session found for the given email.')

            # Encrypt and return the generated content in the response
            encrypted_response_content = encrypt_data({'generated_content': content})
            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        logger.error("Failed to generate content.")
        return JsonResponse({'error': 'Failed to generate content. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@require_POST
@csrf_exempt
def rephrasely_view_guest(request):
    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            request_body = json.loads(request.body.decode('utf-8'))
            encrypted_content = request_body.get('encrypted_content')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            # Extract required fields
            text_to_rephrase = data.get('text_to_rephrase')
            tone = data.get('tone')
            target_audience = data.get('target_audience')
            num_words = data.get('num_words', "default")  # Optional, default is "default"
            guest_email = data.get('email')  # Extract email from the payload

            # Retrieve the guest email and existing word count if it exists
            if guest_email:
                guest_record = GuestLogin.objects.filter(email=guest_email).first()
                if guest_record and guest_record.word_count >= 2000:
                    logger.info('Word Count limit already exceeded')
                    return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

            # Call the rephrasely function
            rephrased_text = rephrasely(text_to_rephrase, tone, target_audience, num_words)

            # Calculate word count of rephrased content
            word_count = len(rephrased_text.split())
            logger.debug(f'Rephrased content word count: {word_count}')

            # Restrict word count to a maximum of 2000
            if word_count > 2000:
                return JsonResponse({'error': 'The rephrased content exceeds the maximum allowed word count of 2000.'}, status=400)

            # Update the word count for the guest email if provided
            if guest_email:
                if guest_record:
                    # Update the word count in the record
                    new_word_count = guest_record.word_count + word_count
                    guest_record.word_count = new_word_count

                    guest_record.save()
                    logger.info(f'Word count for {guest_email} updated successfully. New word count: {new_word_count}')
                else:
                    logger.warning('No active session found for the given email.')

            # Encrypt and return the rephrased content in the response
            encrypted_response = encrypt_data({'rephrased_text': rephrased_text})
            logger.info('Rephrased content generated successfully.')

            # Return the encrypted response
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def generate_blog_view_guest(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract and decrypt the incoming payload
        data = json.loads(body)
        encrypted_content = data.get('encrypted_content')

        if not encrypted_content:
            logger.warning("No encrypted content found in the request.")
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        logger.debug(f"Encrypted content received: {encrypted_content}")
        decrypted_content = decrypt_data(encrypted_content)
        logger.debug(f"Decrypted content: {decrypted_content}")

        data = json.loads(decrypted_content)

        # Extract the required fields
        title = data.get('title')
        tone = data.get('tone')
        keywords = data.get('keywords', None)  # Optional
        guest_email = data.get('email')  # Extract email from the payload

        # Ensure required fields are present
        if not title or not tone:
            return JsonResponse({"error": "Missing 'title' or 'tone'."}, status=400)

        # Retrieve the guest email and existing word count if it exists
        if guest_email:
            guest_record = GuestLogin.objects.filter(email=guest_email).first()
            if guest_record and guest_record.word_count >= 2000:
                logger.info('Word count limit already exceeded')
                return JsonResponse({'error': 'Word Count limit already exceeded'}, status=400)

        # Call the generate_blog function
        blog_content = generate_blog(title, tone, keywords)

        # Calculate word count of generated blog content
        word_count = len(blog_content.split())
        logger.debug(f'Generated blog content word count: {word_count}')

        # Restrict word count to a maximum of 2000
        if word_count > 2000:
            return JsonResponse({'error': 'The generated blog content exceeds the maximum allowed word count of 2000.'}, status=400)

        # Update the word count for the guest email if provided
        if guest_email:
            if guest_record:
                # Update the word count in the record
                new_word_count = guest_record.word_count + word_count
                guest_record.word_count = new_word_count

                guest_record.save()
                logger.info(f'Word count for {guest_email} updated successfully. New word count: {new_word_count}')
            else:
                logger.warning('No active session found for the given email.')

        # Encrypt the response content
        encrypted_response_content = encrypt_data({'blog_content': blog_content})

        # Return the encrypted blog content as a JSON response
        return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=500)

    # If not a POST request, return an error
    return JsonResponse({"error": "Only POST method is allowed."}, status=405)

@csrf_exempt
@require_POST
def translate_content_guest(request):
    translated_content = None
    error = None
    language = ""

    if request.method == 'POST':
        try:
            # Extract and decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            logger.debug(f'Encrypted content received: {encrypted_content}')

            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            logger.debug(f'Decrypted content: {decrypted_content}')
            data = json.loads(decrypted_content)

            generated_content = data.get('generated_content')
            language = data.get('language')

            if not generated_content or not language:
                logger.warning('Both generated_content and language are required fields.')
                return JsonResponse({'error': 'Both generated_content and language are required fields.'}, status=400)

            logger.info(f'Translating content: {generated_content} to language: {language}')
            response = bhashini_translate(generated_content, language)

            if response["status_code"] == 200:
                translated_content = response["translated_content"]
                logger.info('Content translated successfully.')
                # Encrypt the response content
                encrypted_response = encrypt_data({
                    'generated_content': generated_content,
                    'translated_content': translated_content,
                    'selected_language': language
                })
                logger.debug(f'Encrypted response: {encrypted_response}')
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            else:
                logger.error('Translation failed with status code: {}'.format(response["status_code"]))
                return JsonResponse({'error': 'Translation failed.'}, status=500)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.warning(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Unexpected error: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.error('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)



@csrf_exempt
@require_POST
def guest_send_otp(request):
    if request.method == 'POST':
        try:
            # Load untrusted domains from the text file
            with open('./domains.txt', 'r') as file:
                untrusted_domains = {line.strip().lower() for line in file}

            # Decrypt incoming request data
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
            
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            email = data.get('email')  # Use email instead of mobile number

            if not email:
                return JsonResponse({'error': 'Email is required.'}, status=400)

            # Extract domain from the email
            try:
                email_domain = email.split('@')[1].lower()
            except IndexError:
                return JsonResponse({'error': 'Invalid email format.'}, status=400)


            # Check if the email domain is in the untrusted list
            if email_domain in untrusted_domains:
                return JsonResponse({
                    'error': 'It seems you are using an untrusted email domain service. Please try with another email.'}, 
                    status=400)

            # Check if the email is already registered in the GuestLogin table
            existing_guest = GuestLogin.objects.filter(email=email).first()
            if existing_guest:
                return JsonResponse({'error': 'This email is already registered. Please use a different email.'}, 
                                     status=400)

            # Generate OTP and set expiry time
            otp = guest_generate_otp()
            valid_till = guest_otp_expiry_time()

            # Create or update the GuestLogin entry
            guest_login, created = GuestLogin.objects.update_or_create(
                email=email,
                defaults={
                    'otp': otp,
                    'valid_till': valid_till,
                    'is_active': True
                }
            )

            # Prepare the email content
            subject = 'Confidential OTP for Guest Login'
            plain_message = f"""
Dear Sir/Madam,

We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is {otp} and will remain valid for a period of 10 minutes.

Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.

This notification is intended to provide you with the necessary information to complete Guest Login. If you have any concerns or require assistance, please contact our support team through the appropriate channels.

Thank you for your understanding and cooperation.

Sincerely,
The ProdigiDesk Team
"""
            html_message = f"""
<p>Dear Sir/Madam,</p>
<p>We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is <strong>{otp}</strong> and will remain valid for a period of 10 minutes.</p>
<p>Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.</p>
<p>This notification is intended to provide you with the necessary information to complete your Guest Login. If you have any concerns or require assistance, please contact our support team through the appropriate channels.</p>
<p>Thank you for your understanding and cooperation.</p>
<p>Sincerely,<br>The ProdigiDesk Team</p>
"""

            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [email]

            # Send the OTP via email
            try:
                send_mail(
                    subject, 
                    plain_message, 
                    from_email, 
                    recipient_list, 
                    fail_silently=False,
                    html_message=html_message  # Send both plain and HTML content
                )

                # Encrypt the response
                encrypted_response = encrypt_data({'message': f'OTP sent to {email}.'})
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            except Exception as e:
                return JsonResponse({'error': f'Error sending email: {str(e)}'}, status=500)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Invalid HTTP method.'}, status=405)


@csrf_exempt
@require_POST
def guest_validate_otp(request):
    if request.method == 'POST':
        try:
            # Load and decode the request body
            body = request.body.decode('utf-8')
            data = json.loads(body)

            encrypted_content = data.get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the incoming payload
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            email = data.get('email')
            otp = data.get('otp')

            if not email or not otp:
                return JsonResponse({'error': 'Email and OTP are required.'}, status=400)

            # Find the GuestLogin entry
            try:
                guest_login = GuestLogin.objects.get(email=email, otp=otp)
            except GuestLogin.DoesNotExist:
                return JsonResponse({'error': 'Invalid OTP.'}, status=400)

            # Check if OTP is valid and session is active
            if not guest_login.is_valid() or not guest_login.is_active:
                guest_login.deactivate_session()  # Deactivate if expired
                return JsonResponse({'error': 'OTP has expired or session is inactive.'}, status=400)

            # If valid, mark the session as inactive and return success
            guest_login.deactivate_session()
            # Encrypt the response content
            response_content = {'message': 'OTP is valid. Redirecting to homepage'}
            encrypted_response_content = encrypt_data(response_content)

            return JsonResponse({'encrypted_content': encrypted_response_content}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    else:
        return JsonResponse({'error': 'Invalid HTTP method.'}, status=405)


@csrf_exempt
@require_GET
def get_word_count(request):
    if request.method == 'GET':
        try:
            # Get encrypted content from query parameters
            encrypted_content = request.GET.get('encrypted_content', None)
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            # Decrypt the received content
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')

            # Retrieve email from decrypted data
            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email parameter is required.'}, status=400)

            # Fetch word count from GuestLogin model
            guest_login = GuestLogin.objects.filter(email=email).first()
            if guest_login:
                # Encrypt the response content
                encrypted_response = encrypt_data({'email': email, 'word_count': guest_login.word_count})
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)
            else:
                return JsonResponse({'error': 'No record found for the provided email.'}, status=404)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)
    else:
        logger.warning('Method not allowed.')
        return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_http_methods(["PUT"])
def create_cart(request):
    if request.method == 'PUT':  # Changed to PUT method for updating
        try:
            # Decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email is required'}, status=400)

            # Fetch selected services from the request
            selected_services = data.get('selected_services', [])

            # Check if the cart exists, otherwise create a new one
            cart, created = Cart.objects.get_or_create(email=email)

            # Update the cart services based on the selected services
            cart.email_service = 1 in selected_services
            cart.offer_letter_service = 2 in selected_services
            cart.business_proposal_service = 3 in selected_services
            cart.sales_script_service = 4 in selected_services
            cart.content_generation_service = 5 in selected_services
            cart.summarize_service = 6 in selected_services
            cart.ppt_generation_service = 7 in selected_services
            cart.blog_generation_service = 9 in selected_services
            cart.rephrasely_service = 10 in selected_services

            # Save the updated cart
            cart.save()

            # Prepare the response content
            response_data = {
                'email': cart.email,
                'selected_services': selected_services,
                'created_at': cart.created_at.isoformat(),
                'updated_at': cart.updated_at.isoformat(),
            }

            # Encrypt the response content
            encrypted_response = encrypt_data(response_data)

            # Return the encrypted response
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)  # Use 200 status code for successful updates

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
@require_http_methods(["PUT"])
def create_cart_yearly(request):
    if request.method == 'PUT':  # Changed to PUT method for updating
        try:
            # Decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email is required'}, status=400)

            # Fetch selected services from the request
            selected_services = data.get('selected_services', [])

            # Check if the cart exists, otherwise create a new one
            cart, created = YearlyCart.objects.get_or_create(email=email)

            # Update the cart services based on the selected services
            cart.email_service = 1 in selected_services
            cart.offer_letter_service = 2 in selected_services
            cart.business_proposal_service = 3 in selected_services
            cart.sales_script_service = 4 in selected_services
            cart.content_generation_service = 5 in selected_services
            cart.summarize_service = 6 in selected_services
            cart.ppt_generation_service = 7 in selected_services
            cart.blog_generation_service = 9 in selected_services
            cart.rephrasely_service = 10 in selected_services

            # Save the updated cart
            cart.save()

            # Prepare the response data
            response_data = {
                'email': cart.email,
                'selected_services': selected_services,
                'created_at': cart.created_at.isoformat(),
                'updated_at': cart.updated_at.isoformat(),
            }

            # Encrypt the response data
            encrypted_response = encrypt_data(response_data)

            return JsonResponse({'encrypted_content': encrypted_response}, status=200)  # Send back the encrypted response

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
@require_POST
def remove_service(request):
    if request.method == "POST":
        try:
            # Decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                logger.warning('No encrypted content found in the request.')
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)  # Implement your decrypt_data function
            data = json.loads(decrypted_content)
            logger.debug(f'Decrypted content: {data}')
            email = data.get("email")
            service_id = data.get("service_id")  # Change service to service_id

            if not email or service_id is None:  # Check for service_id being None
                return JsonResponse({"error": "Missing email or service ID"}, status=400)

            # Get the cart for the given email
            cart = get_object_or_404(Cart, email=email)

            # Map service IDs to their corresponding attribute names in the Cart model
            service_mapping = {
                1: 'email_service',
                2: 'offer_letter_service',
                3: 'business_proposal_service',
                4: 'sales_script_service',
                5: 'content_generation_service',
                6: 'summarize_service',
                7: 'ppt_generation_service',
                9: 'blog_generation_service',
                10: 'rephrasely_service',
            }

            # Get the service name corresponding to the given service ID
            service_name = service_mapping.get(service_id)

            if service_name is None:
                return JsonResponse({"error": f"Invalid service ID: {service_id}"}, status=400)

            # Set the service to False to remove it
            setattr(cart, service_name, False)
            cart.save()

            # Encrypt the response message
            response_message = {"message": f"{service_name} removed successfully"}
            encrypted_response = encrypt_data(response_message)  # Implement your encrypt_data function
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)
    else:
        logger.warning('Method not allowed.')
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
@require_POST
def remove_service_yearly(request):
    if request.method == "POST":
        try:
            # Extract and decrypt the encrypted content from the request
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)

            # Process the decrypted data
            email = data.get("email")
            service_id = data.get("service_id")

            if not email or service_id is None:
                return JsonResponse({"error": "Missing email or service ID"}, status=400)

            # Get the cart for the given email
            cart = get_object_or_404(YearlyCart, email=email)

            # Map service IDs to their corresponding attribute names in the Cart model
            service_mapping = {
                1: 'email_service',
                2: 'offer_letter_service',
                3: 'business_proposal_service',
                4: 'sales_script_service',
                5: 'content_generation_service',
                6: 'summarize_service',
                7: 'ppt_generation_service',
                9: 'blog_generation_service',
                10: 'rephrasely_service',
            }

            # Get the service name corresponding to the given service ID
            service_name = service_mapping.get(service_id)

            if service_name is None:
                return JsonResponse({"error": f"Invalid service ID: {service_id}"}, status=400)

            # Set the service to False to remove it
            setattr(cart, service_name, False)
            cart.save()

            # Prepare and encrypt the response content
            response_data = {"message": f"{service_name} removed successfully"}
            encrypted_response = encrypt_data(response_data)

            return JsonResponse({"encrypted_content": encrypted_response}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
        except ValueError as e:
            return JsonResponse({"error": str(e)}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    else:
        return JsonResponse({"error": "Invalid request method"}, status=405)

@csrf_exempt
@require_POST
def get_cart(request):
    if request.method == 'POST':
        try:
            # Decrypt the incoming payload
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            email = data.get('email')

            if email:
                # Get or create the cart for the provided email
                cart, created = Cart.objects.get_or_create(email=email)

                # Prepare the service details with IDs
                cart_services = {
                    "email_service": {
                        "id": 1,
                        "is_active": cart.email_service,
                    },
                    "offer_letter_service": {
                        "id": 2,
                        "is_active": cart.offer_letter_service,
                    },
                    "business_proposal_service": {
                        "id": 3,
                        "is_active": cart.business_proposal_service,
                    },
                    "sales_script_service": {
                        "id": 4,
                        "is_active": cart.sales_script_service,
                    },
                    "content_generation_service": {
                        "id": 5,
                        "is_active": cart.content_generation_service,
                    },
                    "summarize_service": {
                        "id": 6,
                        "is_active": cart.summarize_service,
                    },
                    "ppt_generation_service": {
                        "id": 7,
                        "is_active": cart.ppt_generation_service,
                    },
                    "blog_generation_service": {
                        "id": 9,
                        "is_active": cart.blog_generation_service,
                    },
                    "rephrasely_service": {
                        "id": 10,
                        "is_active": cart.rephrasely_service,
                    },
                }

                # Prepare the response data
                response_data = {
                    'email': cart.email,
                    'services': cart_services,
                    'created_at': cart.created_at.isoformat(),
                    'updated_at': cart.updated_at.isoformat(),
                }

                # Encrypt the response
                encrypted_response = encrypt_data(response_data)
                return JsonResponse({'encrypted_content': encrypted_response}, status=200)

            else:
                return JsonResponse({'error': 'Email parameter is missing in the request.'}, status=400)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format.'}, status=400)
        except ValueError as e:
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    elif request.method == 'GET':
        try:
            carts = Cart.objects.all()
            cart_list = []
            for cart in carts:
                cart_services = {
                    "email_service": {
                        "id": 1,
                        "is_active": cart.email_service,
                    },
                    "offer_letter_service": {
                        "id": 2,
                        "is_active": cart.offer_letter_service,
                    },
                    "business_proposal_service": {
                        "id": 3,
                        "is_active": cart.business_proposal_service,
                    },
                    "sales_script_service": {
                        "id": 4,
                        "is_active": cart.sales_script_service,
                    },
                    "content_generation_service": {
                        "id": 5,
                        "is_active": cart.content_generation_service,
                    },
                    "summarize_service": {
                        "id": 6,
                        "is_active": cart.summarize_service,
                    },
                    "ppt_generation_service": {
                        "id": 7,
                        "is_active": cart.ppt_generation_service,
                    },
                    "blog_generation_service": {
                        "id": 9,
                        "is_active": cart.blog_generation_service,
                    },
                    "rephrasely_service": {
                        "id": 10,
                        "is_active": cart.rephrasely_service,
                    },
                }

                cart_list.append({
                    'email': cart.email,
                    'services': cart_services,
                    'created_at': cart.created_at.isoformat(),
                    'updated_at': cart.updated_at.isoformat(),
                })

            # Encrypt the response
            encrypted_response = encrypt_data(cart_list)
            return JsonResponse({'encrypted_content': encrypted_response}, safe=False, status=200)

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    else:
        return JsonResponse({'error': 'Method not allowed.'}, status=405)


@csrf_exempt
@require_POST
def get_cart_yearly(request):
    try:
        # Step 1: Decrypt the incoming payload
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)

        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        email = data.get('email')

        if email:
            # Try to get or create the cart for the provided email
            cart, created = YearlyCart.objects.get_or_create(email=email)

            if created:
                cart.save()  # Save the newly created cart instance

            # Step 2: Prepare the cart services data
            cart_services = {
                "email_service": {"id": 1, "is_active": cart.email_service},
                "offer_letter_service": {"id": 2, "is_active": cart.offer_letter_service},
                "business_proposal_service": {"id": 3, "is_active": cart.business_proposal_service},
                "sales_script_service": {"id": 4, "is_active": cart.sales_script_service},
                "content_generation_service": {"id": 5, "is_active": cart.content_generation_service},
                "summarize_service": {"id": 6, "is_active": cart.summarize_service},
                "ppt_generation_service": {"id": 7, "is_active": cart.ppt_generation_service},
                "blog_generation_service": {"id": 9, "is_active": cart.blog_generation_service},
                "rephrasely_service": {"id": 10, "is_active": cart.rephrasely_service},
            }

            # Step 3: Encrypt the response data
            response_data = {
                'email': cart.email,
                'services': cart_services,
                'created_at': cart.created_at.isoformat(),
                'updated_at': cart.updated_at.isoformat(),
            }
            encrypted_response = encrypt_data(response_data)
            return JsonResponse({'encrypted_content': encrypted_response}, status=200)

        else:
            # If no email provided, fetch all carts and encrypt response
            carts = YearlyCart.objects.all()
            cart_list = []
            for cart in carts:
                cart_services = {
                    "email_service": {"id": 1, "is_active": cart.email_service},
                    "offer_letter_service": {"id": 2, "is_active": cart.offer_letter_service},
                    "business_proposal_service": {"id": 3, "is_active": cart.business_proposal_service},
                    "sales_script_service": {"id": 4, "is_active": cart.sales_script_service},
                    "content_generation_service": {"id": 5, "is_active": cart.content_generation_service},
                    "summarize_service": {"id": 6, "is_active": cart.summarize_service},
                    "ppt_generation_service": {"id": 7, "is_active": cart.ppt_generation_service},
                    "blog_generation_service": {"id": 9, "is_active": cart.blog_generation_service},
                    "rephrasely_service": {"id": 10, "is_active": cart.rephrasely_service},
                }
                cart_list.append({
                    'email': cart.email,
                    'services': cart_services,
                    'created_at': cart.created_at.isoformat(),
                    'updated_at': cart.updated_at.isoformat(),
                })

            encrypted_response = encrypt_data(cart_list)
            return JsonResponse({'encrypted_content': encrypted_response}, safe=False, status=200)

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f'ValueError: {str(e)}')
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f'Exception: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_POST
def empty_cart(request):
    try:
        # Retrieve and decrypt the incoming encrypted content
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            logger.warning('No encrypted content found in the request.')
            return JsonResponse({'error': 'No encrypted content found in the request.'}, status=400)
        
        # Decrypt the incoming content
        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        email = data.get('email')

        if not email:
            return JsonResponse({'error': 'Email parameter is required'}, status=400)

        # Get or create a Cart for the provided email
        cart, created = Cart.objects.get_or_create(email=email)

        # Set all service flags to False
        cart.email_service = False
        cart.offer_letter_service = False
        cart.business_proposal_service = False
        cart.sales_script_service = False
        cart.content_generation_service = False
        cart.summarize_service = False
        cart.ppt_generation_service = False
        cart.blog_generation_service = False
        cart.rephrasely_service = False

        # Save the cart with updated values
        cart.save()

        # Prepare response data
        response_data = {
            'email': cart.email,
            'services': {
                "email_service": cart.email_service,
                "offer_letter_service": cart.offer_letter_service,
                "business_proposal_service": cart.business_proposal_service,
                "sales_script_service": cart.sales_script_service,
                "content_generation_service": cart.content_generation_service,
                "summarize_service": cart.summarize_service,
                "ppt_generation_service": cart.ppt_generation_service,
                "blog_generation_service": cart.blog_generation_service,
                "rephrasely_service": cart.rephrasely_service,
            },
            'created_at': cart.created_at.isoformat(),
            'updated_at': cart.updated_at.isoformat(),
        }

        # Encrypt the response data
        encrypted_response = encrypt_data(response_data)
        return JsonResponse({'encrypted_content': encrypted_response}, status=200)

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f'ValueError: {str(e)}')
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f'Exception: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_POST
def empty_cart_yearly(request):
    try:
        # Get the encrypted content from the request
        encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
        if not encrypted_content:
            return JsonResponse({"error": "No encrypted content found in the request."}, status=400)

        # Decrypt the request data
        decrypted_content = decrypt_data(encrypted_content)
        data = json.loads(decrypted_content)
        email = data.get('email')

        if not email:
            encrypted_response = encrypt_data({"error": "Email parameter is required"})
            return JsonResponse({"encrypted_content": encrypted_response}, status=400)

        # Get or create a Cart for the provided email
        cart, created = YearlyCart.objects.get_or_create(email=email)

        # Set all service flags to False
        cart.email_service = False
        cart.offer_letter_service = False
        cart.business_proposal_service = False
        cart.sales_script_service = False
        cart.content_generation_service = False
        cart.summarize_service = False
        cart.ppt_generation_service = False
        cart.blog_generation_service = False
        cart.rephrasely_service = False

        # Save the cart with updated values
        cart.save()

        # Encrypt the response data
        response_data = {
            'email': cart.email,
            'services': {
                "email_service": cart.email_service,
                "offer_letter_service": cart.offer_letter_service,
                "business_proposal_service": cart.business_proposal_service,
                "sales_script_service": cart.sales_script_service,
                "content_generation_service": cart.content_generation_service,
                "summarize_service": cart.summarize_service,
                "ppt_generation_service": cart.ppt_generation_service,
                "blog_generation_service": cart.blog_generation_service,
                "rephrasely_service": cart.rephrasely_service,
            },
            'created_at': cart.created_at.isoformat(),
            'updated_at': cart.updated_at.isoformat(),
        }
        encrypted_response = encrypt_data(response_data)

        return JsonResponse({"encrypted_content": encrypted_response}, status=200)

    except json.JSONDecodeError:
        encrypted_response = encrypt_data({"error": "Invalid JSON format. Please provide valid JSON data."})
        return JsonResponse({"encrypted_content": encrypted_response}, status=400)
    except Exception as e:
        encrypted_response = encrypt_data({"error": str(e)})
        return JsonResponse({"encrypted_content": encrypted_response}, status=500)


@csrf_exempt
@require_http_methods(["DELETE"])
@api_view(['DELETE'])
@permission_classes([IsAuthenticated,HasAPIKey])
def delete_user_account(request):
    if request.method == "DELETE":
        try:
            # Get the encrypted content from the request
            encrypted_content = json.loads(request.body.decode('utf-8')).get('encrypted_content')
            if not encrypted_content:
                return JsonResponse({"error": "No encrypted content found in the request."}, status=400)

            # Decrypt the request data
            decrypted_content = decrypt_data(encrypted_content)
            data = json.loads(decrypted_content)
            email = data.get('userEmail')

            if not email:
                return JsonResponse({"error": "Email is required"}, status=400)

            # Find the user based on the email
            try:
                user = User.objects.get(email=email)
            except User.DoesNotExist:
                encrypted_response = encrypt_data({"error": "User not found"})
                return JsonResponse({"encrypted_content": encrypted_response}, status=404)

            # Delete associated records
            UserService.objects.filter(user=user).delete()  # Delete user services
            Profile.objects.filter(user=user).delete()      # Delete user profile
            Payment.objects.filter(email=email).delete()    # Delete payments associated with the user
            UserSession.objects.filter(user=user).delete()  # Delete user sessions
            EmailVerificationOTP.objects.filter(user=user).delete()  # Delete email verification OTPs
            TemporaryEmailVerificationOTP.objects.filter(email=email).delete()  # Delete temp email verifications
            Cart.objects.filter(email=email).delete()  # Delete the cart

            # Finally, delete the user
            user.delete()

            # Send account deletion email
            subject = "Account Deletion Confirmation"
            message = f"""
Dear {email},

We regret to inform you that your account with us has been deleted successfully as per your request. 

All your associated data, including services, profiles, and payment records, have been removed from our system. Please note that this action is irreversible.

If you believe this was a mistake or have any questions, feel free to contact our support team for assistance.

Thank you for using our services.

Sincerely,
The ProdigiDesk Team
"""
            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [email]

            try:
                send_mail(subject, message, from_email, recipient_list, fail_silently=False)
                logger.info(f"Account deletion email sent to {email}")
            except Exception as e:
                logger.error(f"Error sending account deletion email: {str(e)}")
                encrypted_response = encrypt_data({"error": f"Error sending email: {str(e)}"})
                return JsonResponse({"encrypted_content": encrypted_response}, status=500)

            # Encrypt the success message
            encrypted_response = encrypt_data({"message": "User account and associated data deleted successfully"})
            return JsonResponse({"encrypted_content": encrypted_response}, status=200)

        except json.JSONDecodeError:
            encrypted_response = encrypt_data({"error": "Invalid JSON format. Please provide valid JSON data."})
            return JsonResponse({"encrypted_content": encrypted_response}, status=400)
        except Exception as e:
            encrypted_response = encrypt_data({"error": str(e)})
            return JsonResponse({"encrypted_content": encrypted_response}, status=500)

    encrypted_response = encrypt_data({"error": "Invalid request method"})
    return JsonResponse({"encrypted_content": encrypted_response}, status=400)



#Mobile App API's
@csrf_exempt
@require_POST
def signin_android(request):
    if request.method == 'POST':
        try:
            # Parse the request body as JSON
            body = json.loads(request.body.decode('utf-8'))
            login_input = body.get('login_input', '').lower()
            password = body.get('password')
            print(body)
            logout_from_all = body.get('logout_from_all', False)  # Check if the checkbox is set

            if not login_input or not password:
                logger.warning('Login input and password are required')
                return JsonResponse({'error': 'Login input and password are required'}, status=400)

            try:
                if '@' in login_input:
                    user = User.objects.get(email=login_input)
                else:
                    user = User.objects.get(username=login_input)
            except User.DoesNotExist:
                logger.warning('Username or email not found')
                return JsonResponse({'error': 'Username or email not found'}, status=401)

            # Check if the user has an active session
            if logout_from_all:
                # Mark all previous sessions as inactive
                UserSession.objects.filter(user=user, active=True).update(active=False)

            active_session = UserSession.objects.filter(user=user, active=True).first()
            if active_session and not logout_from_all:
                logger.warning('User already logged in with an active session')
                return JsonResponse({'error': 'User already logged in'}, status=403)

            user = authenticate(request, username=user.username, password=password)
            if user is not None:
                login(request, user)
                refresh = RefreshToken.for_user(user)
                logger.info(f'User {user.username} authenticated successfully')

                # Create a new session
                session_id = get_random_string(length=32)
                UserSession.objects.create(user=user, session_id=session_id, email=user.email, active=True)  # Include email here

                logger.debug(f'Session created with ID: {session_id}')

                response_data = {
                    'success': 'User authenticated',
                    'access': str(refresh.access_token),
                    'refresh': str(refresh),
                    'user_id': user.id,
                    'session_id': session_id  # Include session_id in the response
                }
                print(response_data)
                return JsonResponse(response_data, status=200)
            else:
                logger.warning('Password not correct')
                return JsonResponse({'error': 'Password not correct'}, status=401)
        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            return JsonResponse({'error': 'Internal server error'}, status=500)
    else:
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)


@csrf_exempt
@require_POST
def create_razorpay_order_android(request):
    if request.method == "POST":
        try:
            # Load the request body
            data = json.loads(request.body.decode('utf-8'))

            # Extract amount and email
            amount = data.get('amount', 0)  # Get the amount in rupees
            email = data.get('email')  # Extract email from the request
            
            if not email:
                return JsonResponse({"error": "Email is required"}, status=400)

            # Convert amount to paise
            amount_in_paise = int(amount * 100)

            # Create Razorpay order
            razorpay_order = razorpay_client.order.create({
                "amount": amount_in_paise,  # Use amount in paise
                "currency": "INR",
                "payment_capture": "1"
            })

            print(razorpay_order)

            # Save order details to the Payment table including the email
            Payment.objects.create(
                order_id=razorpay_order['id'],
                amount=amount,  # Store amount in rupees in the database
                currency="INR",
                payment_capture=True,
                email=email  # Store the email
            )

            # Prepare the response data
            response_data = {
                "order_id": razorpay_order['id'],
                "amount": amount,  # Return amount in rupees for response
                "currency": "INR",
                "razorpay_key_id": settings.RAZORPAY_KEY_ID
            }
            
            return JsonResponse(response_data, status=200)

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON format."}, status=400)
        except ValueError as e:
            return JsonResponse({"error": str(e)}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)

@csrf_exempt
def verify_payment_android(request):
    try:
        # Load the request body
        data = json.loads(request.body.decode('utf-8'))

        razorpay_order_id = data.get('razorpay_order_id')
        razorpay_payment_id = data.get('razorpay_payment_id')
        razorpay_signature = data.get('razorpay_signature')
        selected_services = data.get('selected_services')
        email = data.get('email')

        if not razorpay_order_id or not razorpay_payment_id or not razorpay_signature:
            return JsonResponse({'error': 'Missing required payment data'}, status=400)

        logger.info(f"Received payment verification request with order_id: {razorpay_order_id}, payment_id: {razorpay_payment_id}, signature: {razorpay_signature}")

        # Verify payment signature
        params_dict = {
            'razorpay_order_id': razorpay_order_id,
            'razorpay_payment_id': razorpay_payment_id,
            'razorpay_signature': razorpay_signature
        }

        try:
            # Verify the payment signature
            razorpay_client.utility.verify_payment_signature(params_dict)
            logger.info("Payment signature verification successful")

            # Fetch and update the payment record
            payment = Payment.objects.get(order_id=razorpay_order_id)
            payment.payment_id = razorpay_payment_id
            payment.signature = razorpay_signature
            payment.email = email
            payment.verified = True

            # Process selected services
            if not selected_services or not email:
                return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

            user = get_object_or_404(User, email=email)
            user_services, created = UserService.objects.get_or_create(user=user)
            subscribed_services = []

            # Check if "Introductory Offer" is selected
            if selected_services.get("introductory_offer_service", False):
                user_services.email_service = user_services.offer_letter_service = 1
                user_services.business_proposal_service = user_services.sales_script_service = 1
                user_services.content_generation_service = user_services.summarize_service = 1
                user_services.ppt_generation_service = user_services.blog_generation_service = 1
                user_services.rephrasely_service = 1
                subscribed_services = [
                    "Email Service", "Offer Letter Service", "Business Proposal Service",
                    "Sales Script Service", "Content Generation Service", "Summarize Service",
                    "PPT Generation Service", "Blog Generation Service", "Rephrasely Service"
                ]
            else:
                # Update services based on data
                for service_key, service_attr in [
                    ("email_service", "Email Service"), ("offer_letter_service", "Offer Letter Service"),
                    ("business_proposal_service", "Business Proposal Service"), ("sales_script_service", "Sales Script Service"),
                    ("content_generation_service", "Content Generation Service"), ("summarize_service", "Summarize Service"),
                    ("ppt_generation_service", "PPT Generation Service"), ("blog_generation_service", "Blog Generation Service"),
                    ("rephrasely_service", "Rephrasely Service")
                ]:
                    if selected_services.get(service_key, 0) > 0:
                        setattr(user_services, service_key, 1)
                        subscribed_services.append(service_attr)

            # Save the updated services
            user_services.save()
            payment.order_datetime = datetime.now()
            payment.subscribed_services = selected_services
            payment.service = user_services
            payment.subscription_duration = 'monthly'
            payment.save()

            # Send subscription confirmation email
            subject = 'Subscription Confirmation - ProdigiDesk Services'
            services_list = ''.join(f"<li>{service}</li>" for service in subscribed_services)
            message = f"""
            <html><body>
            <p>Dear {user.get_full_name()},</p>
            <p>We are pleased to inform you that your purchase has been successfully processed. Below is a summary of the services activated for you:</p>
            <p><strong>Activated Services for a month:</strong></p>
            <ul>{services_list}</ul>
            <p><strong>Order Details:</strong></p>
            <ul>
                <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                <li><strong>Order Date and Time:</strong> {payment.order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                <li><strong>Registered Email:</strong> {email}</li>
            </ul>
            <p>Should you have any queries or require assistance, feel free to contact us at contact@espritanalytique.com. We look forward to serving you!</p>
            <p>Thank you for choosing ProdigiDesk.</p>
            <br>Best regards
            <br>The ProdigiDesk Team
            <br><a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>   
            </body></html>
            """

            email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
            email_message.content_subtype = 'html'
            email_message.send(fail_silently=False)

            # Return success response
            response_data = {'message': 'Payment and service save successful'}
            return JsonResponse(response_data, status=200)

        except razorpay.errors.SignatureVerificationError:
            logger.error("Payment signature verification failed")
            return JsonResponse({'error': 'Payment verification failed'}, status=400)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format")
        return JsonResponse({'error': 'Invalid JSON format'}, status=400)
    except Exception as e:
        logger.error(f"Exception occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=400)

@csrf_exempt
@require_POST
def verify_payment_yearly_android(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Content received: {data}')

            # Extract payment details
            razorpay_order_id = data.get('razorpay_order_id')
            razorpay_payment_id = data.get('razorpay_payment_id')
            razorpay_signature = data.get('razorpay_signature')
            selected_services = data.get('selected_services')
            email = data.get('email')

            logger.info(f"Received yearly payment verification request with order_id: {razorpay_order_id}")

            # Verify payment signature
            params_dict = {
                'razorpay_order_id': razorpay_order_id,
                'razorpay_payment_id': razorpay_payment_id,
                'razorpay_signature': razorpay_signature
            }

            try:
                razorpay_client.utility.verify_payment_signature(params_dict)
                logger.info("Payment signature verification successful for yearly subscription")

                # Update payment and user service details
                payment = Payment.objects.get(order_id=razorpay_order_id)
                payment.payment_id = razorpay_payment_id
                payment.signature = razorpay_signature
                payment.email = email
                payment.verified = True

                if not selected_services or not email:
                    return JsonResponse({'error': 'No services or email found in the request.'}, status=400)

                user = get_object_or_404(User, email=email)
                user_services, created = UserService.objects.get_or_create(user=user)
                subscribed_services = []
                expiration_date = timezone.now().date() + relativedelta(years=1)  # Default expiration date set to 1 year

                # Check if the introductory offer is selected
                if selected_services.get("introductory_offer_service", False):
                    # Activate multiple services as part of the introductory offer
                    services = [
                        "email_service", "offer_letter_service", "business_proposal_service",
                        "sales_script_service", "content_generation_service", "summarize_service",
                        "ppt_generation_service", "blog_generation_service", "rephrasely_service"
                    ]
                    for service in services:
                        setattr(user_services, service, 1)
                        setattr(user_services, f"{service.replace('_service', '')}_end_date", expiration_date)
                    subscribed_services = [service.replace("_service", "").replace("_", " ").title() for service in services]
                else:
                    # Activate individual services based on selection
                    for service, end_date_field in [
                        ("email_service", "email_end_date"),
                        ("offer_letter_service", "offer_letter_end_date"),
                        ("business_proposal_service", "business_proposal_end_date"),
                        ("sales_script_service", "sales_script_end_date"),
                        ("content_generation_service", "content_generation_end_date"),
                        ("summarize_service", "summarize_end_date"),
                        ("ppt_generation_service", "ppt_generation_end_date"),
                        ("blog_generation_service", "blog_generation_end_date"),
                        ("rephrasely_service", "rephrasely_end_date"),
                    ]:
                        if selected_services.get(service, 0) > 0:
                            setattr(user_services, service, 1)
                            setattr(user_services, end_date_field, expiration_date)
                            subscribed_services.append(service.replace("_service", "").replace("_", " ").title())

                user_services.save()
                payment.order_datetime = datetime.now()
                payment.subscribed_services = selected_services
                payment.subscription_duration = 'yearly'  # Set subscription_duration as yearly
                payment.service = user_services
                payment.save()

                # Send subscription confirmation email
                subject = 'Subscription Confirmation - ProdigiDesk Services'
                services_list = ''.join(f"<li>{service}</li>" for service in subscribed_services)
                message = f"""
                <html><body>
                <p>Dear {user.get_full_name()},</p>
                <p>We are pleased to inform you that your subscription has been successfully processed. Below is a summary of the services activated for you:</p>
                <p><strong>Activated Services for a year:</strong></p>
                <ul>{services_list}</ul>
                <p><strong>Order Details:</strong></p>
                <ul>
                    <li><strong>Order Number:</strong> {razorpay_order_id}</li>
                    <li><strong>Order Date and Time:</strong> {payment.order_datetime.strftime("%Y-%m-%d %H:%M:%S")}</li>
                    <li><strong>Payment Amount:</strong> {payment.amount} {payment.currency}</li>
                    <li><strong>Registered Email:</strong> {email}</li>
                </ul>
                <p>Should you have any queries or require assistance, feel free to contact us at contact@espritanalytique.com. We look forward to serving you!</p>
                <p>Thank you for choosing ProdigiDesk.</p>
                <br>
                <br>Best regards
                <br>The ProdigiDesk Team
                <br><a href="http://www.prodigidesk.ai/">http://www.prodigidesk.ai/</a>   
                
                </body></html>
                """

                email_message = EmailMessage(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
                email_message.content_subtype = 'html'
                email_message.send(fail_silently=False)

                # Return success response
                return JsonResponse({
                    'message': 'Yearly payment and service save successful',
                    'subscribed_services': subscribed_services,
                    'subscription_duration': 'yearly'
                }, status=200)

            except razorpay.errors.SignatureVerificationError:
                logger.error("Yearly payment signature verification failed")
                return JsonResponse({"status": "Payment verification failed"}, status=400)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format")
            return JsonResponse({"error": "Invalid JSON format"}, status=400)
        except Exception as e:
            logger.error(f"Exception occurred: {str(e)}")
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request method"}, status=400)


@csrf_exempt
@require_POST
def add_user_android(request):
    if request.method == 'POST':
        try:
            # Load untrusted domains from the text file
            with open('./domains.txt', 'r') as file:
                untrusted_domains = {line.strip().lower() for line in file}

            # Load and decode the request body
            body = request.body.decode('utf-8')
            logger.debug(f"Request body received: {body}")

            # Parse the JSON request body
            data = json.loads(body)
            
            # Extract required fields
            first_name = data.get('first_name')
            last_name = data.get('last_name')
            username = data.get('username')
            email = data.get('email')
            password = data.get('password')
            confirm_password = data.get('confirm_password')

            # Check if username and email are provided
            if not username:
                return JsonResponse({'error': 'Username is required.'}, status=400)
            if not email:
                return JsonResponse({'error': 'Email is required.'}, status=400)

            # Normalize username and email to lowercase
            username = username.lower()
            email = email.lower()

            # Extract domain from the email
            try:
                email_domain = email.split('@')[1].lower()
            except IndexError:
                return JsonResponse({'error': 'Invalid email format.'}, status=400)

            # Check if the email domain is in the untrusted list
            if email_domain in untrusted_domains:
                return JsonResponse({
                    'error': 'It seems you are using an untrusted email domain service. Please try with another email.'}, 
                    status=400)

            # Check if passwords match
            if password != confirm_password:
                return JsonResponse({'error': 'Passwords do not match.'}, status=400)

            # Check if username already exists
            if User.objects.filter(username=username).exists():
                return JsonResponse({'error': 'Username already exists.'}, status=400)

            # Validate email
            try:
                validate_email(email)
            except ValidationError:
                return JsonResponse({'error': 'Invalid email address.'}, status=400)

            # Check if email already exists
            if User.objects.filter(email=email).exists():
                return JsonResponse({'error': 'Email already exists.'}, status=400)

            # Create user
            user = User.objects.create(
                first_name=first_name,
                last_name=last_name,
                username=username,
                email=email,
                password=make_password(password)  # Hash the password
            )
            user.save()

            # Prepare the response data
            response_data = {
                'message': 'User created successfully',
                'user_id': user.id,
                'email': email
            }

            return JsonResponse(response_data, status=201)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format in request")
            return JsonResponse({'error': 'Invalid JSON format'}, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)
    else:
        logger.error("Invalid request method")
        return JsonResponse({'error': 'Invalid request method'}, status=405)


@csrf_exempt
def logout_view_android(request):
    try:
        if not request.user.is_authenticated:
            # If the user is not authenticated, return an error message
            print("Attempted logout by an unauthenticated user.")
            return JsonResponse({'error': 'User is not authenticated'}, status=401)

        # Fetch the active session for the logged-in user
        user_session = UserSession.objects.filter(user=request.user, active=True).first()
        print(user_session)
        if user_session:
            # Set the active field to False
            user_session.active = False
            user_session.save()
            print(f"Session for user {request.user.username} marked as inactive.")
        else:
            print(f"No active session found for user {request.user.username}.")

        # Perform Django logout operation
        logout(request)
        print(f"User {request.user.username} logged out successfully.")
        logger.info(f"User {request.user.username} logged out successfully.")
        return JsonResponse({'success': 'Logged out successfully'}, status=200)
    except Exception as e:
        print(f"Error during logout: {str(e)}")
        logger.error(f"Error during logout: {str(e)}")
        return JsonResponse({'error': 'An error occurred during logout.'}, status=500)


@csrf_exempt
@require_POST
def send_email_verification_otp_android(request):
    if request.method == 'POST':
        try:
            # Load untrusted domains from the text file
            with open('./domains.txt', 'r') as file:
                untrusted_domains = {line.strip().lower() for line in file}

            # Load and decode the request body
            body = request.body.decode('utf-8')
            logger.debug(f"Request body received: {body}")

            data = json.loads(body)

            # Extract email from the payload
            email = data.get('email')
            if not email:
                return JsonResponse({'error': 'Email is required'}, status=400)

            # Extract domain from the email
            try:
                email_domain = email.split('@')[1].split('.')[0].lower()
            except IndexError:
                return JsonResponse({'error': 'Invalid email format.'}, status=400)

            # Check if the email domain is in the untrusted list
            if email_domain in untrusted_domains:
                return JsonResponse({
                    'error': 'It seems you are using an untrusted email domain service. Please try with another email.'},
                    status=400)

            # Check if the email is already registered
            if User.objects.filter(email=email).exists():
                return JsonResponse({'error': 'Email is already registered'}, status=400)

            # Generate OTP and set expiry time
            otp = generate_otp()
            expiry_time = timezone.now() + timedelta(minutes=10)

            # Save OTP and its expiry in the database
            TemporaryEmailVerificationOTP.objects.update_or_create(
                email=email,
                defaults={
                    'otp': otp,
                    'expiry_time': expiry_time
                }
            )

            # Send OTP via email
            subject = 'Welcome to ProdigiDesk'
            plain_message = f"""
Dear Sir/Madam,

We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is {otp} and will remain valid for a period of 10 minutes.

Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.

This notification is intended to provide you with the necessary information to complete your Email Verification. If you have any concerns or require assistance, please contact our support team through the appropriate channels.

Thank you for your understanding and cooperation.

Sincerely,
The ProdigiDesk Team
"""
            html_message = f"""
<p>Dear Sir/Madam,</p>
<p>We are writing to inform you that a confidential One-Time Password (OTP) has been generated by our system. The OTP is <strong>{otp}</strong> and will remain valid for a period of 10 minutes.</p>
<p>Please be advised that this email has been generated automatically by our system and does not require a response. We kindly request that you refrain from replying to this email.</p>
<p>This notification is intended to provide you with the necessary information to complete your Email Verification. If you have any concerns or require assistance, please contact our support team through the appropriate channels.</p>
<p>Thank you for your understanding and cooperation.</p>
<p>Sincerely,<br>The ProdigiDesk Team</p>
"""

            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [email]

            try:
                send_mail(
                    subject, 
                    plain_message, 
                    from_email, 
                    recipient_list, 
                    fail_silently=False,
                    html_message=html_message  # Add HTML content here
                )
            except Exception as e:
                return JsonResponse({'error': f'Error sending email: {str(e)}'}, status=500)

            # Return success response
            return JsonResponse({'success': 'OTP sent successfully'}, status=200)

        except json.JSONDecodeError:
            logger.error("Invalid JSON format received.")
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            logger.error(f"An unexpected error occurred: {str(e)}")
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=405)

@csrf_exempt
@require_POST
def otp_verify_android(request):
    if request.method == 'POST':
        try:
            # Load and decode the request body
            body = request.body.decode('utf-8')
            data = json.loads(body)

            # Extract OTP and email from the request payload
            otp = data.get('otp')
            email = data.get('email')

            if not otp or not email:
                return JsonResponse({'error': 'OTP and email are required'}, status=400)

            # Fetch the OTP entry from the database
            try:
                otp_entry = TemporaryEmailVerificationOTP.objects.get(email=email)
            except TemporaryEmailVerificationOTP.DoesNotExist:
                return JsonResponse({'error': 'Incorrect OTP'}, status=400)

            # Check if the OTP is correct and not expired
            if otp_entry.otp != otp:
                return JsonResponse({'error': 'Incorrect OTP'}, status=400)

            if otp_entry.expiry_time < timezone.now():
                return JsonResponse({'error': 'OTP Expired'}, status=400)

            # OTP is correct and not expired
            return JsonResponse({'success': 'OTP verified successfully'}, status=200)

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

    return JsonResponse({'error': 'Invalid request method'}, status=405)

@csrf_exempt
@require_POST
def check_session_status_android(request):
    if request.method == 'POST':
        try:
            # Parse incoming payload
            data = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Request payload: {data}')

            # Extract session ID from the request payload
            session_id = data.get('session_id')
            print(session_id)
            if not session_id:
                logger.warning('Session ID not provided')
                return JsonResponse({'error': 'Session ID not provided'}, status=400)

            session = UserSession.objects.filter(session_id=session_id).first()

            if not session:
                logger.warning('Session not found')
                return JsonResponse({'error': 'Session not found'}, status=404)

            # Log the session status
            logger.info(f'Session {session_id} status is {session.active}')

            # Return the session status
            return JsonResponse({
                'session_id': session_id,
                'active': session.active
            }, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format in request')
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        except Exception as e:
            logger.error(f'Internal server error: {str(e)}')
            return JsonResponse({'error': 'Internal server error'}, status=500)

    else:
        logger.error('Invalid request method')
        return JsonResponse({'error': 'Invalid request method'}, status=405)

@api_view(['POST'])
@permission_classes([])
def speech_api_android(request):
    if request.method == 'POST':
        try:
            # Parse the JSON data from the request body
            data = json.loads(request.body)
            text = data.get('text', 'No text provided')
            
            # Generate speech from text using gTTS
            tts = gTTS(text=text, lang='en', slow=False)
 
            # In-memory file object for the audio
            mp3_fp = io.BytesIO()
            tts.write_to_fp(mp3_fp)
 
            mp3_fp.seek(0)  # Reset the pointer to the beginning of the file
 
            # Prepare the HTTP response with the audio file
            response = HttpResponse(mp3_fp, content_type='audio/mpeg')
            response['Content-Disposition'] = 'attachment; filename="speech.mp3"'
 
            return response
       
        except Exception as e:
            # Handle any exception and return error message
            return JsonResponse({'status': 'error', 'message': str(e)})
   
    else:
        # Handle non-POST requests
        return JsonResponse({'status': 'error', 'message': 'Only POST requests are allowed.'})
    
@api_view(['GET', 'POST'])
@permission_classes([])
def profile_android(request):
    user = request.user
    profile = Profile.objects.get(user=user)
    errors = []

    if request.method == 'POST':
        try:
            # Parse the incoming JSON payload
            data = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Received content: {data}')

            print(data)
            # Update user and profile data based on received JSON
            user.first_name = data.get('first_name', user.first_name)
            user.last_name = data.get('last_name', user.last_name)
            user.email = data.get('email', user.email)
            profile.bio = data.get('bio', profile.bio)
            profile.location = data.get('location', profile.location)
            profile.user_gst = data.get('user_gst', profile.user_gst)

            birth_date = data.get('birth_date')
            if birth_date:
                parsed_date = parse_date(birth_date)
                if parsed_date:
                    profile.birth_date = parsed_date
                else:
                    errors.append("Invalid date format for birth date.")
                    profile.birth_date = None

            if not user.first_name:
                errors.append("First name is required.")
            if not user.last_name:
                errors.append("Last name is required.")
            if not user.email:
                errors.append("Email is required.")

            if not errors:
                user.save()
                profile.save()
                response_data = {'message': 'Profile updated successfully.'}
            else:
                response_data = {'errors': errors}

            logger.info('Profile updated successfully.')

            # Return the response in plain format (no encryption)
            return JsonResponse(response_data, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    # Handle GET request
    response_data = {
        'user': {
            'first_name': user.first_name,
            'last_name': user.last_name,
            'email': user.email
        },
        'profile': {
            'bio': profile.bio,
            'location': profile.location,
            'birth_date': profile.birth_date.isoformat() if profile.birth_date else None,
            'user_gst': profile.user_gst
        }
    }

    logger.info('Profile data retrieved successfully.')

    # Return the response in plain format (no encryption)
    return JsonResponse(response_data)

@api_view(['GET'])
@permission_classes([])
def profile_info_android(request):
    try:
        # Get the email from the query parameters
        email = request.GET.get('email')
        if not email:
            return JsonResponse({'error': 'Email parameter is required.'}, status=400)

        # Fetch the user based on the provided email
        user = User.objects.get(email=email)

        # Get the user's subscribed services
        user_service = UserService.objects.get(user=user)

        # Get all payments made by the user
        payments = Payment.objects.filter(email=user.email)

        # Prepare the payment details
        payment_info = []
        for payment in payments:
            payment_info.append({
                'order_id': payment.order_id,
                'payment_id': payment.payment_id,
                'amount': str(payment.amount),
                'currency': payment.currency,
                'created_at': payment.created_at.isoformat(),  # Convert to ISO string
                'verified': payment.verified,
            })

        # Prepare the response data with date fields converted to ISO format
        response_data = {
            'user_info': {
                'first_name': user.first_name,
                'last_name': user.last_name,
                'email': user.email,
                'last_login': user.last_login.isoformat() if user.last_login else None,
                'username': user.username,
                'date_joined': user.date_joined.isoformat(),
            },
            'services': {
                'email_service': user_service.email_service,
                'offer_letter_service': user_service.offer_letter_service,
                'business_proposal_service': user_service.business_proposal_service,
                'sales_script_service': user_service.sales_script_service,
                'content_generation_service': user_service.content_generation_service,
                'summarize_service': user_service.summarize_service,
                'ppt_generation_service': user_service.ppt_generation_service,
                'blog_generation_service': user_service.blog_generation_service,
                'rephrasely_service': user_service.rephrasely_service,
                'service_start_dates': {
                    'email_service_start': user_service.email_end_date.isoformat() if user_service.email_end_date else None,
                    'offer_letter_service_start': user_service.offer_letter_end_date.isoformat() if user_service.offer_letter_end_date else None,
                    'business_proposal_service_start': user_service.business_proposal_end_date.isoformat() if user_service.business_proposal_end_date else None,
                    'sales_script_service_start': user_service.sales_script_end_date.isoformat() if user_service.sales_script_end_date else None,
                    'content_generation_service_start': user_service.content_generation_end_date.isoformat() if user_service.content_generation_end_date else None,
                    'summarize_service_start': user_service.summarize_end_date.isoformat() if user_service.summarize_end_date else None,
                    'ppt_generation_service_start': user_service.ppt_generation_end_date.isoformat() if user_service.ppt_generation_end_date else None,
                    'blog_generation_service_start': user_service.blog_generation_end_date.isoformat() if user_service.blog_generation_end_date else None,
                    'rephrasely_service_start': user_service.rephrasely_end_date.isoformat() if user_service.rephrasely_end_date else None,
                },
                'service_end_dates': {
                    'email_service_end': user_service.email_end_date.isoformat() if user_service.email_end_date else None,
                    'offer_letter_service_end': user_service.offer_letter_end_date.isoformat() if user_service.offer_letter_end_date else None,
                    'business_proposal_service_end': user_service.business_proposal_end_date.isoformat() if user_service.business_proposal_end_date else None,
                    'sales_script_service_end': user_service.sales_script_end_date.isoformat() if user_service.sales_script_end_date else None,
                    'content_generation_service_end': user_service.content_generation_end_date.isoformat() if user_service.content_generation_end_date else None,
                    'summarize_service_end': user_service.summarize_end_date.isoformat() if user_service.summarize_end_date else None,
                    'ppt_generation_service_end': user_service.ppt_generation_end_date.isoformat() if user_service.ppt_generation_end_date else None,
                    'blog_generation_service_end': user_service.blog_generation_end_date.isoformat() if user_service.blog_generation_end_date else None,
                    'rephrasely_service_end': user_service.rephrasely_end_date.isoformat() if user_service.rephrasely_end_date else None,
                },
            },
            'payments': payment_info,
        }

        return JsonResponse(response_data, status=200)

    except ObjectDoesNotExist:
        return JsonResponse({'error': 'User or user services not found.'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def email_generator_android(request):
    if request.method == 'POST':
        try:
            # Extract incoming payload
            content = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Received content: {content}')

            if not content:
                logger.warning('No content found in the request.')
                return JsonResponse({'error': 'No content found in the request.'}, status=status.HTTP_400_BAD_REQUEST)

            # Extract data from the request content
            purpose = content.get('purpose')
            if purpose == 'Other':
                purpose = content.get('otherPurpose')
            num_words = content.get('num_words')
            subject = content.get('subject')
            rephrase = content.get('rephraseSubject', False)
            to = content.get('to')
            tone = content.get('tone')
            keywords = content.get('keywords', [])
            contextual_background = content.get('contextualBackground')
            call_to_action = content.get('callToAction')
            if call_to_action == 'Other':
                call_to_action = content.get('otherCallToAction')
            additional_details = content.get('additionalDetails')
            priority_level = content.get('priorityLevel')
            closing_remarks = content.get('closingRemarks')

            logger.info(f'Generating email with the following data: {content}')

            # Generate email content
            generated_content = generate_email(
                purpose, num_words, subject, rephrase, to, tone, keywords,
                contextual_background, call_to_action, additional_details,
                priority_level, closing_remarks
            )

            if generated_content:
                logger.info('Email content generated successfully.')
                return JsonResponse({'generated_content': generated_content})
            else:
                logger.error('Failed to generate email content.')
                return JsonResponse({'error': 'Failed to generate email content.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            logger.error(f'Error processing request: {e}')
            return JsonResponse({'error': 'An error occurred while processing the request.'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
def sales_script_generator_android(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Extract the incoming payload
        data = json.loads(body)
        
        # Extract fields from the JSON data
        company_details = data.get('company_details', '')
        product_descriptions = data.get('product_descriptions', '')
        features_benefits = data.get('features_benefits', '')
        pricing_info = data.get('pricing_info', '')
        promotions = data.get('promotions', '')
        target_audience = data.get('target_audience', '')
        sales_objectives = data.get('sales_objectives', '')
        competitive_advantage = data.get('competitive_advantage', '')
        compliance = data.get('compliance', '')
        num_words = data.get('num_words', '')


        logger.debug(f"Data extracted for sales script generation: num_words={num_words}, company_details={company_details}")

        # Generate the sales script
        logger.info("Generating sales script...")
        sales_script = generate_sales_script(
            company_details,
            num_words,
            product_descriptions,
            features_benefits,
            pricing_info,
            promotions,
            target_audience,
            sales_objectives,
            competitive_advantage,
            compliance,
        )

        if sales_script:
            logger.info("Sales script generated successfully.")
            return JsonResponse({'generated_content': sales_script}, status=200)

        logger.error("Failed to generate sales script.")
        return JsonResponse({'error': 'Failed to generate sales script. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
def content_generator_android(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Parse the incoming JSON payload
        data = json.loads(body)

        # Extract fields from the JSON data
        company_info = data.get('company_info')
        content_purpose = data.get('content_purpose')
        desired_action = data.get('desired_action')
        topic_details = data.get('topic_details')
        keywords = data.get('keywords')
        audience_profile = data.get('audience_profile')
        format_structure = data.get('format_structure')
        num_words = data.get('num_words')
        seo_keywords = data.get('seo_keywords')
        references = data.get('references')

        logger.debug(f"Data extracted for content generation: company_info={company_info}, content_purpose={content_purpose}, desired_action={desired_action}")

        # Generate the content
        logger.info("Generating content...")
        content = generate_content(
            company_info,
            content_purpose,
            desired_action,
            topic_details,
            keywords,
            audience_profile,
            format_structure,
            num_words,
            seo_keywords,
            references
        )

        if content:
            logger.info("Content generated successfully.")
            return JsonResponse({'generated_content': content}, status=200)

        logger.error("Failed to generate content.")
        return JsonResponse({'error': 'Failed to generate content. Please try again.'}, status=500)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

    logger.error("Method not allowed.")
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
def generate_blog_view_android(request):
    try:
        # Load and decode the request body
        body = request.body.decode('utf-8')
        logger.debug(f"Request body received: {body}")

        # Parse the incoming payload
        data = json.loads(body)

        # Extract the required fields
        title = data.get('title')
        tone = data.get('tone')
        keywords = data.get('keywords', None)  # Optional

        # Ensure required fields are present
        if not title or not tone:
            return JsonResponse({"error": "Missing 'title' or 'tone'."}, status=400)

        # Call the generate_blog function
        blog_content = generate_blog(title, tone, keywords)

        # Return the blog content as a JSON response
        return JsonResponse({'blog_content': blog_content}, status=200)

    except json.JSONDecodeError:
        logger.error("Invalid JSON format received.")
        return JsonResponse({"error": "Invalid JSON format. Please provide valid JSON data."}, status=400)
    except ValueError as e:
        logger.error(f"ValueError occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=400)
    except Exception as e:
        logger.error(f"An unexpected error occurred: {str(e)}")
        return JsonResponse({"error": str(e)}, status=500)

    # If not a POST request, return an error
    return JsonResponse({"error": "Only POST method is allowed."}, status=405)

# @csrf_exempt
# def create_presentation_android(request):
#     try:
#         # Handle the multipart form data
#         content = request.POST.get('content')
#         if not content:
#             return JsonResponse({'error': 'No content found in the request.'}, status=400)

#         # Parse JSON data
#         data = json.loads(content)

#         # Extract fields from the data
#         title = data.get('title')
#         num_slides = data.get('num_slides')
#         bg_image_path = request.FILES.get('background_image')  # bg_image as a file
#         document = request.FILES.get('document')  # document as a file

#         print(f"Title: {title}, Number of Slides: {num_slides}, Background Image: {bg_image_path}, Document: {document}")

#         if not title or not num_slides:
#             return JsonResponse({'error': 'Title and number of slides are required.'}, status=400)

#         # Handle document content optionally
#         document_content = extract_document_content(document) if document else ""

#         # Generate presentation logic
#         prs = Presentation()
#         slide_titles = generate_slide_titles(document_content, num_slides, None, title)
#         slide_titles = slide_titles.replace('[', '').replace(']', '').replace('"', '').split(',')

#         slide_contents = {}
#         error_messages = []

#         # Function to generate slide content in a separate thread
#         def generate_and_store_slide_content(slide_title):
#             try:
#                 content = generate_slide_content(document_content, slide_title, None).replace("*", '').split('\n')
#                 current_content = [point.strip() for point in content if len(point.strip()) > 0]
#                 if len(current_content) > 4:
#                     current_content = current_content[:4]  # Limit to only 4 points
#                 slide_contents[slide_title] = current_content
#             except Exception as e:
#                 error_messages.append(f"Error generating content for '{slide_title}': {str(e)}")

#         # Start threads for generating slide content
#         threads = []
#         for st in slide_titles:
#             thread = Thread(target=generate_and_store_slide_content, args=(st.strip(),))
#             thread.start()
#             threads.append(thread)

#         # Wait for all threads to finish
#         for thread in threads:
#             thread.join()

#         # Check for any errors that occurred during content generation
#         if error_messages:
#             return JsonResponse({'error': error_messages}, status=500)

#         # Add slides to the presentation
#         for slide_title, slide_content in slide_contents.items():
#             add_slide(prs, slide_title, slide_content, bg_image_path)

#         # Save presentation to a BytesIO object
#         buffer = BytesIO()
#         prs.save(buffer)
#         buffer.seek(0)  # Rewind the buffer

#         # Return file response
#         response = FileResponse(buffer, as_attachment=True, filename='SmartOffice_Assistant_Presentation.pptx')
#         return response

#     except json.JSONDecodeError:
#         return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
#     except Exception as e:
#         return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
@require_GET
def get_user_services_android(request, email):
    if request.method == "GET":
        try:
            email = email.lower()
            user = get_object_or_404(User, email=email)
            user_services = get_object_or_404(UserService, user=user)

            services = {
                "email_service": {
                    "id": 1,
                    "is_active": user_services.email_service and (user_services.email_end_date >= date.today()),
                    "end_date": user_services.email_end_date.isoformat() if user_services.email_end_date else None,
                },
                "offer_letter_service": {
                    "id": 2,
                    "is_active": user_services.offer_letter_service and (user_services.offer_letter_end_date >= date.today()),
                    "end_date": user_services.offer_letter_end_date.isoformat() if user_services.offer_letter_end_date else None,
                },
                "business_proposal_service": {
                    "id": 3,
                    "is_active": user_services.business_proposal_service and (user_services.business_proposal_end_date >= date.today()),
                    "end_date": user_services.business_proposal_end_date.isoformat() if user_services.business_proposal_end_date else None,
                },
                "sales_script_service": {
                    "id": 4,
                    "is_active": user_services.sales_script_service and (user_services.sales_script_end_date >= date.today()),
                    "end_date": user_services.sales_script_end_date.isoformat() if user_services.sales_script_end_date else None,
                },
                "content_generation_service": {
                    "id": 5,
                    "is_active": user_services.content_generation_service and (user_services.content_generation_end_date >= date.today()),
                    "end_date": user_services.content_generation_end_date.isoformat() if user_services.content_generation_end_date else None,
                },
                "summarize_service": {
                    "id": 6,
                    "is_active": user_services.summarize_service and (user_services.summarize_end_date >= date.today()),
                    "end_date": user_services.summarize_end_date.isoformat() if user_services.summarize_end_date else None,
                },
                "ppt_generation_service": {
                    "id": 7,
                    "is_active": user_services.ppt_generation_service and (user_services.ppt_generation_end_date >= date.today()),
                    "end_date": user_services.ppt_generation_end_date.isoformat() if user_services.ppt_generation_end_date else None,
                },
                "blog_generation_service": {
                    "id": 9,
                    "is_active": user_services.blog_generation_service and (user_services.blog_generation_end_date >= date.today()),
                    "end_date": user_services.blog_generation_end_date.isoformat() if user_services.blog_generation_end_date else None,
                },
                "rephrasely_service": {
                    "id": 10,
                    "is_active": user_services.rephrasely_service and (user_services.rephrasely_end_date >= date.today()),
                    "end_date": user_services.rephrasely_end_date.isoformat() if user_services.rephrasely_end_date else None,
                },
            }

            return JsonResponse({"user_id": user.id, "services": services}, status=200)

        except User.DoesNotExist:
            return JsonResponse({"error": "User not found"}, status=404)
        except UserService.DoesNotExist:
            return JsonResponse({"error": "User services not found"}, status=404)

    return JsonResponse({"error": "Invalid request method"}, status=400)

@csrf_exempt
@require_POST
def rephrasely_view_android(request):
    if request.method == 'POST':
        try:
            # Extract the incoming payload
            data = json.loads(request.body.decode('utf-8'))
            logger.debug(f'Incoming content: {data}')

            # Extract required fields
            text_to_rephrase = data.get('text_to_rephrase')
            tone = data.get('tone')
            target_audience = data.get('target_audience')
            num_words = data.get('num_words', "default")  # Optional, default is "default"

            # Call the rephrasely function
            rephrased_text = rephrasely(text_to_rephrase, tone, target_audience, num_words)

            # Return the rephrased content directly
            return JsonResponse({'rephrased_text': rephrased_text}, status=200)

        except json.JSONDecodeError:
            logger.error('Invalid JSON format received.')
            return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
        except ValueError as e:
            logger.error(f'ValueError: {str(e)}')
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            logger.error(f'Exception: {str(e)}')
            return JsonResponse({'error': str(e)}, status=500)

    logger.warning('Method not allowed.')
    return JsonResponse({'error': 'Method not allowed.'}, status=405)

@csrf_exempt
def summarize_document_android(request):
    try:
        # Extract the JSON payload from request.POST
        data = json.loads(request.body)
        logger.debug(f'Received content: {data}')

        # Extract form fields from data
        document_context = data.get('documentContext')
        main_subject = data.get('mainSubject')
        summary_purpose = data.get('summaryPurpose')
        length_detail = data.get('lengthDetail')
        important_elements = data.get('importantElements')
        audience = data.get('audience')
        tone = data.get('tone')
        format_ = data.get('format')
        additional_instructions = data.get('additionalInstructions')

        # Extract the uploaded file from request.FILES
        document_file = request.FILES.get('documentFile')
        if not document_file:
            logger.warning('No document file provided.')
            return JsonResponse({'error': 'No document file provided.'}, status=400)

        # Generate summary using provided data and the uploaded document
        summary = generate_summary(
            document_context, main_subject, summary_purpose, length_detail,
            important_elements, audience, tone, format_, additional_instructions, document_file
        )

        if summary.startswith("Error:"):
            logger.error(summary)
            return JsonResponse({'error': summary}, status=500)

        logger.info('Summary generated successfully.')

        return JsonResponse({'summary': summary}, status=200)

    except json.JSONDecodeError:
        logger.error('Invalid JSON format received.')
        return JsonResponse({'error': 'Invalid JSON format. Please provide valid JSON data.'}, status=400)
    except ValueError as e:
        logger.error(f'ValueError: {str(e)}')
        return JsonResponse({'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f'Exception: {str(e)}')
        return JsonResponse({'error': str(e)}, status=500)


import zipfile
from io import BytesIO
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt

# Function to sort JSON data by keys
def sort_json_data(data):
    if isinstance(data, dict):
        return {key: sort_json_data(data[key]) if isinstance(data[key], dict) else data[key] for key in sorted(data)}
    return data



@csrf_exempt
def translate_json_files(request):
    if request.method == 'POST':
        try:
            # Get the uploaded JSON file, target languages, and translation option
            json_file = request.FILES.get('file')
            translate_to = request.POST.get('translate_to')
            translation_option = request.POST.get('translation_option')
            keys = request.POST.get('keys') if translation_option == 'specific' else None
            if not json_file:
                return JsonResponse({'error': 'No JSON file provided.'}, status=400)
            if not translate_to:
                return JsonResponse({'error': 'No target language provided.'}, status=400)
            # Parse target languages and original JSON content
            translate_to_list = [lang.strip() for lang in translate_to.split(',')]
            file_content = json_file.read().decode('utf-8')
            original_json = json.loads(file_content)
            # Recursive function to traverse and translate JSON
            def translate_nested_json(data, specific_keys=None):
                if isinstance(data, dict):
                    translated_dict = {}
                    for key, value in data.items():
                        if isinstance(value, str) and (specific_keys is None or key in specific_keys):
                            translated_dict[key] = value  # This value will be translated later
                        else:
                            translated_dict[key] = translate_nested_json(value, specific_keys)
                    return translated_dict
                elif isinstance(data, list):
                    return [translate_nested_json(item, specific_keys) for item in data]
                else:
                    return data

            # Function to fill in translated values
            def apply_translations(data, translations, translation_iter):
                if isinstance(data, dict):
                    for key, value in data.items():
                        if isinstance(value, str):
                            data[key] = next(translation_iter).get("translated_content", value)
                        else:
                            apply_translations(value, translations, translation_iter)
                elif isinstance(data, list):
                    for item in data:
                        apply_translations(item, translations, translation_iter)
            # Prepare zip buffer to store translated files
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_archive:
                for target_lang in translate_to_list:
                    if translation_option == 'entire':
                        # Translate entire JSON file
                        nested_to_translate = translate_nested_json(original_json)
                    elif translation_option == 'specific' and keys:
                        specific_keys = [key.strip() for key in keys.split(',')]
                        nested_to_translate = translate_nested_json(original_json, specific_keys=specific_keys)
                    else:
                        nested_to_translate = {}
                    # Extract all strings to translate
                    strings_to_translate = []
                    def collect_strings(data):
                        if isinstance(data, dict):
                            for key, value in data.items():
                                if isinstance(value, str):
                                    strings_to_translate.append(value)
                                else:
                                    collect_strings(value)
                        elif isinstance(data, list):
                            for item in data:
                                collect_strings(item)
                    collect_strings(nested_to_translate)
                    # Batch size definition
                    batch_size = 50  # Define how many strings to process per batch
                    def process_in_batches(strings, batch_size):
                        for i in range(0, len(strings), batch_size):
                            yield strings[i:i + batch_size]
                    # Translate string fields in batches using the API
                    translated_texts = []
                    for batch in process_in_batches(strings_to_translate, batch_size):
                        translated_batch = translate_multiple_texts(
                            batch,
                            from_code="English",  # You can customize this
                            to_code=target_lang,
                            user_id=BHASHINI_USER_ID,
                            api_key=BHASHINI_API_KEY
                        )
                        translated_texts.extend(translated_batch)
                    # Reapply translations to the JSON structure
                    apply_translations(nested_to_translate, translated_texts, iter(translated_texts))
                    # Sort the translated JSON by keys
                    sorted_translated_json = sort_json_data(nested_to_translate)
                    # Write the sorted translated JSON to a file in the zip archive
                    translated_json_str = json.dumps(sorted_translated_json, ensure_ascii=False, indent=4)
                    translated_file_name = f"translated_{target_lang}_sorted.json"
                    zip_archive.writestr(translated_file_name, translated_json_str)
            # Return the zip file containing translated JSON files
            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="translated_sorted_files.zip"'
            return response
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON file format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f"Error during translation: {str(e)}"}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)


import asyncio
semaphore = asyncio.Semaphore(10)  
from .email_llama3 import bhashini_translate


import json
import asyncio
import zipfile
from io import BytesIO
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt

# Initialize line counter
global line_number
line_number = 1

indian_languages = {
    "English": "en",
    "Hindi": "hi",
    "Tamil": "ta",
    "Telugu": "te",
    "Marathi": "mr",
    "Kannada": "kn",
    "Bengali": "bn",
    "Odia": "or",
    "Assamese": "as",
    "Punjabi": "pa",
    "Malayalam": "ml",
    "Gujarati": "gu",
    "Urdu": "ur",
    "Sanskrit": "sa",
    "Nepali": "ne",
    "Bodo": "brx",
    "Maithili": "mai",
    "Sindhi": "sd",
    "Kashmiri": "ks",
    "Konkani": "kok",
    "Dogri": "doi",
    "Goan Konkani": "gom",
    "Santali": "sat"
}


@csrf_exempt
def translate_json_files_new(request):
    global line_number
    translated_json = {}
    if request.method == 'POST':
        try:
            json_file = request.FILES.get('file')
            translate_to = request.POST.get('translate_to')

            if not json_file:
                return JsonResponse({'error': 'No JSON file provided.'}, status=400)

            if not translate_to:
                return JsonResponse({'error': 'No target language provided.'}, status=400)

            file_content = json_file.read().decode('utf-8')
            original_json = json.loads(file_content)

            translation_tasks = [(key, value) for key, value in original_json.items() if isinstance(value, str)]
            translated_json = {key: value for key, value in original_json.items() if not isinstance(value, str)}

            async def translate_key_value(key, value, target_lang):
                global line_number
                try:
                    print(f"Line {line_number}: Translating key '{key}' with value '{value}'")
                    translation_result = bhashini_translate(value, target_lang)
                    translated_json[key] = translation_result["translated_content"]
                    print(f"Line {line_number}: Translated value '{translated_json[key]}'")
                    line_number += 1
                except Exception as e:
                    print(f"Line {line_number}: Error translating key '{key}' - {str(e)}")
                    translated_json[key] = f"Translation Error: {str(e)}"

            async def trans_main(translation_tasks, translate_to):
                tasks = [translate_key_value(key, value, translate_to) for key, value in translation_tasks]
                await asyncio.gather(*tasks)

            asyncio.run(trans_main(translation_tasks, translate_to))

            # Sort the keys alphabetically
            translated_json = dict(sorted(translated_json.items()))

            translated_json_str = json.dumps(translated_json, ensure_ascii=False, indent=4)
            translated_file_name = f"translated_{translate_to}.json"
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_archive:
                zip_archive.writestr(translated_file_name, translated_json_str)

            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="translated_sorted_files.zip"'
            return response

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON file format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Error during translation: {str(e)}'}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)

@csrf_exempt
def fix_null_values_in_translation(request):
    global line_number
    if request.method == 'POST':
        try:
            english_file = request.FILES.get('english_file')
            translated_file = request.FILES.get('translated_file')

            if not english_file or not translated_file:
                return JsonResponse({'error': 'Both English and Translated JSON files are required.'}, status=400)

            english_content = english_file.read().decode('utf-8')
            translated_content = translated_file.read().decode('utf-8')

            english_json = json.loads(english_content)
            translated_json = json.loads(translated_content)

            for key, value in translated_json.items():
                if value is None and key in english_json:
                    try:
                        print(f"Line {line_number}: Found null value for key '{key}'. Translating value '{english_json[key]}'")
                        translation_result = bhashini_translate(english_json[key], request.POST.get('translate_to'))
                        translated_json[key] = translation_result["translated_content"]
                        print(f"Line {line_number}: Translated value '{translated_json[key]}'")
                        line_number += 1
                    except Exception as e:
                        print(f"Line {line_number}: Error translating key '{key}' - {str(e)}")
                        translated_json[key] = f"Translation Error: {str(e)}"

            fixed_json_str = json.dumps(translated_json, ensure_ascii=False, indent=4)
            fixed_file_name = "fixed_translated.json"

            response = HttpResponse(fixed_json_str, content_type='application/json')
            response['Content-Disposition'] = f'attachment; filename="{fixed_file_name}"'
            return response

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON file format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Error during fixing null values: {str(e)}'}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)


@csrf_exempt
def translate_json_from_english_to_other(request):
    global line_number
    translated_json = {}
    if request.method == 'POST':
        try:
            json_file = request.FILES.get('file')
            source_language = request.POST.get('source_language')
            translate_to = request.POST.get('translate_to')

            if not json_file:
                return JsonResponse({'error': 'No JSON file provided.'}, status=400)

            if not source_language:
                return JsonResponse({'error': 'No source language provided.'}, status=400)

            if not translate_to:
                return JsonResponse({'error': 'No target language provided.'}, status=400)

            file_content = json_file.read().decode('utf-8')
            original_json = json.loads(file_content)

            translation_tasks = [(key, value) for key, value in original_json.items() if isinstance(value, str)]
            translated_json = {key: value for key, value in original_json.items() if not isinstance(value, str)}

            async def translate_key_value(key, value, source_lang, target_lang):
                global line_number
                try:
                    print(f"Line {line_number}: Translating key '{key}' with value '{value}'")
                    translation_result = bhashini_translate(value, target_lang, source_lang)
                    translated_json[key] = translation_result["translated_content"]
                    print(f"Line {line_number}: Translated value '{translated_json[key]}'")
                    line_number += 1
                except Exception as e:
                    print(f"Line {line_number}: Error translating key '{key}' - {str(e)}")
                    translated_json[key] = f"Translation Error: {str(e)}"

            async def trans_main(translation_tasks, source_lang, target_lang):
                tasks = [translate_key_value(key, value, source_lang, target_lang) for key, value in translation_tasks]
                await asyncio.gather(*tasks)

            asyncio.run(trans_main(translation_tasks, source_language, translate_to))

            # Sort the keys alphabetically
            translated_json = dict(sorted(translated_json.items()))

            translated_json_str = json.dumps(translated_json, ensure_ascii=False, indent=4)
            translated_file_name = f"translated_{translate_to}.json"
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_archive:
                zip_archive.writestr(translated_file_name, translated_json_str)

            zip_buffer.seek(0)
            response = HttpResponse(zip_buffer, content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="translated_sorted_files.zip"'
            return response

        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON file format.'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Error during translation: {str(e)}'}, status=500)
    else:
        return JsonResponse({'error': 'Invalid request method'}, status=400)
    

from django.http import JsonResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
import zipfile
import asyncio
from django.http import JsonResponse, FileResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from docx import Document
import zipfile
import asyncio

@csrf_exempt
def translate_and_download_document(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed.'}, status=405)

    try:
        # Extract the uploaded file and target languages from the request
        document = request.FILES.get('document')
        languages = request.POST.get('languages')

        if not document or not languages:
            return JsonResponse({'error': 'Both document and target languages are required.'}, status=400)

        # Parse languages into a list
        target_languages = [lang.strip() for lang in languages.split(',')]

        # Set up a temporary directory and zip file to store translated files
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, "translated_documents.zip")

        with zipfile.ZipFile(zip_path, 'w') as zip_file:
            for lang in target_languages:
                # Reset file pointer and read content based on file type
                document.seek(0)

                if document.name.endswith('.docx'):
                    # Load the .docx document
                    doc = Document(document)
                    translated_doc = Document()  # New doc for translated content

                    # Translate each paragraph
                    for para in doc.paragraphs:
                        if para.text.strip():  # Avoid translating empty paragraphs
                            response = bhashini_translate(para.text, to_code=lang)
                            if response["status_code"] != 200:
                                return JsonResponse({'error': f'Translation failed for {lang}.'}, status=500)
                            translated_text = response["translated_content"]
                            translated_doc.add_paragraph(translated_text)
                        else:
                            translated_doc.add_paragraph('')  # Keep empty paragraphs

                    # Save translated document to temporary directory
                    filename = f'translated_{lang}.docx'
                    file_path = os.path.join(temp_dir, filename)
                    translated_doc.save(file_path)
                    zip_file.write(file_path, filename)

                else:
                    return JsonResponse({'error': 'Only .docx files are supported.'}, status=400)

        # Return the zip file as a response
        with open(zip_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="translated_documents.zip"'
            return response

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)